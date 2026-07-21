"""
Finova / BitWealth — Client Trading Mandate Addendum
Annexure [●]: BitWealth Bitcoin LTH PVR Strategy
Output: docs/FSCA Compliance/BitWealth_LTH_PVR_Client_Addendum_v1.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("docs") / "FSCA Compliance"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Colours (matched to Finova's Annexure example) ───────────────────────────
BLUE_DARK  = RGBColor(0x17, 0x3F, 0x5F)   # deep navy title text
BLUE_MED   = RGBColor(0x1F, 0x5C, 0x8B)   # section heading blue
DARK       = RGBColor(0x1A, 0x1A, 0x1A)
GREY       = RGBColor(0x55, 0x55, 0x55)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

HDR_BG_HEX = "D9E2F3"   # light blue-grey — header banner background
BODY_FONT  = "Calibri"
BODY_SZ    = 10


# ══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ══════════════════════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def para(doc, text="", bold=False, italic=False, size=BODY_SZ, color=None,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=2, sa=4, underline=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.alignment = align
    if text:
        run = p.add_run(text)
        run.bold = bold; run.italic = italic; run.underline = underline
        run.font.name = BODY_FONT; run.font.size = Pt(size)
        run.font.color.rgb = color if color else DARK
    return p


def mixed_para(doc, parts, align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=2, sa=4):
    """parts = list of (text, bold, italic). Returns paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    p.alignment = align
    for text, bold, italic in parts:
        r = p.add_run(text)
        r.bold = bold; r.italic = italic
        r.font.name = BODY_FONT; r.font.size = Pt(BODY_SZ); r.font.color.rgb = DARK
    return p


def bullet(doc, text, bold_lead=None, indent=Cm(0.8)):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = indent
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead:
        rb = p.add_run(bold_lead)
        rb.bold = True; rb.font.name = BODY_FONT
        rb.font.size = Pt(BODY_SZ); rb.font.color.rgb = DARK
    r = p.add_run(text)
    r.font.name = BODY_FONT; r.font.size = Pt(BODY_SZ); r.font.color.rgb = DARK
    return p


def section_head(doc, text, sb=8, sa=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(sa)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True; run.font.name = BODY_FONT; run.font.size = Pt(BODY_SZ)
    run.font.color.rgb = DARK
    return p


def hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), "173F5F")
    pBdr.append(b); pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles["Normal"]
style.font.name = BODY_FONT; style.font.size = Pt(BODY_SZ)
for section in doc.sections:
    section.top_margin = section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.2)

# ── Title banner (matches the blue-grey box in Finova's example) ─────────────
tbl = doc.add_table(rows=1, cols=1)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.columns[0].width = Cm(16.6)
c = tbl.cell(0, 0)
set_cell_bg(c, HDR_BG_HEX)
c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
r = c.paragraphs[0].add_run("Annexure [\u25cf]: BitWealth Bitcoin LTH PVR Strategy")
r.bold = False; r.italic = True; r.font.name = BODY_FONT
r.font.size = Pt(12); r.font.color.rgb = BLUE_DARK
c.paragraphs[0].paragraph_format.space_before = Pt(4)
c.paragraphs[0].paragraph_format.space_after  = Pt(4)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ── Parties ──────────────────────────────────────────────────────────────────
para(doc, 'entered into between Finova (PTY) Ltd ("Finova") and',
     align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=2)
para(doc, "\u2026" * 48, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=0)
para(doc, '("The Client")', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=8)

# ── Preamble ─────────────────────────────────────────────────────────────────
para(doc,
    "BitWealth Asset Managers (Pty) Ltd (\u201cBitWealth\u201d), a Juristic Representative of "
    "Finova (PTY) Ltd (FSP No. 21095), offers a proprietary Bitcoin investment management "
    "service known as the BitWealth Bitcoin LTH PVR Strategy (\u201cthe Strategy\u201d). This Strategy "
    "operates as a fully rules-based, algorithm-driven Bitcoin accumulation and risk-management "
    "service executed on the VALR exchange within a segregated client sub-account. No "
    "cross-border transfer of funds or forex externalization is required.", sb=0, sa=6)

# ── Preamble ─────────────────────────────────────────────────────────────────
para(doc,
    "BitWealth Asset Managers (Pty) Ltd (\u201cBitWealth\u201d), a Juristic Representative of "
    "Finova (PTY) Ltd (FSP No. 21095), offers a proprietary Bitcoin investment management "
    "service known as the BitWealth Bitcoin LTH PVR Strategy (\u201cthe Strategy\u201d). The Strategy "
    "combines three core components \u2014 an on-chain signal layer, a dollar-cost averaging "
    "accumulation framework, and a dynamic grid execution engine \u2014 into a single, fully-automated "
    "Bitcoin allocation engine. An optional fourth component (idle-cash yield via USDPC) is "
    "available and must be explicitly opted into. The Strategy is executed entirely on the VALR "
    "exchange within a segregated, dedicated client sub-account. No cross-border transfer of "
    "funds or forex externalization is required.", sb=0, sa=6)

# ── The Strategy ─────────────────────────────────────────────────────────────
section_head(doc, "The Strategy \u2014 How it Works")

section_head(doc, "Component 1: The LTH PVR On-Chain Signal", sb=4)
para(doc,
    "The core signal of the Strategy is the Long-Term Holder Profit-to-Volatility Ratio "
    "(\u201cLTH PVR\u201d), a proprietary on-chain analytical metric. It measures the deviation between "
    "the current Bitcoin market price and the average acquisition cost (realised price) of "
    "Bitcoin\u2019s long-term holders \u2014 defined as holders who have not moved their coins for "
    "155 days or more. This cohort, often called the \u201csmart money\u201d of the Bitcoin network, "
    "represents over 70% of circulating supply and reflects the conviction of the most "
    "committed participants in the market.")
para(doc,
    "The deviation of price from the LTH realised price is normalised by the rolling "
    "volatility of that cohort\u2019s cost basis and expressed in standard deviations (\u03c3) from "
    "its historical mean. A negative \u03c3 reading indicates Bitcoin is statistically cheap "
    "relative to what long-term holders paid; a positive \u03c3 reading indicates it is expensive. "
    "Because LTH realised price is computed directly from blockchain UTXO data, the signal "
    "reflects actual economic conviction \u2014 it cannot be manipulated by short-term flows, "
    "derivative positioning, or sentiment.")

section_head(doc, "Component 2: Dollar-Cost Averaging (DCA)", sb=4)
para(doc,
    "The Strategy commits the client\u2019s regular contributions to Bitcoin at a disciplined "
    "cadence regardless of short-term market conditions. By spreading entries over time, DCA "
    "eliminates the single largest risk in volatile assets: the timing of a lump-sum entry. "
    "Average cost converges to the mean traded price over the contribution window, not to "
    "whatever price prevailed on any one day. The accumulation is fully automated; the client "
    "does not need to monitor the market or make day-to-day decisions.")

section_head(doc, "Component 3: Dynamic Grid Execution", sb=4)
para(doc,
    "The LTH PVR signal feeds a dynamic grid engine that determines how much Bitcoin to buy "
    "or sell on any given day. The Strategy uses an eleven-band confidence-interval framework: "
    "eleven grid levels are drawn as \u03c3 bands around the daily LTH PVR mean, with buy orders "
    "becoming progressively larger as the signal moves deeper into statistically cheap territory "
    "(\u221e2\u03c3 or lower) and sell orders becoming progressively larger as it moves into overheated "
    "territory (+2\u03c3 or higher). Unlike a traditional static grid bot \u2014 whose fixed price levels "
    "quickly become stranded in a trending market \u2014 the sigma bands recalibrate automatically "
    "every day as the on-chain mean drifts, ensuring orders always fire relative to current "
    "market conditions.")
para(doc,
    "The grid also includes a Bear Market Pause: when momentum indicators confirm a sustained "
    "downtrend, the engine suspends new buy orders until conditions stabilise, protecting "
    "against deploying capital into a persistent decline.")
para(doc,
    "All orders are placed as LIMIT orders on the VALR exchange. If a limit order is not "
    "filled within 5 minutes, it is automatically cancelled and replaced with a MARKET order "
    "to ensure execution. The client\u2019s assets are held at all times in a segregated, dedicated "
    "VALR sub-account; BitWealth does not receive or hold client funds.")
para(doc,
    "The Strategy may be funded by a once-off lump-sum of at least R\u202950\u202f000 and / or a "
    "recurring monthly DCA contribution of at least R\u20095\u202f000 per month. Withdrawals may be "
    "requested at any time with 5 (five) business days\u2019 written notice.")

# ── Automated nature + discretion clause ─────────────────────────────────────
section_head(doc, "Automated Execution and Reserved Discretion")
para(doc,
    "The Strategy is fully automated. Once the Client\u2019s account is configured and funded, "
    "signals are generated and orders placed without manual intervention. However, BitWealth "
    "expressly reserves discretion to override, pause, or modify automated execution in the "
    "following exceptional circumstance:")
bullet(doc,
    "MARKET STRUCTURE CHANGE \u2014 a fundamental, structural change in the Bitcoin market, "
    "on-chain data architecture, or exchange infrastructure that renders the LTH PVR signals "
    "unreliable or inapplicable (for example, a protocol-level hard fork that materially alters "
    "the on-chain data on which the algorithm depends).",
    bold_lead="")
para(doc,
    "Any exercise of discretion will be: (a) documented and motivated in writing; "
    "(b) limited to the minimum action necessary to protect the Client\u2019s interests; and "
    "(c) notified to the Client within 2 (two) business days of the action being taken. "
    "Outside of the above exception, the Strategy will not deviate from its rules-based signals.")

# ── USDPC Idle-Cash Yield (Optional Component 4) ─────────────────────────────
section_head(doc, "Optional Component 4: Idle-Cash Yield via USDPC")
para(doc,
    "Between Bitcoin positions \u2014 during HOLD periods and the Bear Market Pause \u2014 the client\u2019s "
    "portfolio may hold a material balance of idle USDT. By default this cash earns no return. "
    "BitWealth offers an optional enhancement: idle USDT can be automatically swept into "
    "USDPC, the USD Private Credit Token issued by RainFin (Pty) Ltd, backed by the "
    "Garrington Private Credit Strategy \u2014 a portfolio of senior secured, asset-backed loans "
    "with a 10-year track record.")
para(doc,
    "USDPC targets a net USD return of approximately 8\u201310% per annum. The sweep is fully "
    "automated and designed never to delay a Bitcoin trade: the instant the LTH PVR engine "
    "generates a buy signal, any USDPC balance is converted back to USDT (sub-second "
    "settlement) before the order is placed. Each conversion costs approximately 0.1% of "
    "the amount swept. This feature is OPTIONAL and disabled by default.")

tbl_usdpc = doc.add_table(rows=1, cols=2)
tbl_usdpc.style = "Table Grid"
tbl_usdpc.alignment = WD_TABLE_ALIGNMENT.LEFT
tbl_usdpc.columns[0].width = Cm(1.0); tbl_usdpc.columns[1].width = Cm(15.0)
set_cell_bg(tbl_usdpc.cell(0, 0), "FFFDE7")
c_box = tbl_usdpc.cell(0, 0)
c_box.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
rb = c_box.paragraphs[0].add_run("\u2610")
rb.font.name = BODY_FONT; rb.font.size = Pt(16); rb.font.color.rgb = BLUE_DARK
c_txt = tbl_usdpc.cell(0, 1); set_cell_bg(c_txt, "FFFDE7"); c_txt.paragraphs[0].clear()
rl = c_txt.paragraphs[0].add_run(
    "I / We opt IN to the USDPC idle-cash yield feature (Optional Component 4). "
    "I / We acknowledge the USDPC-specific risks listed below and agree that idle USDT in "
    "my / our portfolio may be automatically swept into USDPC between Bitcoin positions.   "
    "Client initials: __________")
rl.font.name = BODY_FONT; rl.font.size = Pt(BODY_SZ); rl.font.color.rgb = DARK

para(doc, "USDPC-specific risks (applicable only if opted in above):", sb=6, sa=2)
bullet(doc,
    "USDPC is NOT a bank deposit and is not protected by the Deposit Insurance Scheme. "
    "The 8\u201310% per annum target return is not guaranteed and may vary.")
bullet(doc,
    "The client is exposed to credit risk (borrower default within the Garrington "
    "Private Credit portfolio), liquidity risk (redemption delays in certain conditions), "
    "and smart-contract / technology risk associated with the USDPC token.")
bullet(doc,
    "USDPC is a third-party product issued by RainFin (Pty) Ltd. BitWealth and Finova "
    "are not responsible for the performance, solvency, or continued operation of RainFin "
    "or the Garrington Private Credit Strategy.")
bullet(doc,
    "Each automated USDT\u21d4USDPC conversion costs approximately 0.1% of the amount swept. "
    "These costs are for the client\u2019s account.")

# ── Risks ─────────────────────────────────────────────────────────────────────
section_head(doc, "Risks")
para(doc,
    "Bitcoin is a high-risk, highly volatile asset. Past or back-tested performance of the "
    "LTH PVR Strategy is not a guarantee of future results. The client acknowledges the "
    "following risks, which are not exhaustive:")
bullet(doc,
    "Bitcoin price may decline sharply and rapidly. The client may receive back less "
    "than the amount invested.")
bullet(doc,
    "The LTH PVR algorithm uses on-chain data sourced from Research Bitcoin, a third-party "
    "data provider. Inaccurate, delayed, or unavailable data may cause incorrect signals or "
    "missed trades.")
bullet(doc,
    "All trades are executed on the VALR exchange. The client is exposed to exchange "
    "operational risk, including downtime, technical failure, and counterparty risk.")
bullet(doc,
    "Bitcoin is a relatively new asset class. Regulatory frameworks governing crypto assets "
    "are evolving. Future changes may affect the Strategy, VALR, or the client\u2019s access to "
    "their assets.")
bullet(doc,
    "Client assets are held in a VALR sub-account. VALR is a licensed crypto asset service "
    "provider, but assets held on an exchange are not covered by the Investor Protection "
    "Compensation Scheme applicable to traditional securities.")
bullet(doc,
    "The Dynamic Grid and Bear Market Pause features reduce but do not eliminate the risk "
    "of loss during sustained bear markets or extreme volatility events.")
bullet(doc,
    "No warranties, guarantees or representations regarding returns are made. The client "
    "should consider the appropriateness of this Strategy for their personal financial "
    "circumstances before investing.")

# ── Fees ─────────────────────────────────────────────────────────────────────
section_head(doc, "Fees")
para(doc,
    "The following fees apply. Applicable rates are agreed between the client and BitWealth "
    "and recorded in the client\u2019s individual fee schedule.")
mixed_para(doc, [
    ("Performance Fee with High-Water Mark (HWM): ", True, False),
    ("[\u25cf]% of net profits above the HWM, calculated and deducted ", False, False),
    ("quarterly", False, True),
    (". Standard rate: 10%.", False, False),
])
para(doc,
    "The HWM is the highest NAV the client\u2019s portfolio has previously achieved. Performance "
    "fees are charged only on new profits above that mark. If the portfolio falls and "
    "recovers, no fee is charged on the recovery \u2014 only on gains that set a new high. "
    "Example: NAV reaches $100k, drops to $80k, recovers to $100k (no fee), then grows to "
    "$120k \u2014 fee charged on $20k of new profit only.", sb=0, sa=4)
mixed_para(doc, [
    ("Management Fee: ", True, False),
    ("[\u25cf]% per annum on portfolio NAV, calculated and deducted monthly. "
     "Standard rate: 1% p.a. (\u224820.83 basis points per month on average NAV).", False, False),
])
para(doc,
    "Exchange fees (VALR): approximately 8 basis points (0.08%) per BTC/USDT trade "
    "and 18 basis points (0.18%) per ZAR/USDT conversion. These are for the client\u2019s "
    "account. BitWealth receives a revenue-share from VALR which does not result in any "
    "additional charge to the client.", sb=0)
mixed_para(doc, [
    ("USDPC conversion fee (if opted in only): ", True, False),
    ("approximately 0.1% per automated USDT\u21d4USDPC sweep. Charged to the client\u2019s "
     "account on each conversion.", False, False),
])
para(doc, "No platform fee on contributions. No withdrawal fees (beyond exchange costs). "
          "No inactivity fees. Full fee transparency in the BitWealth client portal.", sb=0)

# ── Authorization and Power of Attorney ──────────────────────────────────────
section_head(doc, "Authorization and Power of Attorney")
para(doc,
    "By signing below, the Client authorises and instructs Finova and BitWealth to act on the "
    "Client\u2019s behalf in connection with the Strategy for so long as this Addendum remains in "
    "force:")
bullet(doc,
    "The Client authorises BitWealth and / or Finova to open and administer a dedicated "
    "sub-account on the VALR exchange in the Client\u2019s name, including completing all "
    "required FICA / KYC documentation on the Client\u2019s behalf to give effect to the above.")
bullet(doc,
    "The Client authorises BitWealth to execute buy and sell orders in Bitcoin on the "
    "VALR exchange on the Client\u2019s behalf, in accordance with the LTH PVR algorithm "
    "signals, without requiring a separate instruction for each individual trade.")
bullet(doc,
    "The Client authorises BitWealth to move funds between the Client\u2019s ZAR (cash) "
    "wallet and Bitcoin wallet on VALR as required by the Strategy, within the limits "
    "of the Client\u2019s funded balance.")
bullet(doc,
    "Where the Client has opted into Optional Component 4, the Client authorises BitWealth "
    "to automatically sweep idle USDT into USDPC and convert it back to USDT upon a buy "
    "signal, incurring the 0.1% conversion fee for the Client\u2019s account. If the USDPC "
    "opt-in box above has NOT been initialled, this authorisation does not apply.")
bullet(doc,
    "The Client acknowledges the fully automated and rules-based nature of the "
    "Strategy and accepts that trades will be executed by the algorithm without prior "
    "notification of each individual transaction.")
bullet(doc,
    "The Client acknowledges the High-Water Mark fee structure described above and "
    "authorises the deduction of fees from the portfolio at the intervals stated.")
bullet(doc,
    "The Client confirms that they have read and understood the risks set out in this "
    "Addendum, have considered the appropriateness of the Strategy for their personal "
    "financial circumstances, and have obtained or waived independent financial advice.")
para(doc,
    "This Addendum is incorporated into and forms part of the Finova Client Mandate "
    "Agreement signed by the Client. In the event of any conflict between this Addendum "
    "and the main Mandate Agreement, this Addendum prevails in respect of the "
    "BitWealth Bitcoin LTH PVR Strategy.", sb=6, sa=8)

hrule(doc)

# ── Signature block ───────────────────────────────────────────────────────────
sig = doc.add_table(rows=5, cols=2)
sig.style = "Table Grid"
sig.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in sig.rows:
    row.cells[0].width = Cm(8.0); row.cells[1].width = Cm(8.0)


def sig_cell(cell, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]; p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.name = BODY_FONT; r.font.size = Pt(BODY_SZ); r.font.color.rgb = DARK


sig_cell(sig.cell(0, 0), "For and on behalf of by the Client:", bold=False)
p = sig.cell(0, 0).paragraphs[0]
# Bold "the Client" part
p.clear(); r1 = p.add_run("For and on behalf of ")
r1.font.name = BODY_FONT; r1.font.size = Pt(BODY_SZ); r1.font.color.rgb = DARK
r2 = p.add_run("by the Client"); r2.bold = True
r2.font.name = BODY_FONT; r2.font.size = Pt(BODY_SZ); r2.font.color.rgb = DARK
r3 = p.add_run(":"); r3.font.name = BODY_FONT; r3.font.size = Pt(BODY_SZ); r3.font.color.rgb = DARK

p2 = sig.cell(0, 1).paragraphs[0]
p2.clear(); r1 = p2.add_run("For and on behalf of ")
r1.font.name = BODY_FONT; r1.font.size = Pt(BODY_SZ); r1.font.color.rgb = DARK
r2 = p2.add_run("Finova"); r2.bold = True
r2.font.name = BODY_FONT; r2.font.size = Pt(BODY_SZ); r2.font.color.rgb = DARK
r3 = p2.add_run(","); r3.font.name = BODY_FONT; r3.font.size = Pt(BODY_SZ); r3.font.color.rgb = DARK

sig_cell(sig.cell(1, 0), "")
sig.cell(1, 0).paragraphs[0].paragraph_format.space_before = Pt(20)
sig_cell(sig.cell(1, 1), "")
sig.cell(1, 1).paragraphs[0].paragraph_format.space_before = Pt(20)

sig_cell(sig.cell(2, 0), "\u2026" * 38)
sig_cell(sig.cell(2, 1), "\u2026" * 38)

sig_cell(sig.cell(3, 0), "the signatory warranting that he / she is duly authorised:", italic=True)
sig_cell(sig.cell(3, 1), "the signatory warranting that he / she is duly authorised:", italic=True)

sig_cell(sig.cell(4, 0), "Signed at \u2026\u2026\u2026\u2026\u2026\u2026  On date: \u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026")
sig_cell(sig.cell(4, 1), "Signed at \u2026\u2026\u2026\u2026\u2026\u2026\u2026  On date:\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026\u2026")

# ── Footer note ───────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Finova Mandate \u2013 BitWealth Bitcoin LTH PVR Strategy Addendum")
r.italic = True; r.font.name = BODY_FONT; r.font.size = Pt(8); r.font.color.rgb = GREY

out = OUT_DIR / "BitWealth_LTH_PVR_Client_Addendum_v1.docx"
doc.save(out)
print(f"Saved: {out}")
