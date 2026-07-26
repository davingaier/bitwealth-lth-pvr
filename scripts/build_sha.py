"""
BitWealth Asset Managers (Pty) Ltd — Shareholders' Agreement
Between: Davin Harald Gaier & Simon Henry Newbold Hobday & the Company
Output: docs/Shareholding/BitWealth_SHA_v1.docx

Gives binding legal effect to all terms agreed in the Equity Proposal v4.
References: Companies Act 71 of 2008 · FAIS Act 37 of 2002 · STT Act 25 of 2007

NOT legal advice — must be reviewed and executed under the supervision of a
qualified South African attorney before it creates any binding obligations.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("docs") / "Shareholding"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Parties ──────────────────────────────────────────────────────────────────
FOUNDER       = "Davin Harald Gaier"
FOUNDER_ID    = "8405025239081"
PARTNER       = "Simon Henry Newbold Hobday"
PARTNER_ID    = "6806175080088"
CO_NAME       = "BitWealth Asset Managers (Pty) Ltd"
CO_REG        = "2026/090346/07"
FSP_NAME      = "Finova (Pty) Ltd"
FSP_NO        = "21095"

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0A, 0x2A, 0x43)
DARK   = RGBColor(0x1A, 0x1A, 0x1A)
GREY   = RGBColor(0x55, 0x55, 0x55)
RED    = RGBColor(0xB7, 0x1C, 0x1C)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX = "0A2A43"
LGREY    = "F2F2F2"

BODY_FONT = "Calibri"
BODY_SZ   = 11


# ═══════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def para(doc, text="", bold=False, italic=False, size=BODY_SZ, color=None,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=2, sa=4, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.name = BODY_FONT; run.font.size = Pt(size)
        run.font.color.rgb = color if color else DARK
    return p


def clause(doc, number, text, level=0, bold_num=True):
    """Clause with hanging indent and numbered lead."""
    p = doc.add_paragraph()
    indent = level * 1.1
    p.paragraph_format.left_indent      = Cm(indent + 1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if number:
        rn = p.add_run(f"{number}\t")
        rn.bold = bold_num; rn.font.name = BODY_FONT; rn.font.size = Pt(BODY_SZ)
        rn.font.color.rgb = DARK
    rt = p.add_run(text)
    rt.font.name = BODY_FONT; rt.font.size = Pt(BODY_SZ); rt.font.color.rgb = DARK
    return p


def section(doc, number, title, sb=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True; p.paragraph_format.page_break_before = (number != 1)
    num_txt = f"{number}." if number else ""
    if number:
        r = p.add_run(f"{number}.\t{title.upper()}")
    else:
        r = p.add_run(title.upper())
    r.bold = True; r.font.name = BODY_FONT; r.font.size = Pt(13)
    r.font.color.rgb = NAVY
    return p


def hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), NAVY_HEX)
    pBdr.append(b); pPr.append(pBdr)


def defn(doc, term, definition):
    """Definition clause pair — term in bold, definition following."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent      = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rt = p.add_run(f'"{term}"\t')
    rt.bold = True; rt.font.name = BODY_FONT; rt.font.size = Pt(BODY_SZ); rt.font.color.rgb = DARK
    rd = p.add_run(definition)
    rd.font.name = BODY_FONT; rd.font.size = Pt(BODY_SZ); rd.font.color.rgb = DARK


# ═══════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles["Normal"]
style.font.name = BODY_FONT; style.font.size = Pt(BODY_SZ)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.5)

# ── Cover / Title block ───────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER; t.paragraph_format.space_before = Pt(0)
r = t.add_run("SHAREHOLDERS' AGREEMENT")
r.bold = True; r.font.size = Pt(20); r.font.color.rgb = NAVY; r.font.name = BODY_FONT

para(doc, CO_NAME, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13, sb=4, sa=2)
para(doc, f"(Registration No. {CO_REG})", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
     size=11, color=GREY, sb=0, sa=12)
hrule(doc)
para(doc, "PRIVATE AND CONFIDENTIAL", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
para(doc,
    "THIS AGREEMENT HAS NOT BEEN REVIEWED BY AN ATTORNEY. IT IS PROVIDED AS A WORKING DRAFT "
    "BASED ON THE TERMS AGREED IN THE PARTIES\u2019 SIGNED EQUITY PROPOSAL (v4). IT MUST BE "
    "REVIEWED, FINALISED AND EXECUTED UNDER THE SUPERVISION OF A QUALIFIED SOUTH AFRICAN "
    "ATTORNEY AND COMPLIANCE OFFICER BEFORE IT CREATES ANY BINDING OBLIGATIONS.",
    bold=True, size=9, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, sb=4, sa=10)
hrule(doc)

# ── Parties block ─────────────────────────────────────────────────────────────
para(doc, "ENTERED INTO BY AND BETWEEN:", bold=True, sb=10)

para(doc, f"1.\t{FOUNDER}, South African, ID {FOUNDER_ID} (\u201cDavin\u201d or the \u201cFounder\u201d);", indent=0.5)
para(doc, f"2.\t{PARTNER}, South African, ID {PARTNER_ID} (\u201cSimon\u201d or the \u201cBD Partner\u201d);", indent=0.5)
para(doc, f"3.\t{CO_NAME}, Registration No. {CO_REG}, a private company incorporated under the laws of the "
     f"Republic of South Africa (\u201cthe Company\u201d).", indent=0.5)

para(doc, "(Davin, Simon and the Company are collectively referred to as \u201cthe Parties\u201d "
     "and individually as a \u201cParty\u201d.)", italic=True, sb=0)
hrule(doc)

# ── Preamble ──────────────────────────────────────────────────────────────────
para(doc, "PREAMBLE", bold=True, size=12, color=NAVY)
for txt in [
    "A.\tThe Company is a private company that operates the BitWealth LTH PVR bitcoin DCA "
    "investment management platform and renders Category II discretionary financial services "
    "as a Juristic Representative of " + FSP_NAME + " (FSP No. " + FSP_NO + ").",
    "B.\tDavin has contributed intellectual property to the Company under a signed IP Transfer "
    "Agreement at an agreed value of R2\u202f500\u202f000 and holds all 1\u202f000 issued ordinary shares.",
    "C.\tThe Parties have agreed that Simon will join the Company as Head of Business "
    "Development (non-executive consulting capacity), receive a 10% founding equity grant, "
    "and have the opportunity to earn up to a further 30% equity through AUM milestones.",
    "D.\tThe Parties intend this arrangement to constitute a genuine equity partnership between "
    "co-owners and NOT an employee share-incentive scheme.",
    "E.\tThe Parties wish to record the terms of their relationship in this Agreement in "
    "accordance with the Companies Act, 71 of 2008.",
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(txt); r.font.name = BODY_FONT; r.font.size = Pt(BODY_SZ); r.font.color.rgb = DARK

hrule(doc)
para(doc, "THE PARTIES AGREE AS FOLLOWS:", bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# 1. DEFINITIONS AND INTERPRETATION
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 1, "Definitions and Interpretation", sb=8)

clause(doc, "1.1", "In this Agreement, unless the context indicates otherwise:")

defs = [
    ("Agreement",        "this Shareholders\u2019 Agreement and all its Annexures;"),
    ("AUM",              "the aggregate Rand value of assets under management across all active "
                         "client portfolios managed by the Company, as reported in the BitWealth "
                         "platform NAV report at the relevant month-end;"),
    ("Bad Leaver",       "a Party who leaves or is removed in the circumstances defined in clause 12;"),
    ("Board",            "the board of directors of the Company;"),
    ("Business Day",     "any day other than a Saturday, Sunday or recognised South African public holiday;"),
    ("Commencement Date","the date on which the last Party signs this Agreement;"),
    ("Companies Act",    "the Companies Act, 71 of 2008, as amended;"),
    ("Davin",            FOUNDER + ", ID " + FOUNDER_ID + ";"),
    ("Earn-In Shares",   "the shares to be transferred to Simon under the Milestones in clause 8;"),
    ("FAIS Act",         "the Financial Advisory and Intermediary Services Act, 37 of 2002;"),
    ("Founding Shares",  "100 ordinary shares (10%) to be transferred to Simon under clause 7;"),
    ("FSP",              FSP_NAME + " (FSP No. " + FSP_NO + ");"),
    ("Good Leaver",      "a Party who leaves or is removed in the circumstances defined in clause 12;"),
    ("HWM",              "High-Water Mark, the highest NAV previously achieved by a client\u2019s portfolio, "
                         "used to calculate performance fees;"),
    ("Innovation Revenue Share",
                         "the ring-fenced 10% share of net revenue from a Simon-originated, non-AUM, "
                         "inside-BitWealth revenue stream, as defined in clause 14;"),
    ("IP",               "all intellectual property in and to the BitWealth LTH PVR strategy, algorithm, "
                         "software, platform, data, trade marks, branding and related materials;"),
    ("LTH PVR",          "the Long-Term Holder Price Variance Ratio proprietary bitcoin DCA strategy "
                         "developed by Davin and contributed to the Company;"),
    ("Milestone",        "each of the three AUM targets defined in clause 8;"),
    ("Milestone Deadline", "the deadline applicable to each Milestone as set out in clause 8;"),
    ("NAV",              "net asset value of a client\u2019s portfolio;"),
    ("NewCo",            "any new company formed to develop a new product, business line or revenue stream "
                         "separable from the core AUM business;"),
    ("Ordinary Shares",  "ordinary shares of R1.00 par value in the Company, each carrying one vote, "
                         "equal dividend rights and equal rights upon winding-up;"),
    ("Party",            "each of Davin, Simon and the Company, as the context requires;"),
    ("Registrar",        "the Companies and Intellectual Property Commission (CIPC);"),
    ("SHA",              "this Shareholders\u2019 Agreement;"),
    ("Simon",            PARTNER + ", ID " + PARTNER_ID + ";"),
    ("STT Act",          "the Securities Transfer Tax Act, 25 of 2007;"),
    ("Transfer",         "any transfer, cession, sale, pledge, hypothecation or other disposition of shares;"),
    ("VALR",             "VALR Financial Technologies (Pty) Ltd, a registered crypto asset exchange;"),
]
for term, definition in defs:
    defn(doc, term, definition)

clause(doc, "1.2", "Words importing the singular include the plural and vice versa; words importing "
        "any gender include all genders; and references to a person include a natural person, "
        "company, close corporation, trust, partnership or other legal entity.")
clause(doc, "1.3", "Clause headings are for convenience only and do not affect interpretation. "
        "References to clauses and annexures are to clauses and annexures of this Agreement.")
clause(doc, "1.4", "Any reference to legislation includes any amendment, re-enactment or "
        "subordinate legislation thereunder from time to time.")
clause(doc, "1.5", "If any provision of this Agreement conflicts with the Company\u2019s "
        "Memorandum of Incorporation (MOI), the Parties undertake to amend the MOI to give "
        "effect to this Agreement to the extent permitted by the Companies Act.")

# ══════════════════════════════════════════════════════════════════════════════
# 2. SHARE CAPITAL AND STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 2, "Share Capital and Structure")
clause(doc, "2.1", "The Company has 1\u202f000 (one thousand) ordinary shares of R1.00 par value "
        "each in issue. All shares carry equal voting, dividend and liquidation rights.")
clause(doc, "2.2", "As at the Commencement Date, all 1\u202f000 Ordinary Shares are held by Davin. "
        "The IP asset (LTH PVR platform) was contributed by Davin at an agreed value of "
        "R2\u202f500\u202f000 and is recorded on the Company\u2019s balance sheet.")
clause(doc, "2.3", "This Agreement constitutes a genuine equity partnership between co-owners "
        "of the Company. Simon\u2019s Founding Shares and Earn-In Shares are and remain Ordinary "
        "Shares carrying full ownership, voting and economic rights. They may NOT be "
        "reclassified, recharacterised, converted or expropriated without Simon\u2019s prior "
        "written consent.")
clause(doc, "2.4", "The anticipated shareholding at each stage is set out in Annexure A.")

# ══════════════════════════════════════════════════════════════════════════════
# 3. BOARD OF DIRECTORS AND MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 3, "Board of Directors and Management")
clause(doc, "3.1", "The Board shall comprise at least one director. Davin is appointed Managing "
        "Director with full executive authority over all day-to-day operations of the Company.")
clause(doc, "3.2", "Simon is entitled, but not obliged, to serve as a non-executive director of "
        "the Company for as long as he holds any Ordinary Shares. Simon\u2019s appointment as "
        "director shall not impose any minimum time commitment beyond that required by the "
        "Companies Act.")
clause(doc, "3.3", "Board decisions shall be taken by resolution of a majority of directors "
        "present at a duly constituted meeting. Davin, as Managing Director, has a casting vote "
        "in the event of a tied vote on any operational or strategic matter.")
clause(doc, "3.4", "The following matters are RESERVED MATTERS that require the written consent "
        "of both Shareholders (regardless of the Board\u2019s decision):")
for sub, txt in [
    ("3.4.1", "Any amendment to the Company\u2019s Memorandum of Incorporation;"),
    ("3.4.2", "Any winding-up, dissolution or deregistration of the Company;"),
    ("3.4.3", "Any reclassification of Ordinary Shares or creation of any new class of share;"),
    ("3.4.4", "Any action that would dilute Simon\u2019s economic entitlement below its then-current "
              "basis, made for the purpose of such dilution;"),
    ("3.4.5", "Any sale of all or substantially all the Company\u2019s assets or business;"),
    ("3.4.6", "Any change to the Company\u2019s principal business (from LTH PVR bitcoin DCA "
              "investment management)."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "3.5", "Once Simon holds 25% or more of the Ordinary Shares, the following additional "
        "matters become Reserved Matters requiring Simon\u2019s written consent:")
for sub, txt in [
    ("3.5.1", "Any change of FSP / JR arrangement affecting the Company\u2019s FAIS authorisation;"),
    ("3.5.2", "Any change of primary trading exchange (from VALR) or custody arrangement;"),
    ("3.5.3", "Any change to the core client fee structure (performance, management and "
              "exchange fees), where such change is made for the purpose of, or would have the "
              "effect of, materially reducing Simon\u2019s economic entitlement."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "3.6", "Subject to the Reserved Matters above, Davin has full authority to manage the "
        "business. The Parties agree that Simon\u2019s primary contribution is business development "
        "and client acquisition, not operational management.")
clause(doc, "3.7", "An anti-prejudice covenant applies at all times: no operational or structural "
        "change may be made for the purpose of, or with the material effect of, reducing "
        "Simon\u2019s economic entitlement below its then-current basis without Simon\u2019s consent, "
        "regardless of whether it would otherwise constitute a Reserved Matter.")

# ══════════════════════════════════════════════════════════════════════════════
# 4. SIMON'S ROLE AND OBLIGATIONS
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 4, "Simon\u2019s Role and Obligations")
clause(doc, "4.1", "Simon serves as non-executive Head of Business Development in a consulting "
        "capacity, with no minimum time commitment. His equity participation is linked solely "
        "to AUM milestone achievement.")
clause(doc, "4.2", "Simon\u2019s primary responsibilities are:")
for sub, txt in [
    ("4.2.1", "origination and onboarding of new clients;"),
    ("4.2.2", "management and growth of client relationships;"),
    ("4.2.3", "achievement of the AUM Milestones; and"),
    ("4.2.4", "such other business development activities as the Board may reasonably agree."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "4.3", "The Company will provide Simon with reasonable sales, marketing and "
        "compliance/admin support and an annual marketing budget approved by the Board. "
        "Obligations on both sides are defined and not open-ended.")
clause(doc, "4.4", "Simon will use reasonable endeavours to obtain and maintain any FAIS "
        "accreditation required for his representative role. Delays caused by the FSCA, "
        "the FSP, paperwork or the Company are excluded from Simon\u2019s obligation. Only a "
        "wilful or negligent failure by Simon constitutes a breach of this clause.")

# ══════════════════════════════════════════════════════════════════════════════
# 5. INTRODUCED CLIENT REGISTER
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 5, "Introduced Client Register")
clause(doc, "5.1", "The Company will maintain an \u201cIntroduced Client Register\u201d in the format set "
        "out in Annexure C. Any client Simon introduces is logged within 5 (five) Business "
        "Days of first meaningful contact.")
clause(doc, "5.2", "An \u201cIntroduced Client\u201d is one where Simon makes the first meaningful contact "
        "and who is not already in the Company\u2019s active pipeline. Each client is credited once. "
        "Disputed entries are resolved by the Board acting reasonably, failing which by "
        "independent expert determination.")
clause(doc, "5.3", "The Register is maintained for revenue-share and record purposes. Simon\u2019s "
        "AUM milestone vesting is based on TOTAL Company AUM (see clause 8.4) and is NOT "
        "limited to clients he personally introduced.")
clause(doc, "5.4", "An Introduced Client who converts within 18 months of introduction, or within "
        "12 months of Simon\u2019s exit, counts toward any applicable Innovation Revenue Share.")

# ══════════════════════════════════════════════════════════════════════════════
# 6. INFORMATION RIGHTS AND REPORTING
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 6, "Information Rights and Reporting")
clause(doc, "6.1", "While Simon holds any Ordinary Shares, the Company shall provide him with:")
for sub, txt in [
    ("6.1.1", "monthly management accounts within 15 Business Days of month-end;"),
    ("6.1.2", "a monthly AUM report showing total assets under management, per-client balances "
              "(anonymised if required by POPIA), and progress toward the current Milestone;"),
    ("6.1.3", "annual financial statements within 120 days of the Company\u2019s financial year-end; and"),
    ("6.1.4", "read-only access to the BitWealth platform\u2019s management dashboard."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "6.2", "The Company shall hold an annual meeting of shareholders within 6 (six) months "
        "of each financial year-end, at which the annual financial statements are presented "
        "and key business decisions are discussed.")

# ══════════════════════════════════════════════════════════════════════════════
# 7. FOUNDING EQUITY GRANT
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 7, "Founding Equity Grant")
clause(doc, "7.1", "Within 5 (five) Business Days of the Commencement Date (and subject to clause "
        "7.4), Davin shall transfer 100 Ordinary Shares (10% of the issued share capital) to "
        "Simon for nil monetary consideration.")
clause(doc, "7.2", "The Founding Shares carry no performance condition. Simon holds them as a full "
        "equity owner from the date of transfer, subject only to the leaver provisions in "
        "clause 12.")
clause(doc, "7.3", "At the IP-implied price of R2\u202f500 per share, the notional value of the "
        "Founding Shares is R250\u202f000 (Two Hundred and Fifty Thousand Rand). This is an agreed "
        "commercial figure, not a guaranteed or market-derived valuation.")
clause(doc, "7.4", "The transfer of Founding Shares is subject to the following conditions "
        "precedent being fulfilled:")
for sub, txt in [
    ("7.4.1", "Simon having obtained independent legal advice and confirming this in writing;"),
    ("7.4.2", "both Parties having obtained independent tax advice regarding the income-tax, "
              "CGT and donations-tax consequences of a nil-consideration share transfer;"),
    ("7.4.3", "the FSP having provided the Company with a JR appointment or a letter of intent "
              "to appoint; and"),
    ("7.4.4", "confirmation that Simon meets the minimum FAIS fit-and-proper requirements "
              "applicable to his representative role."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "7.5", "The transfer is effected by signed share transfer form, updating the "
        "Company\u2019s share register and lodging the required CIPC forms. All costs, including "
        "STT payable under the STT Act, are borne by the Company.")

# ══════════════════════════════════════════════════════════════════════════════
# 8. AUM MILESTONE EARN-IN
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 8, "AUM Milestone Earn-In (Additional 30%)")
clause(doc, "8.1", "Simon may earn up to a further 30% equity in three equal tranches of 10% "
        "(100 Ordinary Shares per tranche), each constituting a transfer of existing shares "
        "from Davin. No new shares are issued. The tranches and Milestones are:")
ml = doc.add_table(rows=4, cols=5)
ml.style = "Table Grid"; ml.alignment = WD_TABLE_ALIGNMENT.CENTER
ml_h = ["Tranche", "AUM Target", "Equity", "Milestone Deadline", "Target Date"]
ml_w = [Cm(2.4), Cm(3.8), Cm(2.5), Cm(4.0), Cm(3.2)]
for i, col in enumerate(ml.columns):
    for cell in col.cells: cell.width = ml_w[i]
for i, h in enumerate(ml_h):
    c = ml.cell(0, i); c.text = ""
    run = c.paragraphs[0].add_run(h); run.bold = True
    run.font.size = Pt(9.5); run.font.name = BODY_FONT; run.font.color.rgb = WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; set_cell_bg(c, NAVY_HEX)
ml_rows = [
    ["Tranche 1","R50 000 000","+10% (→20%)","24 months from Commencement","~July 2028"],
    ["Tranche 2","R100 000 000","+10% (→30%)","48 months from Commencement","~July 2030"],
    ["Tranche 3","R200 000 000","+10% (→40%)","72 months from Commencement","~July 2032"],
]
for r_idx, row in enumerate(ml_rows, 1):
    bg = "F5F5F5" if r_idx % 2 == 0 else "FFFFFF"
    for c_idx, val in enumerate(row):
        c = ml.cell(r_idx, c_idx); c.text = ""
        run = c.paragraphs[0].add_run(val); run.font.size = Pt(10)
        run.font.name = BODY_FONT; run.font.color.rgb = DARK
        run.bold = (c_idx == 2)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; set_cell_bg(c, bg)
para(doc, "", sb=6)

clause(doc, "8.2", "AUM means the aggregate Rand value of all client portfolios managed by the "
        "Company at month-end, as reported in the BitWealth platform NAV report and confirmed "
        "by the Company\u2019s accounting records.")
clause(doc, "8.3", "Milestones are sequential: a later tranche cannot vest before the earlier "
        "tranche is earned or resolved under clause 9.")
clause(doc, "8.4", "TOTAL-AUM BASIS. Simon\u2019s milestone vesting is tied to TOTAL Company AUM, "
        "including clients introduced by Davin, the marketing investor, third parties or "
        "otherwise. The Introduced Client Register is maintained for revenue-share and "
        "record purposes only and does not limit milestone vesting.")
clause(doc, "8.5", "OBJECTIVE VERIFICATION. A Milestone is achieved when total AUM meets the "
        "target at a month-end AND remains at or above it to the following month-end (two "
        "consecutive month-ends), verified from platform, custodian, exchange and accounting "
        "records.")
clause(doc, "8.6", "ADMINISTRATIVE CONFIRMATION. Board confirmation of Milestone achievement is "
        "administrative only and may not be unreasonably withheld or delayed. If the evidence "
        "shows the threshold met and sustained and the Board has not confirmed within 10 "
        "Business Days, the Milestone is deemed achieved.")
clause(doc, "8.7", "Once a Milestone is confirmed or deemed achieved, Davin shall effect the "
        "relevant share transfer within 5 Business Days in the same manner as clause 7.5.")

# ══════════════════════════════════════════════════════════════════════════════
# 9. MILESTONE SHORTFALL — NEAR-MISS VESTING
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 9, "Milestone Shortfall \u2014 Near-Miss Vesting")
clause(doc, "9.1", "This clause applies equally to all three Milestone Tranches.")
clause(doc, "9.2", "STEP 1 \u2014 EXTENSION ELIGIBILITY (measured at the Milestone Deadline): if "
        "total AUM is at least 80% of the relevant target, Simon automatically receives a "
        "single 6-month extension for that tranche. If AUM is below 80% at the Milestone "
        "Deadline, the tranche lapses immediately and no extension is granted.")
clause(doc, "9.3", "STEP 2 \u2014 FULL VESTING: if AUM reaches 100% of the target at any point "
        "during the 6-month extension, the full 10% tranche vests and is transferred to "
        "Simon within 5 Business Days.")
clause(doc, "9.4", "STEP 3 \u2014 PRO-RATA VESTING: if, at the end of the 6-month extension, AUM "
        "is at least 80% but below 100% of the target, Simon vests a pro-rata portion "
        "= (AUM at extension end \u00f7 Milestone target) \u00d7 10%. Example: R45m of R50m target "
        "(90%) \u2192 Simon vests 9%; the residual 1% remains with Davin.")
clause(doc, "9.5", "STEP 4 \u2014 LAPSE: if, at the end of the extension, AUM has fallen below 80% "
        "of the target, the tranche lapses in full and the shares remain with Davin. Simon "
        "retains all previously vested tranches and his Founding Shares.")
clause(doc, "9.6", "If a Milestone Deadline is extended under clause 9.2, the deadlines of all "
        "subsequent tranches are extended by the same period, because tranches must be "
        "achieved in sequence.")

# ══════════════════════════════════════════════════════════════════════════════
# 10. INVESTOR CAPITAL AND DILUTION
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 10, "Investor Capital and Dilution of Earn-In")
clause(doc, "10.1", "The Parties acknowledge that a marketing investor may be introduced to provide "
        "Year 1 working capital in exchange for equity.")
clause(doc, "10.2", "By agreement between the Parties, any investor equity shall be carved "
        "exclusively from Simon\u2019s 30% earn-in pool — not from Davin\u2019s 60%, and not from "
        "Simon\u2019s Founding Shares. Davin\u2019s 60% and Simon\u2019s 10% founding equity are "
        "unaffected by any investor dilution.")
clause(doc, "10.3", "The AUM Milestone targets (R50m / R100m / R200m) are not affected by investor "
        "dilution. Only the equity reward per tranche reduces proportionally. For example, "
        "with a 15% investor: each tranche becomes approximately 5% (Simon\u2019s maximum "
        "earn-in reduces from 30% to 15%).")
clause(doc, "10.4", "No investor equity may be issued without Simon\u2019s prior written consent, "
        "which shall not be unreasonably withheld.")

# ══════════════════════════════════════════════════════════════════════════════
# 11. TRANSFER RESTRICTIONS AND PRE-EMPTIVE RIGHTS
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 11, "Transfer Restrictions and Pre-Emptive Rights")
clause(doc, "11.1", "No Shareholder may Transfer any Ordinary Shares to a third party without "
        "first offering them to the other Shareholder(s) at the same price and on the same "
        "terms (right of first refusal).")
clause(doc, "11.2", "The selling Shareholder shall give written notice of the proposed Transfer, "
        "including the price, terms and identity of the proposed transferee. The non-selling "
        "Shareholder has 20 Business Days to accept the offer in writing.")
clause(doc, "11.3", "If the non-selling Shareholder does not accept within 20 Business Days, the "
        "selling Shareholder may proceed with the Transfer to the third party at no less "
        "than the offered price and on no better terms, within 60 days of the notice.")
clause(doc, "11.4", "PRE-EMPTIVE SUBSCRIPTION RIGHTS. If the Company proposes to issue new shares, "
        "each Shareholder has a pro-rata right of first offer to subscribe for new shares at "
        "the same price as the incoming subscriber (not a full ratchet or weighted-average "
        "anti-dilution).")
clause(doc, "11.5", "ANTI-CIRCUMVENTION. No shares or options may be issued to Davin, his "
        "associates or staff below independently-determined fair value for the purpose of "
        "diluting Simon, without Simon\u2019s written consent.")
clause(doc, "11.6", "Permitted transfers to a holding company, trust or family member of a "
        "Shareholder are allowed without pre-emptive rights, provided the transferee signs "
        "a deed of adherence to this Agreement.")

# ══════════════════════════════════════════════════════════════════════════════
# 12. GOOD LEAVER AND BAD LEAVER
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 12, "Good Leaver, Bad Leaver and Mutual Protections")
clause(doc, "12.1", "RECIPROCAL APPLICATION. The Good Leaver and Bad Leaver principles apply "
        "EQUALLY and RECIPROCALLY to both Davin and Simon. Each reference to a Party "
        "\u2018leaving\u2019 means that Party ceasing to be actively involved in the Company. "
        "Neither Party enjoys protections the other does not.")
clause(doc, "12.2", "GOOD LEAVER. A Party is a Good Leaver if they leave or are removed by "
        "reason of:")
for sub, txt in [
    ("12.2.1", "death or permanent total disability;"),
    ("12.2.2", "retirement at or after age 65;"),
    ("12.2.3", "removal or termination without cause or material breach by the other Party;"),
    ("12.2.4", "constructive dismissal or a material reduction of their agreed role;"),
    ("12.2.5", "the other Party or the Company failing to provide agreed support;"),
    ("12.2.6", "loss of FSP/JR status not caused by the departing Party; or"),
    ("12.2.7", "material breach of this Agreement by the other Party."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "12.3", "GOOD LEAVER CONSEQUENCE. A Good Leaver retains all vested/earned Ordinary "
        "Shares. Simon\u2019s unearned tranches at the time of exit follow the shortfall mechanism "
        "in clause 9.")
clause(doc, "12.4", "BAD LEAVER. A Party is a Bad Leaver ONLY on:")
for sub, txt in [
    ("12.4.1", "proven fraud, theft or dishonesty directed against the Company or the other Party;"),
    ("12.4.2", "wilful misconduct or gross negligence causing material harm to the Company; or"),
    ("12.4.3", "a final (non-appealable) regulatory finding against that Party in connection "
               "with the Company\u2019s business."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "12.5", "BAD LEAVER CONSEQUENCE. For the Bad Leaver:")
for sub, txt in [
    ("12.5.1", "any unearned/unvested equity lapses immediately;"),
    ("12.5.2", "EARNED shares are acquired by the Company or the remaining Shareholder at FAIR "
               "MARKET VALUE, EXCEPT where the trigger is fraud, theft or dishonesty, in which "
               "case earned shares are acquired at nominal par value of R1.00 per share;"),
    ("12.5.3", "any forced transfer requires PRIOR INDEPENDENT DETERMINATION by a qualified "
               "auditor, expert or arbitrator; and"),
    ("12.5.4", "NEITHER Party may unilaterally declare the other a Bad Leaver. A declaration is "
               "of no effect until confirmed by independent determination or arbitration."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "12.6", "ANTI-RECLASSIFICATION. Neither Party\u2019s shares may be reclassified, converted, "
        "recharacterised (including into any share-incentive or employee scheme), diluted, "
        "cancelled, repurchased or otherwise expropriated without that Party\u2019s prior written "
        "consent. Any purported action in breach is void.")
clause(doc, "12.7", "FOUNDER CONTINUITY. If Davin exits as a Good Leaver (including death or "
        "incapacity), Simon shall retain full access to all IP, systems and credentials "
        "under clause 16 to enable continued operation of the Company.")
clause(doc, "12.8", "RESTRAINT ON COMPETING RE-LAUNCH. A Party who exits as a Bad Leaver, or "
        "whose earned shares are bought back under clause 12.5.2, shall not, for 24 months "
        "within the Republic of South Africa, directly or indirectly re-launch, replicate or "
        "operate the LTH PVR strategy or a materially similar Bitcoin DCA strategy under any "
        "other name, and shall not use the Company\u2019s confidential information or IP to do so.")
clause(doc, "12.9", "CLAWBACK. Clawback of equity applies only to losses caused by a Party\u2019s "
        "wilful misconduct, fraud, gross negligence or proven regulatory breach. Reciprocally, "
        "if a Party\u2019s or the Company\u2019s acts or omissions cause loss of JR/FSP status or "
        "material damage to client relationships, the other Party is treated as a Good Leaver.")
clause(doc, "12.10", "VALUATION. Where fair market value must be determined for a buyout, it shall "
        "be determined by a mutually agreed independent auditor. If not agreed within 10 "
        "Business Days, either Party may apply to the AFSA Arbitration Foundation to appoint "
        "a suitably qualified expert. The expert\u2019s determination is final and binding.")

# ══════════════════════════════════════════════════════════════════════════════
# 13. DRAG-ALONG AND TAG-ALONG
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 13, "Drag-Along and Tag-Along Rights")
clause(doc, "13.1", "DRAG-ALONG. If Davin and/or his associates receive a bona fide offer from an "
        "independent third party to acquire 75% or more of the Ordinary Shares, Davin may "
        "require Simon to sell his pro-rata shares to the same acquirer on the SAME price "
        "and terms. Simon\u2019s drag obligations are subject to:")
for sub, txt in [
    ("13.1.1", "cash-equivalent consideration (or cash at election if non-cash);"),
    ("13.1.2", "warranties and representations limited to title and capacity;"),
    ("13.1.3", "release from any guarantees or personal liabilities in connection with the "
               "Company; and"),
    ("13.1.4", "no sale to a related party or associate of Davin at an undervalue."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "13.2", "TAG-ALONG. If Davin proposes to sell more than 50% of his Ordinary Shares "
        "to a third party, Simon has the right (but not the obligation) to include his "
        "proportionate share of shares in the sale on equivalent terms and at the same price "
        "per share as Davin. Simon must exercise this right within 10 Business Days of "
        "receiving written notice of the proposed sale.")

# ══════════════════════════════════════════════════════════════════════════════
# 14. NEW VENTURES AND INNOVATION REVENUE
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 14, "New Ventures and Innovation Revenue")
clause(doc, "14.1", "DEFAULT — NEWCO. Any materially new product, business line or revenue "
        "stream separable from the core AUM business shall be developed in a NewCo with a "
        "fresh shareholding structure.")
clause(doc, "14.2", "SIMON\u2019S NEWCO ENTITLEMENT. In each NewCo, Simon has an automatic right to "
        "20% founding equity and may earn up to 45% in total through an AUM- or "
        "revenue-based milestone vesting schedule structured similarly to his BitWealth "
        "earn-in. The specific milestone targets are set per venture. The remaining equity "
        "and specific targets are negotiated in good faith, with Davin\u2019s 80% interest and "
        "55% ceiling as the counterpart.")
clause(doc, "14.3", "WHO DECIDES. The Board (with Davin\u2019s control vote) decides whether a new "
        "idea is pursued inside BitWealth or via a NewCo.")
clause(doc, "14.4", "EXCEPTION \u2014 INSIDE BITWEALTH. Where ALL of the following apply, a "
        "Simon-originated non-AUM revenue stream may be kept inside BitWealth:")
for sub, txt in [
    ("14.4.1", "Simon personally originates a genuinely new, NON-AUM revenue stream "
               "(not performance, platform, management or exchange fees on managed assets);"),
    ("14.4.2", "the Board decides to keep it inside BitWealth; and"),
    ("14.4.3", "Simon actively leads the commercial and business-development delivery of "
               "that stream (concept origination, client/partner relationships, go-to-market, "
               "sales and revenue growth), but is NOT required to perform IT development, "
               "engineering, product build or technical operations."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "14.5", "INNOVATION REVENUE SHARE. Where clause 14.4 applies, Simon receives a "
        "ring-fenced Innovation Revenue Share of 10% of the NET revenue of that specific "
        "stream only, for a period of 48 months from first revenue, payable only while Simon "
        "actively leads the stream, with the stream separately accounted and auditable.")
clause(doc, "14.6", "RING-FENCE. The Innovation Revenue Share applies ONLY to the specific new "
        "non-AUM stream. It NEVER touches AUM-based revenue (performance, platform, management "
        "or exchange fees) from any client.")
clause(doc, "14.7", "STACKING. The Innovation Revenue Share stacks on Simon\u2019s equity dividend: "
        "Simon first receives 10% of net stream revenue, then his equity % of the residual "
        "90%. The effective formula is: Simon\u2019s total = 10% + (equity% \u00d7 90%).")
clause(doc, "14.8", "RECIPROCITY. The same framework applies to Davin-originated ventures. "
        "Simon has no Innovation Revenue Share claim over a stream he did not originate "
        "and does not actively lead.")
clause(doc, "14.9", "ANTI-CIRCUMVENTION. Neither Party may re-label core AUM revenue as a "
        "\u2018new stream\u2019 to avoid the other\u2019s legitimate share. Characterisation disputes "
        "go to independent expert determination.")
clause(doc, "14.10", "IP IN NEW STREAMS. All IP created in connection with any new revenue stream "
        "or NewCo vests in the Company (or NewCo), not in Simon personally.")

# ══════════════════════════════════════════════════════════════════════════════
# 15. INTELLECTUAL PROPERTY AND BUSINESS CONTINUITY
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 15, "Intellectual Property, Access and Business Continuity")
clause(doc, "15.1", "All IP (including the LTH PVR strategy, algorithm, source code, platform, "
        "data, trade marks and branding) is and remains the sole property of the Company. "
        "No right to use the IP is granted to any Party personally.")
clause(doc, "15.2", "IP ACCESS ON SIGNING. Upon the Commencement Date, Davin shall provide Simon "
        "with full access to all IP, systems and credentials used to operate the Company, "
        "including: source code and repositories; email and web-hosting environment; production "
        "database; Research Bitcoin data feed/account; the Administration UI; VALR exchange "
        "API and sub-account administration; and all related logins and documentation. "
        "Such access is maintained for the duration of Simon\u2019s shareholding.")
clause(doc, "15.3", "BUSINESS CONTINUITY. The Parties shall explore and implement continuity "
        "measures including: a technical shadow or second engineer; key-man life and disability "
        "insurance; documented operational runbooks; secure credential vaulting with "
        "emergency (\u2018break-glass\u2019) access for Simon; source-code escrow; and regular "
        "data backups and disaster recovery. The Parties target implementation within 6 "
        "months of the Commencement Date.")
clause(doc, "15.4", "IP PROTECTION. The Parties will jointly investigate additional protection "
        "for the LTH PVR strategy, including trade-secret protection, copyright in source "
        "code, trade-mark registration and (subject to independent patent attorney advice) "
        "patent protection in applicable jurisdictions.")

# ══════════════════════════════════════════════════════════════════════════════
# 16. CONFIDENTIALITY AND NON-SOLICITATION
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 16, "Confidentiality and Non-Solicitation")
clause(doc, "16.1", "Each Party shall keep strictly confidential all non-public information "
        "regarding the other Party, the Company, its strategy, clients, technology and "
        "finances. Confidentiality obligations survive termination of this Agreement.")
clause(doc, "16.2", "NON-SOLICITATION. For 12 months following a Party\u2019s exit from the Company, "
        "that Party shall not, within South Africa:")
for sub, txt in [
    ("16.2.1", "solicit or approach any client of the Company with whom that Party had a material "
               "dealing, for the purpose of moving their assets to a competitor; or"),
    ("16.2.2", "solicit or induce any employee, contractor or key person of the Company to leave "
               "their engagement."),
]:
    clause(doc, sub, txt, level=1)
clause(doc, "16.3", "The Parties acknowledge that the Company\u2019s client list and technology are "
        "legitimate protectable interests justifying the above restraints. The Parties agree "
        "that these restraints are reasonable in scope, duration and area.")

# ══════════════════════════════════════════════════════════════════════════════
# 17. DIVIDENDS
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 17, "Dividends")
clause(doc, "17.1", "The declaration and payment of dividends is at the discretion of the Board, "
        "subject to the Companies Act solvency and liquidity requirements.")
clause(doc, "17.2", "Any dividend declared shall be paid to all Shareholders pro-rata to their "
        "Ordinary Shareholding at the record date.")
clause(doc, "17.3", "The Parties acknowledge that the Company is in an early growth phase and "
        "that retained earnings should generally be reinvested in the business during the "
        "first three years. After that, dividends shall be considered annually based on "
        "cashflow, growth requirements and working capital needs.")

# ══════════════════════════════════════════════════════════════════════════════
# 18. DEADLOCK AND DISPUTE RESOLUTION
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 18, "Deadlock and Dispute Resolution")
clause(doc, "18.1", "GOOD FAITH RESOLUTION. If a dispute or disagreement arises between the "
        "Parties, they shall use good faith efforts to resolve it within 10 Business Days "
        "of either Party giving written notice of the dispute.")
clause(doc, "18.2", "ESCALATION. If not resolved within 10 Business Days, either Party may "
        "escalate to arbitration under clause 18.3.")
clause(doc, "18.3", "ARBITRATION. Any unresolved dispute shall be referred to binding arbitration "
        "in Johannesburg under the AFSA Arbitration Foundation of Southern Africa Rules. "
        "The arbitrator shall be an attorney or advocate of at least 10 years\u2019 standing "
        "with experience in commercial/corporate law, appointed by agreement or, failing "
        "agreement within 5 Business Days, by the AFSA Chairman. The decision is final and "
        "binding. This clause does not preclude urgent court applications.")
clause(doc, "18.4", "CONFIDENTIALITY. All dispute resolution proceedings and outcomes are "
        "confidential between the Parties.")

# ══════════════════════════════════════════════════════════════════════════════
# 19. FSP AND REGULATORY MATTERS
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 19, "FSP and Regulatory Matters")
clause(doc, "19.1", "The Company renders financial services as a Juristic Representative of "
        + FSP_NAME + " (FSP No. " + FSP_NO + ") under the FAIS Act. Simon acknowledges "
        "that the JR arrangement is subject to regulatory approval and ongoing compliance, "
        "and that loss of JR status would materially impact the Company.")
clause(doc, "19.2", "FEE TRANSPARENCY. The Company shall share a full fee/profitability model "
        "with Simon showing net economics after the Finova revenue share (20% of performance "
        "fees + 50% of platform fees + 50% of management fees; no fixed monthly hosting).")
clause(doc, "19.3", "The Parties agree that decisions regarding FSP/JR arrangement, primary "
        "exchange or custody are operational matters within Davin\u2019s authority, subject to "
        "the anti-prejudice covenant in clause 3.7 and Reserved Matters at clause 3.5.")

# ══════════════════════════════════════════════════════════════════════════════
# 20. GENERAL PROVISIONS
# ══════════════════════════════════════════════════════════════════════════════
section(doc, 20, "General Provisions")
clause(doc, "20.1", "GOVERNING LAW. This Agreement is governed by the law of the Republic of "
        "South Africa.")
clause(doc, "20.2", "ENTIRE AGREEMENT. This Agreement, together with its Annexures, constitutes "
        "the entire agreement between the Parties regarding its subject matter and supersedes "
        "all prior negotiations, term sheets and heads of terms.")
clause(doc, "20.3", "AMENDMENTS. No amendment is valid unless in writing and signed by all Parties. "
        "Schedule-only amendments (e.g. Annexure updates) require signatures of all Shareholders.")
clause(doc, "20.4", "NO WAIVER. Failure to enforce any provision does not constitute a waiver. "
        "A waiver is only effective if given in writing.")
clause(doc, "20.5", "SEVERABILITY. If any provision is unlawful or unenforceable, it is severed "
        "and the remainder continues in force.")
clause(doc, "20.6", "NOTICES. Notices must be in writing, delivered by email (with read receipt) "
        "or registered post to the addresses recorded in the signature block. "
        "Email notices are deemed received when sent during Business Hours on a Business Day.")
clause(doc, "20.7", "COSTS. Each Party bears its own legal and other costs of negotiating and "
        "executing this Agreement unless otherwise agreed.")
clause(doc, "20.8", "COUNTERPARTS. This Agreement may be signed in counterparts, each of which "
        "constitutes an original, and all together constitute one agreement.")
clause(doc, "20.9", "DOMICILIA. The Parties choose the following as their domicilia for notices:")
para(doc, f"Davin: [Davin\u2019s residential address to insert] | Email: [email to insert]", indent=2.0)
para(doc, f"Simon: [Simon\u2019s residential address to insert] | Email: [email to insert]", indent=2.0)
para(doc, f"Company: 89 Bloekom Curve, Terenure ext. 16, Kempton Park, 1619 | Email: info@bitwealth.co.za", indent=2.0)
clause(doc, "20.10", "DEED OF ADHERENCE. Any person who acquires shares in the Company must, as a "
        "condition of the transfer, sign a deed of adherence confirming they are bound by "
        "this Agreement in the capacity of a Shareholder.")

# ══════════════════════════════════════════════════════════════════════════════
# SIGNATURES
# ══════════════════════════════════════════════════════════════════════════════
section(doc, None, "SIGNATURES", sb=14)
para(doc,
    "SIGNED BY THE PARTIES at the places and on the dates indicated below, "
    "each signatory warranting full authority:")

doc.add_paragraph()

def sig_block(doc, party_label, name, capacity, email_placeholder):
    p = doc.add_paragraph()
    r = p.add_run(party_label.upper())
    r.bold = True; r.font.name = BODY_FONT; r.font.size = Pt(BODY_SZ); r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)
    for line in [
        f"Signed: __________________________\t\tDate: __________________________",
        f"Full Name: {name}",
        f"Capacity: {capacity}",
        f"At: __________________________",
        f"Email: {email_placeholder}",
        f"Witness 1: ______________________\t\tWitness 2: ______________________",
    ]:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(3)
        p2.paragraph_format.left_indent = Cm(0.5)
        r2 = p2.add_run(line); r2.font.name = BODY_FONT; r2.font.size = Pt(BODY_SZ)
        r2.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

sig_block(doc,
    "For and on behalf of " + CO_NAME + " and in his personal capacity as Founder",
    FOUNDER, "Managing Director / Founder", "[davin@bitwealth.co.za]")
sig_block(doc,
    PARTNER + " in his personal capacity as Business Development Partner",
    PARTNER, "Business Development Partner", "[simon's email to insert]")

hrule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# ANNEXURES
# ══════════════════════════════════════════════════════════════════════════════
# Annexure A — Cap table
doc.add_page_break()
section(doc, None, "ANNEXURE A — CAP TABLE AT EACH STAGE", sb=4)

def cap_tbl(doc, title, rows, hdr_bg=NAVY_HEX):
    para(doc, title, bold=True, sb=8, sa=4)
    t = doc.add_table(rows=len(rows)+1, cols=5)
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(5.5), Cm(2.2), Cm(2.4), Cm(2.2), Cm(4.0)]
    for i, col in enumerate(t.columns):
        for cell in col.cells: cell.width = widths[i]
    for i, h in enumerate(["Shareholder","Shares","% Holding","Change","Notes"]):
        c = t.cell(0, i); c.text = ""
        run = c.paragraphs[0].add_run(h); run.bold = True
        run.font.size = Pt(9); run.font.name = BODY_FONT; run.font.color.rgb = WHITE
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; set_cell_bg(c, hdr_bg)
    for r_idx, row in enumerate(rows, 1):
        bg = "F5EBC8" if row[0].startswith("TOTAL") else ("F5F5F5" if r_idx%2==0 else "FFFFFF")
        bold = row[0].startswith("TOTAL")
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
        for c_idx, (val, al) in enumerate(zip(row, aligns)):
            c = t.cell(r_idx, c_idx); c.text = ""
            run = c.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10); run.font.name = BODY_FONT; run.font.color.rgb = DARK
            run.bold = bold; c.paragraphs[0].alignment = al; set_cell_bg(c, bg)

cap_tbl(doc, "Stage 1 — On Signing (Founding 10%)", [
    [FOUNDER,"900","90%","−100","IP contributor; technology"],
    [PARTNER, "100","10%","+100","Founding grant; nil consideration"],
    ["TOTAL","1,000","100%","—",""],
])
cap_tbl(doc, "Stage 2 — Post Tranche 1 (R50m AUM)", [
    [FOUNDER,"800","80%","−100","Transfers from own holding"],
    [PARTNER, "200","20%","+100","Tranche 1 earned"],
    ["TOTAL","1,000","100%","—",""],
])
cap_tbl(doc, "Stage 3 — Post Tranche 2 (R100m AUM)", [
    [FOUNDER,"700","70%","−100","Transfers from own holding"],
    [PARTNER, "300","30%","+100","Tranche 2 earned"],
    ["TOTAL","1,000","100%","—",""],
])
cap_tbl(doc, "Stage 4 — Full Earn-In (R200m AUM)", [
    [FOUNDER,"600","60%","−100","Final if all milestones met"],
    [PARTNER, "400","40%","+100","Full earn-in — maximum"],
    ["TOTAL","1,000","100%","—",""],
])

# Annexure B — Milestone schedule
doc.add_page_break()
section(doc, None, "ANNEXURE B — MILESTONE SCHEDULE AND EXTENSION TRACKER", sb=4)
para(doc, "To be updated by the Company\u2019s Secretary each time a Milestone is achieved, "
     "extended or lapses. A copy is to be provided to both Shareholders.", italic=True)
mb = doc.add_table(rows=4, cols=7)
mb.style = "Table Grid"; mb.alignment = WD_TABLE_ALIGNMENT.CENTER
mb_h = ["Tranche","Target","Deadline","AUM@Deadline","Extension?","AUM@Ext.End","Outcome"]
mb_w = [Cm(2.0), Cm(2.5), Cm(2.2), Cm(2.6), Cm(2.2), Cm(2.6), Cm(2.4)]
for i, col in enumerate(mb.columns):
    for cell in col.cells: cell.width = mb_w[i]
for i, h in enumerate(mb_h):
    c = mb.cell(0, i); c.text = ""
    run = c.paragraphs[0].add_run(h); run.bold = True
    run.font.size = Pt(8.5); run.font.name = BODY_FONT; run.font.color.rgb = WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; set_cell_bg(c, NAVY_HEX)
for r, row in enumerate([
    ["Tranche 1","R50m","~Jul 2028","...","...","...","..."],
    ["Tranche 2","R100m","~Jul 2030","...","...","...","..."],
    ["Tranche 3","R200m","~Jul 2032","...","...","...","..."],
], 1):
    bg = "F5F5F5" if r % 2 == 0 else "FFFFFF"
    for c_idx, val in enumerate(row):
        c = mb.cell(r, c_idx); c.text = ""
        run = c.paragraphs[0].add_run(val)
        run.font.size = Pt(9); run.font.name = BODY_FONT; run.font.color.rgb = DARK
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; set_cell_bg(c, bg)

# Annexure C — Introduced Client Register format
doc.add_page_break()
section(doc, None, "ANNEXURE C — INTRODUCED CLIENT REGISTER (FORMAT)", sb=4)
para(doc, "Maintained by the Company. Updated within 5 Business Days of first meaningful contact. "
     "Disputes resolved by the Board; failing which by independent expert.", italic=True)
ic = doc.add_table(rows=3, cols=6)
ic.style = "Table Grid"; ic.alignment = WD_TABLE_ALIGNMENT.CENTER
ic_h = ["Date Logged","Client Ref (POPIA)","Introduced By","First Contact","Status","Notes"]
ic_w = [Cm(2.4), Cm(3.0), Cm(2.8), Cm(2.8), Cm(2.4), Cm(3.0)]
for i, col in enumerate(ic.columns):
    for cell in col.cells: cell.width = ic_w[i]
for i, h in enumerate(ic_h):
    c = ic.cell(0, i); c.text = ""
    run = c.paragraphs[0].add_run(h); run.bold = True
    run.font.size = Pt(9); run.font.name = BODY_FONT; run.font.color.rgb = WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER; set_cell_bg(c, NAVY_HEX)
for r in range(1, 3):
    for c_idx in range(6):
        c = ic.cell(r, c_idx); c.text = ""; set_cell_bg(c, "F5F5F5" if r%2==0 else "FFFFFF")
        run = c.paragraphs[0].add_run("...")
        run.font.size = Pt(9); run.font.name = BODY_FONT; run.font.color.rgb = GREY

out = OUT_DIR / "BitWealth_SHA_v1.docx"
doc.save(out)
print(f"Saved: {out}")
