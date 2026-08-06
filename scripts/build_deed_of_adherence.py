"""
BitWealth Asset Managers (Pty) Ltd — Deed of Adherence
Substitutes Mhuri Investment Holdings (Pty) Ltd as registered holder of Simon's
Founding Shares / Earn-In Shares in place of Simon personally, while keeping his
personal obligations (director duties, FAIS, restraint, confidentiality, Leaver
triggers) intact. Also closes the change-of-control gap created by holding
shares via a company, and records spousal consent (Simon is married in
community of property).

Output: docs/Shareholding/BitWealth_Deed_of_Adherence_Mhuri_v1.docx

NOT legal advice — must be reviewed and executed under the supervision of a
qualified South African attorney before it creates any binding obligations.
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("docs") / "Shareholding"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Parties ──────────────────────────────────────────────────────────────────
FOUNDER    = "Davin Harald Gaier"
FOUNDER_ID = "8405025239081"
PARTNER    = "Simon Henry Newbold Hobday"
PARTNER_ID = "6806175080088"
CO_NAME    = "BitWealth Asset Managers (Pty) Ltd"
CO_REG     = "2026/090346/07"
SIMCO_NAME = "Mhuri Investment Holdings (Pty) Ltd"
SIMCO_REG  = "2017/334996/07"
SIMCO_ADDR = "6 Sardinia, Golf Close, Lonehill, Gauteng, 2196"

# ── Colours ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x0A, 0x2A, 0x43)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
GREY  = RGBColor(0x55, 0x55, 0x55)
RED   = RGBColor(0xB7, 0x1C, 0x1C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX = "0A2A43"
LGREY_HEX = "F2F2F2"

BODY_FONT = "Calibri"
BODY_SZ = 11


# ═══════════════════════════════════════════════════════════════
# HELPERS (mirrors scripts/build_sha.py conventions)
# ═══════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def para(doc, text="", bold=False, italic=False, size=BODY_SZ, color=None,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, sb=2, sa=4, indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    p.alignment = align
    if indent is not None:
        p.paragraph_format.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.name = BODY_FONT; run.font.size = Pt(size)
        run.font.color.rgb = color if color else DARK
    return p


def clause(doc, number, text, level=0, bold_num=True):
    p = doc.add_paragraph()
    indent = level * 1.1
    p.paragraph_format.left_indent = Cm(indent + 1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if number:
        rn = p.add_run(f"{number}\t")
        rn.bold = bold_num; rn.font.name = BODY_FONT; rn.font.size = Pt(BODY_SZ)
        rn.font.color.rgb = DARK
    rt = p.add_run(text)
    rt.font.name = BODY_FONT; rt.font.size = Pt(BODY_SZ); rt.font.color.rgb = DARK
    return p


def section(doc, number, title, sb=14, page_break=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.page_break_before = page_break
    if number:
        r = p.add_run(f"{number}.\t{title.upper()}")
    else:
        r = p.add_run(title.upper())
    r.bold = True; r.font.name = BODY_FONT; r.font.size = Pt(13)
    r.font.color.rgb = NAVY
    return p


def hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), NAVY_HEX)
    pBdr.append(b); pPr.append(pBdr)


def defn(doc, term, definition):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rt = p.add_run(f'"{term}"\t')
    rt.bold = True; rt.font.name = BODY_FONT; rt.font.size = Pt(BODY_SZ); rt.font.color.rgb = DARK
    rd = p.add_run(definition)
    rd.font.name = BODY_FONT; rd.font.size = Pt(BODY_SZ); rd.font.color.rgb = DARK


def sig_block(doc, party_label, name, capacity, email_placeholder):
    p = doc.add_paragraph()
    r = p.add_run(party_label.upper())
    r.bold = True; r.font.name = BODY_FONT; r.font.size = Pt(BODY_SZ); r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)
    for line in [
        "Signed: __________________________\t\tDate: __________________________",
        f"Full Name: {name}",
        f"Capacity: {capacity}",
        "At: __________________________",
        f"Email: {email_placeholder}",
        "Witness 1: ______________________\t\tWitness 2: ______________________",
    ]:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(1); p2.paragraph_format.space_after = Pt(3)
        p2.paragraph_format.left_indent = Cm(0.5)
        r2 = p2.add_run(line); r2.font.name = BODY_FONT; r2.font.size = Pt(BODY_SZ)
        r2.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(8)


# ═══════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles["Normal"]
style.font.name = BODY_FONT; style.font.size = Pt(BODY_SZ)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.2)
    s.left_margin = s.right_margin = Cm(2.5)

# ── Cover / Title block ───────────────────────────────────────────────────────
para(doc, "DEED OF ADHERENCE", bold=True, size=20, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=6)
para(doc, "to the Shareholders' Agreement of", bold=False, size=12, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
para(doc, CO_NAME, bold=True, size=14, color=NAVY,
     align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
para(doc, f"(Registration No. {CO_REG})", size=11, color=GREY,
     align=WD_ALIGN_PARAGRAPH.CENTER, sa=10)
hrule(doc)
para(doc, "PRIVATE AND CONFIDENTIAL", bold=True, size=10, color=RED,
     align=WD_ALIGN_PARAGRAPH.CENTER, sa=2)
para(doc,
    "THIS DEED HAS NOT BEEN REVIEWED BY AN ATTORNEY. IT IS PROVIDED AS A WORKING DRAFT "
    "TO GIVE EFFECT TO THE PARTIES\u2019 AGREEMENT THAT MHURI INVESTMENT HOLDINGS (PTY) LTD "
    "WILL HOLD SIMON\u2019S SHARES IN PLACE OF SIMON PERSONALLY. IT MUST BE REVIEWED, FINALISED "
    "AND EXECUTED UNDER THE SUPERVISION OF A QUALIFIED SOUTH AFRICAN ATTORNEY BEFORE IT "
    "CREATES ANY BINDING OBLIGATIONS.",
    size=9, italic=True, color=RED, align=WD_ALIGN_PARAGRAPH.CENTER, sa=12)
hrule(doc)

para(doc, "ENTERED INTO BY AND BETWEEN:", bold=True, sb=10, sa=6)
para(doc, f"1.\t{CO_NAME}, Registration No. {CO_REG} (\u201cthe Company\u201d);", indent=0.5)
para(doc, f"2.\t{FOUNDER}, South African, ID {FOUNDER_ID} (\u201cDavin\u201d);", indent=0.5)
para(doc, f"3.\t{PARTNER}, South African, ID {PARTNER_ID} (\u201cSimon\u201d); and", indent=0.5)
para(doc, f"4.\t{SIMCO_NAME}, Registration No. {SIMCO_REG}, a private company incorporated "
          f"under the laws of the Republic of South Africa, registered address "
          f"{SIMCO_ADDR} (\u201cMhuri\u201d).",
     indent=0.5)
para(doc, "(Davin, Simon, Mhuri and the Company are collectively \u201cthe Parties\u201d and "
          "individually a \u201cParty\u201d.)", sb=4, sa=8)

# ── Preamble ─────────────────────────────────────────────────────────────────
section(doc, None, "PREAMBLE", sb=10)
para(doc, "A.\tDavin, Simon and the Company are parties to a Shareholders\u2019 Agreement in "
          "respect of the Company dated [insert SHA signature date] (\u201cthe SHA\u201d), which "
          "governs Simon\u2019s Founding Shares, Earn-In Shares and all related rights and "
          "obligations.", indent=0.5)
para(doc, "B.\tThe SHA (clauses 11.6 and 20.10) permits Simon to transfer his shares to a "
          "holding company without triggering the other Shareholder\u2019s pre-emptive rights, "
          "provided the transferee signs a deed of adherence confirming it is bound by the "
          "SHA in the capacity of a Shareholder.", indent=0.5)
para(doc, "C.\tSimon has requested, and Davin has agreed, that Mhuri \u2014 a private investment "
          "holding company owned by Simon and his spouse \u2014 be registered as the holder of "
          "record of Simon\u2019s Founding Shares and any Earn-In Shares that vest under the SHA, "
          "instead of Simon holding those shares in his personal capacity.", indent=0.5)
para(doc, "D.\tSimon is married in community of property, such that his interest in Mhuri "
          "(and indirectly in the Ordinary Shares) forms part of the joint estate shared "
          "with his spouse.", indent=0.5)
para(doc, "E.\tThe Parties record in this Deed the terms on which Mhuri adheres to the SHA, "
          "the safeguards required because the shares will be held through a company rather "
          "than by Simon personally, and the spousal consent applicable given Simon\u2019s "
          "matrimonial property regime.", indent=0.5)
para(doc, "THE PARTIES AGREE AS FOLLOWS:", bold=True, sb=8, sa=8)

# ── 1. Definitions ─────────────────────────────────────────────────────────
section(doc, 1, "Definitions and Interpretation", sb=10)
clause(doc, "1.1", "Unless otherwise defined in this Deed, terms defined in the SHA bear the "
       "same meaning when used in this Deed.")
defn(doc, "Deed", "this Deed of Adherence.")
defn(doc, "Mhuri Shares", "all shares in the issued share capital of Mhuri.")
defn(doc, "Simon\u2019s Spouse", "the person to whom Simon is married in community of property, "
     "and who is a co-shareholder of Mhuri.")
defn(doc, "Change of Control Event", "any change, direct or indirect, in the shareholding, "
     "directorship or effective control of Mhuri, including any Transfer of Mhuri Shares, "
     "any allotment of new Mhuri Shares to a person other than Simon or Simon\u2019s Spouse, "
     "any amendment to Mhuri\u2019s MOI affecting control, and any division, attachment or "
     "reallocation of Mhuri Shares arising from divorce, death, sequestration, liquidation "
     "or insolvency of Simon, Simon\u2019s Spouse, or their joint estate.")
clause(doc, "1.2", "This Deed is supplemental to, and forms part of, the SHA. In the event of "
       "any conflict between this Deed and the SHA, this Deed prevails to the extent of the "
       "conflict, but only insofar as it relates to the substitution of Mhuri as registered "
       "shareholder and the matters expressly dealt with in this Deed.")

# ── 2. Adherence and substitution ────────────────────────────────────────────
section(doc, 2, "Adherence and Substitution as Shareholder", sb=14)
clause(doc, "2.1", "Mhuri hereby adheres to, and agrees to be bound by, all of the terms of the "
       "SHA as if it were named as a Party in the capacity of \u201cSimon\u201d, but only to the "
       "extent that those terms relate to the holding, transfer, voting, dividend and "
       "forfeiture/buy-back rights and obligations attaching to Ordinary Shares (including "
       "clauses 2, 3.4\u20133.7, 7, 8, 9, 10, 11, 12.3, 12.5, 13, 14, 17 and 20 of the SHA).")
clause(doc, "2.2", "With effect from the date of this Deed: (a) the Founding Shares shall be "
       "registered in the name of Mhuri and not Simon; and (b) any Earn-In Shares that vest "
       "under clause 8 or 9 of the SHA after the date of this Deed shall likewise be "
       "transferred into and registered in the name of Mhuri, in each case subject to the "
       "conditions precedent in SHA clause 7.4 and payment of any applicable Securities "
       "Transfer Tax by the Company as contemplated in SHA clause 7.5.")
clause(doc, "2.3", "Wherever the SHA refers to shares being transferred to, held by, retained "
       "by, or forfeited or bought back from \u201cSimon\u201d, that reference is read as a reference "
       "to Mhuri for so long as Mhuri is the registered holder of those shares.")
clause(doc, "2.4", "The Company shall update its securities register accordingly and Mhuri "
       "shall be recorded as a Shareholder for all purposes of the Companies Act and the "
       "Company\u2019s MOI.")

# ── 3. Personal obligations remain with Simon ────────────────────────────────
section(doc, 3, "Personal Obligations Remain with Simon", sb=14)
clause(doc, "3.1", "For the avoidance of doubt, the substitution of Mhuri as registered "
       "shareholder does NOT transfer, dilute or affect any obligation of Simon that is "
       "personal to him, including but not limited to: his role and duties as a director "
       "and Head of Business Development (SHA clause 4); his obligation to obtain and "
       "maintain FAIS accreditation and act as a representative of the FSP (SHA clauses "
       "4.4, 7.4.4 and 19); his restraint of trade, non-solicitation and confidentiality "
       "undertakings (SHA clauses 12.8 and 16); and the conduct-based triggers that "
       "determine whether he is a Good Leaver or Bad Leaver (SHA clause 12).")
clause(doc, "3.2", "Where the SHA\u2019s Good Leaver / Bad Leaver provisions are triggered by "
       "Simon\u2019s personal conduct or departure, the CONSEQUENCE of that trigger (retention, "
       "forfeiture or forced buy-back of Ordinary Shares) shall be applied against the "
       "Ordinary Shares then registered in Mhuri\u2019s name, and Mhuri irrevocably undertakes to "
       "sign all documents and take all steps necessary to give effect to that consequence "
       "within the timeframes specified in the SHA.")
clause(doc, "3.3", "Simon remains personally liable, jointly and severally with Mhuri, for any "
       "breach of the SHA or this Deed arising from his personal conduct.")

# ── 4. Change of control / look-through ──────────────────────────────────────
section(doc, 4, "Change of Control \u2014 Deemed Transfer", sb=14)
clause(doc, "4.1", "Simon and Mhuri warrant that, as at the date of this Deed, the Mhuri Shares "
       "are held only by Simon and Simon\u2019s Spouse, and that Mhuri holds no assets other than "
       "its shareholding in the Company (and any other permitted investments agreed with "
       "Davin in writing).")
clause(doc, "4.2", "A Change of Control Event is DEEMED to be a Transfer of the Ordinary Shares "
       "then held by Mhuri for the purposes of clauses 11 (Transfer Restrictions and "
       "Pre-Emptive Rights), 12 (Good Leaver, Bad Leaver and Mutual Protections) and 13 "
       "(Drag-Along and Tag-Along Rights) of the SHA, and those clauses apply mutatis "
       "mutandis as if Mhuri itself were transferring the Ordinary Shares, regardless of "
       "whether any Ordinary Share is itself transferred.")
clause(doc, "4.3", "Simon and Mhuri shall give the Company and Davin at least 20 (twenty) "
       "Business Days\u2019 prior written notice of any proposed Change of Control Event, "
       "save where the event arises from death, sequestration or a court order, in which "
       "case notice shall be given as soon as reasonably possible.")
clause(doc, "4.4", "Mhuri shall procure that its MOI, and any shareholders\u2019 agreement between "
       "its own shareholders, is not adopted or amended in a manner that conflicts with, or "
       "frustrates, this Deed or the SHA.")

# ── 5. Matrimonial property / spousal consent ────────────────────────────────
section(doc, 5, "Matrimonial Property and Spousal Consent", sb=14)
clause(doc, "5.1", "Simon confirms that he is married in community of property to Simon\u2019s "
       "Spouse, such that the Mhuri Shares held by both Simon and Simon\u2019s Spouse form part "
       "of their joint estate.")
clause(doc, "5.2", "Simon\u2019s Spouse signs the Spousal Consent attached as Annexure A to this "
       "Deed, confirming that she: (a) is aware of the terms of the SHA and this Deed; "
       "(b) consents, to the extent required under section 15 of the Matrimonial Property "
       "Act 88 of 1984 or otherwise, to the Mhuri Shares and the underlying Ordinary Shares "
       "being subject to the SHA and this Deed, including the restraint of trade, "
       "confidentiality, transfer restriction, drag-along/tag-along and Leaver provisions; "
       "and (c) agrees that clause 4 of this Deed (Change of Control) applies "
       "notwithstanding any future division of the joint estate on divorce or death.")
clause(doc, "5.3", "Simon and Mhuri shall promptly notify the Company and Davin of any change "
       "in Simon\u2019s marital status or matrimonial property regime, and shall procure an "
       "equivalent consent from any new or substituted spouse before that person acquires "
       "any interest in Mhuri.")

# ── 6. Insolvency ─────────────────────────────────────────────────────────────
section(doc, 6, "Insolvency of Mhuri or the Joint Estate", sb=14)
clause(doc, "6.1", "If Mhuri is placed under business rescue, wound up or liquidated, or if "
       "Simon, Simon\u2019s Spouse or their joint estate is provisionally or finally sequestrated, "
       "this shall be treated as a Bad Leaver event under SHA clause 12.4 in respect of the "
       "Ordinary Shares held by Mhuri, entitling the Company and Davin to invoke the forced "
       "transfer / buy-back mechanism in SHA clause 12.5, subject to the independent "
       "valuation required by SHA clause 12.5.3.")
clause(doc, "6.2", "This clause 6 does not apply to the extent the relevant insolvency event "
       "arises from circumstances that would constitute a Good Leaver event for Simon under "
       "SHA clause 12.2, in which case the Good Leaver consequence in SHA clause 12.3 "
       "applies instead.")

# ── 7. Tax and STT ────────────────────────────────────────────────────────────
section(doc, 7, "Tax and Securities Transfer Tax", sb=14)
clause(doc, "7.1", "The Parties acknowledge that the conditions precedent in SHA clause 7.4 "
       "(including independent legal and tax advice) apply equally to a transfer of shares "
       "to Mhuri, and confirm that such advice shall specifically address: (a) the income "
       "tax, donations tax and CGT consequences of transferring shares to a company rather "
       "than to Simon personally; and (b) any risk that SARS treats the value transferred "
       "as remuneration for Simon\u2019s personal services (assignment of income), rather than "
       "as a transfer to Mhuri, for tax purposes.")
clause(doc, "7.2", "Securities Transfer Tax and other transfer costs remain for the account of "
       "the Company as provided in SHA clause 7.5, irrespective of Mhuri being the "
       "transferee.")

# ── 8. Notices ────────────────────────────────────────────────────────────────
section(doc, 8, "Notices", sb=14)
clause(doc, "8.1", "For all matters relating to Mhuri\u2019s shareholding, Mhuri\u2019s domicilium for "
       "notices under SHA clause 20.6 is:")
para(doc, f"Mhuri Investment Holdings (Pty) Ltd: {SIMCO_ADDR} | Email: [Mhuri email to insert]",
     indent=2.0)
clause(doc, "8.2", "Simon\u2019s domicilium for all matters relating to his personal obligations "
       "(director duties, FAIS representative status, restraint, confidentiality) remains as "
       "recorded in the SHA signature block.")

# ── 9. General ────────────────────────────────────────────────────────────────
section(doc, 9, "General", sb=14)
clause(doc, "9.1", "This Deed, read together with the SHA, constitutes the whole agreement "
       "between the Parties regarding Mhuri\u2019s adherence to the SHA and supersedes any prior "
       "understanding on that subject.")
clause(doc, "9.2", "No amendment to this Deed is valid unless in writing and signed by all "
       "Parties.")
clause(doc, "9.3", "This Deed is governed by the laws of the Republic of South Africa, and the "
       "dispute resolution mechanism in SHA clause 18 applies to any dispute arising under "
       "this Deed as if it were a dispute under the SHA.")
clause(doc, "9.4", "This Deed may be signed in counterparts, each of which constitutes an "
       "original, and all together constitute one Deed.")

# ══════════════════════════════════════════════════════════════════════════════
# SIGNATURES
# ══════════════════════════════════════════════════════════════════════════════
section(doc, None, "SIGNATURES", sb=14)
para(doc, "SIGNED BY THE PARTIES at the places and on the dates indicated below, each "
          "signatory warranting full authority:")
doc.add_paragraph()

sig_block(doc, "For and on behalf of " + CO_NAME, FOUNDER,
           "Managing Director / Founder", "[davin@bitwealth.co.za]")
sig_block(doc, PARTNER + " in his personal capacity", PARTNER,
           "Business Development Partner", "[simon's email to insert]")
sig_block(doc, "For and on behalf of " + SIMCO_NAME, PARTNER,
           "Director, duly authorised", "[Mhuri email to insert]")
sig_block(doc, "For and on behalf of " + SIMCO_NAME + " (second director, if required by its MOI)",
           "[Full name of Simon\u2019s Spouse]", "Director, duly authorised",
           "[Mhuri email to insert]")

hrule(doc)

# ── Annexure A — Spousal Consent ─────────────────────────────────────────────
doc.add_page_break()
section(doc, None, "ANNEXURE A \u2014 SPOUSAL CONSENT", sb=4)
para(doc, "I, [Full name of Simon\u2019s Spouse], ID [spouse\u2019s ID number to insert], being married "
          "in community of property to Simon Henry Newbold Hobday and a shareholder of "
          f"{SIMCO_NAME} (Registration No. {SIMCO_REG}), confirm that:", sb=8)
para(doc, "1.\tI have read and understood the Shareholders\u2019 Agreement of " + CO_NAME + " and "
          "this Deed of Adherence;", indent=0.5)
para(doc, "2.\tI consent, to the extent required under section 15 of the Matrimonial Property "
          "Act 88 of 1984 or otherwise, to the Mhuri Shares and the underlying Ordinary "
          "Shares in " + CO_NAME + " being subject to, and dealt with in accordance with, "
          "the Shareholders\u2019 Agreement and this Deed of Adherence, including the restraint "
          "of trade, confidentiality, transfer restriction, drag-along/tag-along and Leaver "
          "provisions contained in them;", indent=0.5)
para(doc, "3.\tI agree that clause 4 of the Deed of Adherence (Change of Control) will apply "
          "notwithstanding any future division of our joint estate on divorce or death; and",
     indent=0.5)
para(doc, "4.\tI will not, without the prior written consent of Davin Harald Gaier and " +
          CO_NAME + ", deal with the Mhuri Shares in any manner inconsistent with the "
          "Shareholders\u2019 Agreement or this Deed of Adherence.", indent=0.5)
doc.add_paragraph()
sig_block(doc, "Spousal Consent", "[Full name of Simon\u2019s Spouse]", "Shareholder of Mhuri",
           "[spouse's email to insert]")

path = OUT_DIR / "BitWealth_Deed_of_Adherence_Mhuri_v1.docx"
doc.save(path)
print(f"Saved: {path}")
