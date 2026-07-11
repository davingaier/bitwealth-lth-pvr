"""
BitWealth (Pty) Ltd — Simon Hobday Equity Proposal v3 (COMPLETE, self-contained)
Output: docs/Shareholding/BitWealth_Simon_Equity_Proposal_v3.docx

Full updated term sheet built from v2, integrating ALL session changes:
  · Real parties/IDs/reg/FSP numbers carried from Davin's v2 edits
  · Objective 2-month-sustained AUM verification + deemed-achieved
  · Near-miss pro-rata vesting (replaces the 4 v2 options)
  · Introduced Client Register + conversion tail
  · Defined non-exec BD consulting role; FAIS reasonable-endeavours
  · IP evidence pack; independent tax advice CP
  · Fee transparency + consultation/anti-prejudice (no 10% veto; reserved matters at 25%)
  · Expanded Good Leaver; narrowed Bad Leaver (FMV except fraud; independent determination)
  · Confidentiality + non-solicitation (replaces broad non-compete)
  · Drag-along with minority protections; tag-along
  · Narrowed + reciprocal clawback; pro-rata pre-emptive + anti-circumvention
  · NEW: New Ventures & Innovation Revenue framework WITH worked example
        clarifying that the ring-fenced share STACKS on the equity dividend share
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("docs") / "Shareholding"
OUT_DIR.mkdir(parents=True, exist_ok=True)

# ── Parties (from Davin's v2 edits) ─────────────────────────────────────────
FOUNDER      = "Davin Harald Gaier"
FOUNDER_ID   = "8405025239081"
PARTNER      = "Simon Henry Newbold Hobday"
PARTNER_ID   = "6806175080088"
COMPANY_REG  = "2026 / 090346 / 07"
FSP_NO       = "21095"
SIGN_DATE    = "July 2026"
DOC_DATE     = "11 July 2026"

# ── Colours ─────────────────────────────────────────────────────────────────
GOLD  = RGBColor(0xC9, 0xA2, 0x27)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1B, 0x5E, 0x20)
RED   = RGBColor(0xB7, 0x1C, 0x1C)

L_GOLD_HEX = "F5EBC8"
GOLD_HEX   = "C9A227"
DARK_HEX   = "1A1A1A"
TEAL_HEX   = "006064"
GREEN_HEX  = "1B5E20"
LGREY_HEX  = "F5F5F5"
LGREEN_HEX = "E8F5E9"
LGOLD2_HEX = "FFF8E1"


# ═══════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, size=10, color=None, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else DARK
    return run


def add_heading(doc, text, level=1, gold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper() if level == 1 else text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = GOLD if gold else DARK
    run.font.name = "Calibri"
    p.paragraph_format.keep_with_next = True
    return p


def add_paragraph(doc, text, size=10, color=None, italic=False, sb=2, sa=4, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    run = p.add_run(text)
    run.font.size = Pt(size); run.bold = bold
    run.font.color.rgb = color if color else DARK
    run.italic = italic
    return p


def add_bullet(doc, text, size=10, indent=0.6):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.color.rgb = DARK
    return p


def hr(doc, color=GOLD_HEX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color)
    pBdr.append(bottom); pPr.append(pBdr)
    return p


def doc_header(doc, title, subtitle, date_str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = GOLD; run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph()
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(11); run2.font.color.rgb = DARK
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    p3 = doc.add_paragraph()
    run3 = p3.add_run(f"Date: {date_str}  ·  PRIVATE & CONFIDENTIAL  ·  NON-BINDING  ·  SUBJECT TO CONTRACT")
    run3.font.size = Pt(9); run3.font.color.rgb = GREY; run3.italic = True
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(10)
    hr(doc)


def ts_section(doc, title, rows):
    """Term-sheet 2-column table."""
    add_heading(doc, title, level=1)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        row.cells[0].width = Cm(4.8)
        row.cells[1].width = Cm(11.7)
    for i, (label, detail) in enumerate(rows):
        bg = LGREY_HEX if i % 2 == 0 else "FFFFFF"
        cell_text(table.cell(i, 0), label, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(table.cell(i, 0), bg)
        cell = table.cell(i, 1); set_cell_bg(cell, bg); cell.text = ""
        for p_idx, para in enumerate(detail.split("\n")):
            p = cell.paragraphs[0] if p_idx == 0 else cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(para.strip())
            run.font.size = Pt(10); run.font.color.rgb = DARK
            if para.strip().startswith(("(a)", "(b)", "(c)", "(d)", "(e)", "(f)", "(g)",
                                        "Step", "Default rule", "Exception", "Structure",
                                        "Recommended", "Position:", "ADVANTAGE", "DISADVANTAGE")):
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def cap_table(doc, title, rows, note=None, header_bg=TEAL_HEX):
    add_heading(doc, title, level=2, gold=False)
    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(5.2), Cm(2.4), Cm(2.2), Cm(2.6), Cm(4.6)]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = widths[i]
    hdrs = ["Shareholder", "Shares", "% Holding", "Change", "Notes"]
    for i, hdr in enumerate(hdrs):
        cell_text(table.cell(0, i), hdr, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(table.cell(0, i), header_bg)
    for r_idx, row in enumerate(rows, 1):
        bg = L_GOLD_HEX if row[0].startswith("TOTAL") else (LGREY_HEX if r_idx % 2 == 0 else "FFFFFF")
        bold = row[0].startswith("TOTAL")
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
        for c_idx, (val, al) in enumerate(zip(row, aligns)):
            cell_text(table.cell(r_idx, c_idx), str(val), bold=bold, size=10, align=al)
            set_cell_bg(table.cell(r_idx, c_idx), bg)
    if note:
        p = doc.add_paragraph()
        run = p.add_run(f"★  {note}")
        run.italic = True; run.font.size = Pt(9); run.font.color.rgb = GREY
        p.paragraph_format.space_after = Pt(8)


# ═══════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.3)

# ══════════════════════════════════════════════════════════════════════════
# PART A — CAP TABLE MODEL
# ══════════════════════════════════════════════════════════════════════════
doc_header(
    doc,
    "BITWEALTH (PTY) LTD",
    "Part A — Cap Table Model & Earn-In Schedule (v3)",
    DOC_DATE,
)

add_heading(doc, "1.  Background & IP Foundation")
add_paragraph(doc,
    "BitWealth (Pty) Ltd ('the Company', Registration No. " + COMPANY_REG + ") operates a "
    "Bitcoin Dollar-Cost Averaging (DCA) investment management platform regulated under the "
    "Financial Advisory and Intermediary Services Act, 37 of 2002 (FAIS Act). The Company "
    "will be appointed as a Juristic Representative (JR) under the FSP licence of Finova "
    "Capital (Pty) Ltd (FSP No. " + FSP_NO + ").")
add_paragraph(doc,
    "The proprietary technology platform (BitWealth LTH PVR System) has been transferred to "
    "the Company under a signed IP Transfer Agreement at an agreed value of R2,500,000, "
    "contributed by " + FOUNDER + ". The asset is recorded on the Company's balance sheet.")
add_paragraph(doc,
    "The Company has 1,000 ordinary shares of R1.00 par value in issue, all carrying equal "
    "voting, dividend, and liquidation rights. " + FOUNDER + " currently holds all 1,000 shares.",
    italic=True, color=GREY, size=9)
hr(doc)

cap_table(doc, "Stage 0 — Current Shareholding (Pre-Signing)", [
    [FOUNDER, "1,000", "100%", "—", "IP contributor; technology platform"],
    [PARTNER, "Nil", "—", "—", "Not yet a shareholder"],
    ["TOTAL", "1,000", "100%", "—", ""],
], note="IP value R2,500,000 contributed by " + FOUNDER + ".", header_bg=DARK_HEX)

cap_table(doc, "Stage 1 — On Signing (Founding 10%)", [
    [FOUNDER, "900", "90%", "−100", "Technology, strategy, IP"],
    [PARTNER, "100", "10%", "+100", "Founding grant (nil consideration)"],
    ["TOTAL", "1,000", "100%", "—", ""],
], note="Founding 10% carries no performance condition.", header_bg=TEAL_HEX)

cap_table(doc, "Stage 2 — Post Milestone 1 (R50m AUM)", [
    [FOUNDER, "800", "80%", "−100", "Transfers from own holding"],
    [PARTNER, "200", "20%", "+100", "Tranche 1 earned"],
    ["TOTAL", "1,000", "100%", "—", ""],
], header_bg=GREEN_HEX)

cap_table(doc, "Stage 3 — Post Milestone 2 (R100m AUM)", [
    [FOUNDER, "700", "70%", "−100", "Transfers from own holding"],
    [PARTNER, "300", "30%", "+100", "Tranche 2 earned"],
    ["TOTAL", "1,000", "100%", "—", ""],
], header_bg=GREEN_HEX)

cap_table(doc, "Stage 4 — Post Milestone 3 (R200m AUM) — Full Earn-In", [
    [FOUNDER, "600", "60%", "−100", "Final position if all milestones met"],
    [PARTNER, "400", "40%", "+100", "Tranche 3 earned — maximum earn-in"],
    ["TOTAL", "1,000", "100%", "—", ""],
], note="Full earn-in: Simon 40%, Davin 60%.", header_bg=GOLD_HEX)

hr(doc)

# Lapse scenarios
add_heading(doc, "2.  Milestone Outcome Scenarios", level=2, gold=False)
add_paragraph(doc,
    "Resulting shareholding for each combination of milestone outcomes. Under the new "
    "near-miss rule (Clause 7), a narrowly-missed tranche vests pro-rata rather than lapsing "
    "entirely — see the worked example below.")

lapse = doc.add_table(rows=6, cols=6)
lapse.style = "Table Grid"; lapse.alignment = WD_TABLE_ALIGNMENT.CENTER
l_h = ["M1 (R50m)", "M2 (R100m)", "M3 (R200m)", "Davin", "Simon", "Outcome"]
l_w = [Cm(2.4), Cm(2.4), Cm(2.4), Cm(1.9), Cm(1.9), Cm(4.6)]
for i, col in enumerate(lapse.columns):
    for cell in col.cells:
        cell.width = l_w[i]
for i, h in enumerate(l_h):
    cell_text(lapse.cell(0, i), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(lapse.cell(0, i), TEAL_HEX)
lapse_rows = [
    ["Earned", "Earned", "Earned", "60%", "40%", "Best case — full earn-in"],
    ["Earned", "Earned", "Lapsed", "70%", "30%", "M3 missed"],
    ["Earned", "Lapsed", "Lapsed", "80%", "20%", "Only M1 achieved"],
    ["Lapsed", "Lapsed", "Lapsed", "90%", "10%", "No milestones achieved"],
    ["Lapsed", "Earned", "Earned", "70%", "30%", "M1 missed; M2 & M3 achieved"],
]
for r_idx, row in enumerate(lapse_rows, 1):
    bg = L_GOLD_HEX if r_idx == 1 else (LGREY_HEX if r_idx % 2 == 0 else "FFFFFF")
    for c_idx, val in enumerate(row):
        color = DARK
        if val == "Earned":  color = GREEN
        if val == "Lapsed":  color = RED
        cell_text(lapse.cell(r_idx, c_idx), val, size=9, color=color,
                  bold=(c_idx in [3, 4]), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(lapse.cell(r_idx, c_idx), bg)

add_paragraph(doc,
    "Simon retains his founding 10% regardless of milestone outcomes (subject only to "
    "bad-leaver provisions). Unearned tranches revert to Davin.",
    italic=True, color=GREY, size=9)
hr(doc)

# Near-miss worked example
add_heading(doc, "3.  Near-Miss / Pro-Rata Vesting (Clause 7)", level=2, gold=False)
add_paragraph(doc,
    "Worked example for Tranche 1 (R50m target): a narrowly-missed tranche now vests "
    "pro-rata after an automatic 6-month extension, instead of lapsing in full.", sa=4)
pr = doc.add_table(rows=5, cols=4)
pr.style = "Table Grid"; pr.alignment = WD_TABLE_ALIGNMENT.CENTER
pr_h = ["AUM at deadline", "% of target", "Outcome", "Tranche vested"]
pr_w = [Cm(4.0), Cm(2.8), Cm(6.0), Cm(3.2)]
for i, col in enumerate(pr.columns):
    for cell in col.cells:
        cell.width = pr_w[i]
for i, h in enumerate(pr_h):
    cell_text(pr.cell(0, i), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(pr.cell(0, i), TEAL_HEX)
pr_rows = [
    ["R50m or more", "≥100%", "Full tranche vests", "10%"],
    ["R45m", "90%", "6-mo extension, then pro-rata", "9%"],
    ["R40m", "80%", "6-mo extension, then pro-rata", "8%"],
    ["R35m", "70%", "Below 80% — tranche lapses", "0%"],
]
for r, row in enumerate(pr_rows, 1):
    bg = LGREY_HEX if r % 2 == 0 else "FFFFFF"
    for c, val in enumerate(row):
        color = DARK
        if c == 3 and val not in ("0%",): color = GREEN
        if "lapses" in val or val == "0%":  color = RED
        cell_text(pr.cell(r, c), val, size=9, color=color, bold=(c == 3),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(pr.cell(r, c), bg)

hr(doc)
# Milestone deadlines
add_heading(doc, "4.  Milestone Deadlines", level=2, gold=False)
ml = doc.add_table(rows=4, cols=5)
ml.style = "Table Grid"; ml.alignment = WD_TABLE_ALIGNMENT.CENTER
ml_h = ["Tranche", "AUM Target", "Equity", "Deadline", "Trigger Date"]
ml_w = [Cm(2.4), Cm(3.6), Cm(2.9), Cm(3.6), Cm(3.5)]
for i, col in enumerate(ml.columns):
    for cell in col.cells:
        cell.width = ml_w[i]
for i, h in enumerate(ml_h):
    cell_text(ml.cell(0, i), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(ml.cell(0, i), GOLD_HEX)
ml_rows = [
    ["Tranche 1", "R50,000,000", "+10% (→20%)", "24 months", "~July 2028"],
    ["Tranche 2", "R100,000,000", "+10% (→30%)", "48 months", "~July 2030"],
    ["Tranche 3", "R200,000,000", "+10% (→40%)", "72 months", "~July 2032"],
]
for r_idx, row in enumerate(ml_rows, 1):
    bg = LGREY_HEX if r_idx % 2 == 0 else "FFFFFF"
    for c_idx, val in enumerate(row):
        cell_text(ml.cell(r_idx, c_idx), val, size=10, bold=(c_idx == 2),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(ml.cell(r_idx, c_idx), bg)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# PART B — TERM SHEET
# ══════════════════════════════════════════════════════════════════════════
doc_header(
    doc,
    "BITWEALTH (PTY) LTD",
    "Part B — Non-Binding Heads of Terms — Business Development Equity Participation (v3)",
    DOC_DATE,
)
add_paragraph(doc,
    "This term sheet sets out the proposed principal terms of an equity participation "
    "arrangement between " + FOUNDER + " and " + PARTNER + " in respect of BitWealth (Pty) "
    "Ltd. It is non-binding, subject to contract, and for discussion purposes only. Final "
    "terms will be recorded in a Shareholders' Agreement and Equity Participation Agreement "
    "drafted by the Company's attorneys.")

p = doc.add_paragraph()
run = p.add_run("NOT LEGAL ADVICE.  ")
run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RED
run2 = p.add_run(
    "This document is negotiation support only. It must be reviewed, and the definitive "
    "agreements drafted, by a qualified SA attorney, with FAIS input from Finova Capital and "
    "independent tax advice for both parties, before signature.")
run2.font.size = Pt(9); run2.font.color.rgb = GREY; run2.italic = True
hr(doc)

ts_section(doc, "1.  Parties", [
    ("Founder", FOUNDER + ", ID " + FOUNDER_ID + ", Republic of South Africa ('Davin')."),
    ("Partner", PARTNER + ", ID " + PARTNER_ID + ", Republic of South Africa ('Simon')."),
    ("Company", "BitWealth (Pty) Ltd (Registration No. " + COMPANY_REG + "), a private company "
                "incorporated in the Republic of South Africa."),
    ("FSP", "Finova Capital (Pty) Ltd (FSP No. " + FSP_NO + ") — BitWealth to be appointed a "
            "Juristic Representative under Finova's FSP licence in terms of the FAIS Act, 2002."),
])

ts_section(doc, "2.  Company & IP Foundation", [
    ("Business", "Bitcoin DCA investment management for retail and HNWI clients, using the "
                 "Company's proprietary LTH PVR algorithm executed via the VALR exchange."),
    ("IP Asset", "The BitWealth LTH PVR platform, transferred to the Company by " + FOUNDER +
                 " under a signed IP Transfer Agreement; recorded on the balance sheet at an "
                 "agreed value of R2,500,000."),
    ("IP Evidence Pack",
     "Before signing, the Company will make available:\n"
     "(a) the executed IP Transfer Agreement;\n"
     "(b) developer / contractor IP-assignment confirmations;\n"
     "(c) a list of third-party and open-source dependencies and their licences;\n"
     "(d) a warranty of title and no known competing claims."),
    ("Valuation Basis", "R2,500,000 is the agreed commercial consideration, not an independent "
                        "valuation. Davin warrants ownership and title, not market value. Simon "
                        "may commission an independent valuation at his own cost."),
    ("Share Capital", "1,000 ordinary shares of R1.00 par value in issue, all with equal voting, "
                      "dividend, and liquidation rights. Davin currently holds all 1,000 shares."),
])

ts_section(doc, "3.  Simon's Role & Contribution", [
    ("Role", "Non-executive Head of Business Development (consulting engagement), with defined "
             "quarterly client-acquisition deliverables agreed by the Board."),
    ("Capital", "NIL — Simon contributes no monetary capital. His additional equity is earned "
                "solely through the AUM milestones in Clause 6."),
    ("Responsibilities",
     "(a) Origination and onboarding of new clients;\n"
     "(b) management and growth of client relationships;\n"
     "(c) achievement of the AUM milestones;\n"
     "(d) other business-development activities agreed by the Board."),
    ("Support & Budget", "The Company will provide reasonable sales, marketing, and "
                         "compliance/admin support and an annual marketing budget approved by "
                         "the Board. Obligations on both sides are defined, not open-ended."),
    ("Time Commitment", "An agreed minimum monthly time commitment, recorded in the engagement letter."),
    ("FAIS Accreditation", "Simon will use reasonable endeavours to obtain and maintain the FAIS "
                          "accreditation required for his role. Delays caused by the FSCA, Finova, "
                          "paperwork, or the Company are excluded. Only a wilful or negligent "
                          "failure by Simon constitutes a breach."),
])

ts_section(doc, "4.  Introduced Client Register", [
    ("Register", "The Company maintains an 'Introduced Client Register'. Each client Simon "
                 "introduces is logged within 5 business days of first meaningful contact."),
    ("Definition", "An 'Introduced Client' is one where Simon makes the first meaningful contact "
                   "and who is not already in the Company's active pipeline. Each client is "
                   "credited once. Disputed entries are resolved by the Board acting reasonably, "
                   "failing which by independent expert."),
    ("Conversion Tail", "An Introduced Client who converts within 18 months of introduction, or "
                        "within 12 months of Simon's exit, counts toward milestones and any "
                        "applicable revenue share (Clause 8)."),
])

ts_section(doc, "5.  Founding Equity Grant (Initial 10%)", [
    ("Grant", "On signing, Davin transfers 100 ordinary shares (10%) to Simon for nil "
              "consideration. No performance condition attaches to the founding 10%."),
    ("Notional Value", "At the IP-implied price of R2,500 per share, the notional value is "
                       "R250,000. This is not a guaranteed or market valuation."),
    ("Tax Advice", "Both parties will obtain independent tax advice before any nil- or "
                   "below-value transfer, acknowledging potential income-tax, CGT, and "
                   "donations-tax consequences of sweat equity, to be cleared first."),
    ("Effective Date", "Within 5 business days of signature of the Shareholders' Agreement."),
])

ts_section(doc, "6.  Earn-In Milestones (Additional 30%)", [
    ("Mechanism", "Simon may earn up to a further 30% in three equal 10% tranches (100 shares "
                  "each), each a transfer of existing shares from Davin. No new shares are issued."),
    ("Tranche 1", "10% on verified total AUM of R50,000,000 or more.\n"
                  "Deadline: 24 months from the Commencement Date (~July 2028)."),
    ("Tranche 2", "10% on verified total AUM of R100,000,000 or more.\n"
                  "Deadline: 48 months from the Commencement Date (~July 2030)."),
    ("Tranche 3", "10% on verified total AUM of R200,000,000 or more.\n"
                  "Deadline: 72 months from the Commencement Date (~July 2032)."),
    ("Objective Verification", "A milestone is achieved when total AUM meets the target at a "
                              "month-end AND remains at or above it to the next month-end (two "
                              "consecutive month-ends), verified objectively from platform, "
                              "custodian, exchange, and accounting records. This prevents a "
                              "one-day spike from triggering a permanent equity transfer."),
    ("Administrative Confirmation", "Board confirmation is administrative only and may not be "
                                   "unreasonably withheld or delayed. If the evidence shows the "
                                   "threshold met and sustained and the Board has not confirmed "
                                   "within 10 business days, the milestone is deemed achieved."),
    ("Sequence", "Milestones are achieved in sequence; a later tranche cannot vest before the "
                 "earlier tranche is earned (or its shortfall resolved under Clause 7)."),
    ("Maximum / Minimum", "Full earn-in: Davin 60%, Simon 40%. Simon retains his founding 10% "
                          "regardless of milestone outcome (subject to Clause 9)."),
])

ts_section(doc, "7.  Milestone Shortfall — Near-Miss Vesting", [
    ("Replaces v2 options", "This unified mechanism replaces the four options in v2 and applies "
                            "equally to all three tranches."),
    ("Step 1 — Extension", "If AUM reaches at least 80% of the target by the deadline, Simon "
                          "automatically receives a single 6-month extension for that tranche."),
    ("Step 2 — Pro-rata vesting", "If, at the end of the extension, AUM is still below target but "
                                 "at least 80% of it, Simon vests a pro-rata portion of that "
                                 "tranche = (AUM achieved ÷ target) × 10%. Example: R45m of R50m "
                                 "(90%) → Simon vests 9%; the residual 1% remains with Davin."),
    ("Step 3 — Lapse", "If AUM is below 80% of target at the end of the extension, the tranche "
                       "lapses in full and the shares remain with Davin. Simon retains all "
                       "previously-earned tranches and his founding 10%."),
])

doc.add_page_break()

# ── Clause 8 — NEW VENTURES & INNOVATION REVENUE (expanded, with worked example) ──
add_heading(doc, "8.  New Ventures & Innovation Revenue")
add_paragraph(doc,
    "This clause addresses Simon's request for a revenue share on new products or ideas that "
    "generate new revenue. It balances (i) Davin's preference that genuinely new products are "
    "built in a separate company on fresh terms; (ii) fairness to Simon where he personally "
    "originates and runs a new, non-AUM revenue stream; and (iii) protection against Simon "
    "receiving extra revenue share on revenue generated by clients that Davin or others "
    "introduce.", sa=6)

add_heading(doc, "The core principle", level=2, gold=False)
add_paragraph(doc,
    "Simon is already a shareholder, so through dividends and equity value he already shares "
    "in ALL Company revenue — including AUM revenue from clients introduced by anyone. The "
    "framework therefore separates two different rewards: 'ownership reward' (his equity "
    "dividend) from 'origination reward' (a ring-fenced share of only the specific new stream "
    "he personally creates and runs). The ring-fence never touches AUM revenue.", sa=6)

ts_section(doc, "8.1  Default Rule — New Ventures go into a NewCo", [
    ("Default rule", "Any materially new product, business line, or revenue stream that is "
                     "separable from the core BitWealth AUM business is developed in a NEW "
                     "company ('NewCo'), with a fresh shareholding structure negotiated "
                     "specifically for that venture."),
    ("No automatic entitlement", "Simon has NO automatic right to equity or revenue in any NewCo "
                                 "by virtue of his BitWealth shareholding. Each NewCo is "
                                 "negotiated on its own merits (who originates, funds, builds, runs)."),
    ("Right to negotiate", "Where Simon materially originates a NewCo concept, he has a "
                          "good-faith right of first opportunity to negotiate participation — on "
                          "fresh terms, not his BitWealth terms."),
    ("Who decides", "The Board (with Davin's control vote) decides whether a new idea is pursued "
                   "inside BitWealth or via a NewCo."),
])

ts_section(doc, "8.2  Exception — Simon-originated non-AUM revenue kept inside BitWealth", [
    ("When this applies", "Only where ALL of the following are true:\n"
                          "(a) Simon personally originates a genuinely new, NON-AUM revenue "
                          "stream (not performance, platform, or exchange fees on managed assets);\n"
                          "(b) the Board decides to keep it inside BitWealth; and\n"
                          "(c) Simon actively leads the delivery of that stream."),
    ("Innovation Revenue Share", "Simon receives a ring-fenced share of the NET revenue of THAT "
                                "specific stream only:\n"
                                "(a) a negotiated rate (see options in 8.3) of net stream revenue;\n"
                                "(b) for a fixed period (suggested 48 months from first revenue);\n"
                                "(c) payable only while Simon actively leads the stream;\n"
                                "(d) the stream is separately accounted so the ring-fence is auditable."),
    ("Ring-fence — key protection", "The Innovation Revenue Share attaches ONLY to the new "
                                    "non-AUM stream Simon creates. It NEVER touches AUM-based "
                                    "revenue (performance, platform, or exchange fees) from any "
                                    "client — including clients Davin or others introduce."),
    ("IP ownership", "All IP in the new stream vests in the Company (or NewCo), not in Simon "
                     "personally, regardless of the revenue share."),
    ("Exclusions", "Does NOT apply to enhancements or extensions of the core AUM business (new "
                   "client tiers, pricing changes, new asset pairs) — those are core revenue in "
                   "which Simon already shares as a shareholder."),
])

# ── The clarification the user asked for: does it STACK? Worked example ──
add_heading(doc, "8.3  Does Simon get the ring-fenced share PLUS his equity dividend? — Yes, and here is why it matters", level=2, gold=False)
add_paragraph(doc,
    "This is the key question. If the new stream is kept INSIDE BitWealth, then under the "
    "drafted approach Simon receives BOTH: (1) the ring-fenced Innovation Revenue Share off "
    "the top, AND (2) his normal equity-based dividend share of whatever profit remains after "
    "that. The two stack, because the residual profit of the stream is ordinary Company profit "
    "that flows to all shareholders by their equity %.", sa=4)
add_paragraph(doc,
    "Worked example. Assume the new stream produces R1,000,000 net revenue in a year, Simon "
    "holds 40% equity (full earn-in), a 15% Innovation Revenue Share, and no other costs, all "
    "residual profit distributed as dividends:", sa=4, bold=True)

we = doc.add_table(rows=5, cols=5)
we.style = "Table Grid"; we.alignment = WD_TABLE_ALIGNMENT.CENTER
we_h = ["Simon's equity", "Innovation share (15%)", "Equity share of residual 85%",
        "Simon's total from stream", "Effective %"]
we_w = [Cm(2.6), Cm(3.2), Cm(3.6), Cm(3.4), Cm(2.2)]
for i, col in enumerate(we.columns):
    for cell in col.cells:
        cell.width = we_w[i]
for i, h in enumerate(we_h):
    cell_text(we.cell(0, i), h, bold=True, size=8.5, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(we.cell(0, i), TEAL_HEX)
# net = 1,000,000; innovation = 150,000; residual = 850,000
we_rows = [
    ["10%", "R150,000", "R85,000",  "R235,000", "23.5%"],
    ["20%", "R150,000", "R170,000", "R320,000", "32.0%"],
    ["30%", "R150,000", "R255,000", "R405,000", "40.5%"],
    ["40%", "R150,000", "R340,000", "R490,000", "49.0%"],
]
for r, row in enumerate(we_rows, 1):
    bg = L_GOLD_HEX if r == len(we_rows) else (LGREY_HEX if r % 2 == 0 else "FFFFFF")
    for c, val in enumerate(row):
        cell_text(we.cell(r, c), val, size=9, bold=(c in (3, 4)),
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(we.cell(r, c), bg)

add_paragraph(doc,
    "Reading the last row: at 40% equity, Simon receives R150,000 (the 15% innovation share) "
    "PLUS R340,000 (his 40% equity share of the remaining R850,000) = R490,000, or 49% of the "
    "whole R1,000,000 stream. Davin receives the other R510,000 (51%). The formula is: "
    "Simon's effective share = innovation% + equity% × (1 − innovation%).", sa=4, italic=True, color=GREY, size=9)

add_paragraph(doc, "Three ways to structure this — choose one when negotiating:", bold=True, sa=4)

ts_section(doc, "8.4  Structural Options for the Inside-BitWealth Case", [
    ("Structure A — Stacked (as illustrated above)",
     "Simon gets the 15% innovation share PLUS equity dividend on the residual. His effective "
     "share of the stream rises with his equity (23.5% at 10% → 49% at 40%).\n"
     "ADVANTAGE: simple; strongly rewards Simon's origination and effort.\n"
     "DISADVANTAGE: Simon captures more than his equity % of that stream; most generous to Simon."),
    ("Structure B — Lower rate, eyes open",
     "Keep stacking but set a LOWER innovation rate (e.g., 5–10%) precisely because Simon also "
     "enjoys equity upside. At 10% equity + 5% innovation, effective share ≈ 14.5%; at 40% "
     "equity + 5% innovation ≈ 43%.\n"
     "ADVANTAGE: keeps Simon closer to his equity %; still rewards origination.\n"
     "DISADVANTAGE: less upside for Simon; still stacks."),
    ("Structure C — NewCo (Davin's default preference)",
     "Put the new stream in a NewCo and negotiate a single, clean split (e.g., Simon 50% / "
     "Davin 50%) that IS his whole entitlement. No stacking, no ring-fence accounting.\n"
     "ADVANTAGE: cleanest; avoids all double-benefit ambiguity; each venture stands alone.\n"
     "DISADVANTAGE: administrative cost of a second company."),
    ("Recommended", "Default to Structure C (NewCo) for anything separable. If a non-AUM stream "
                    "must live inside BitWealth, use Structure B (a modest 5–10% innovation rate) "
                    "and negotiate with full awareness that it stacks on Simon's equity dividend."),
])

ts_section(doc, "8.5  Reciprocity & Anti-Circumvention", [
    ("Reciprocal", "The same framework applies to Davin- or third-party-originated ventures: "
                   "Simon has no revenue-share claim over a non-AUM stream he did not originate "
                   "and does not lead."),
    ("Anti-circumvention", "Neither party may re-label core BitWealth AUM revenue as a 'new "
                          "stream' to avoid the other's legitimate share. Characterisation "
                          "disputes go to independent expert determination."),
])

doc.add_page_break()

# ── Protective provisions ──
ts_section(doc, "9.  FSP Arrangement, Fee Transparency & Reserved Matters", [
    ("FSP Terms", "Finova Capital (FSP No. " + FSP_NO + "): 20% of performance fees + 50% of "
                  "platform fees payable to Finova; no fixed monthly hosting fee."),
    ("Fee Transparency", "The Company will share a full fee/profitability model showing net "
                        "economics after the Finova revenue share."),
    ("Reserved Matters", "A 10% holder will NOT have a veto over operational/strategic decisions "
                        "(FSP, exchange, custody, execution). Instead:\n"
                        "(a) consultation rights on material changes;\n"
                        "(b) an anti-prejudice covenant — no change may be made for the purpose "
                        "of, or with the effect of, reducing Simon's economic entitlement below "
                        "its then-current basis without his consent;\n"
                        "(c) genuine reserved-matter consent rights that apply only once Simon "
                        "holds 25% or more."),
    ("Regulatory Risk", "Simon acknowledges the JR arrangement is subject to regulatory approval "
                       "and ongoing compliance; loss of JR status would materially impact the "
                       "Company."),
])

ts_section(doc, "10.  Good Leaver & Bad Leaver", [
    ("Good Leaver — expanded", "Simon is a Good Leaver on: death/permanent disability; "
                              "retirement; termination without cause; constructive dismissal; "
                              "material reduction of his role; the Company/Davin failing to "
                              "support onboarding; loss of FSP/JR status not caused by Simon; or "
                              "material breach by Davin/the Company.\n"
                              "Consequence: retains all earned shares; unearned tranches follow Clause 7."),
    ("Bad Leaver — narrowed", "Limited to proven fraud, theft, dishonesty, wilful misconduct, "
                            "gross negligence, or a final (non-appealable) regulatory finding "
                            "against Simon."),
    ("Bad Leaver consequence", "Unearned tranches lapse. EARNED shares are bought at fair market "
                             "value, EXCEPT where the trigger is fraud/theft/dishonesty (then at "
                             "par, R1/share). Any forced transfer requires prior independent "
                             "determination by an auditor/expert or arbitrator."),
    ("Clawback — narrowed + reciprocal", "Clawback applies only to losses caused by Simon's "
                                        "wilful misconduct, fraud, gross negligence, or proven "
                                        "regulatory breach. Reciprocally, if Davin's/the "
                                        "Company's acts or omissions cause loss of JR/FSP status "
                                        "or damage client relationships, Simon is treated as a "
                                        "Good Leaver."),
])

ts_section(doc, "11.  Confidentiality, Non-Solicitation, Drag & Tag", [
    ("Non-compete removed", "The broad non-compete is removed. Replaced with (i) confidentiality "
                          "and (ii) non-solicitation of clients Simon actually dealt with and of "
                          "staff, for 12 months post-exit, within South Africa."),
    ("Drag-Along — with protections", "On a bona-fide sale of ≥ 75%, Simon must sell pro-rata on "
                                     "the SAME price and terms, with: cash-equivalent "
                                     "consideration; warranties limited to title and capacity; "
                                     "release from guarantees/liabilities; and no sale to a "
                                     "related party at an undervalue."),
    ("Tag-Along", "If Davin sells more than 50% of his shares, Simon may include his shares on "
                 "equivalent terms and the same price per share."),
])

ts_section(doc, "12.  Pre-emptive Rights & Anti-Circumvention", [
    ("Nature clarified", "Simon has PRO-RATA pre-emptive rights (a right of first offer on new "
                        "share issues at the same price as the incoming subscriber) — not a full "
                        "ratchet or weighted-average anti-dilution."),
    ("Anti-circumvention", "No shares or options may be issued to Davin, associates, or staff "
                          "below independently-determined fair value for the purpose of diluting "
                          "Simon, without his consent."),
])

ts_section(doc, "13.  Share Transfer Mechanism", [
    ("Transfer Method", "By share transfer form and update of the Company's share register, "
                       "filed with CIPC where required under the Companies Act, 2008."),
    ("Trigger Events", "(a) Signing (founding 10%); (b) Board confirmation of milestone "
                      "achievement; or (c) a leaver event (Clause 10)."),
    ("Transfer Costs", "All costs, including any securities transfer tax, borne by the Company."),
])

ts_section(doc, "14.  Conditions Precedent & General", [
    ("Conditions Precedent", "(a) Shareholders' Agreement and Equity Participation Agreement in "
                           "agreed form;\n"
                           "(b) independent legal advice obtained by Simon;\n"
                           "(c) independent tax advice obtained by both parties;\n"
                           "(d) Finova Capital JR appointment or letter of intent;\n"
                           "(e) confirmation Simon meets minimum FAIS fit-and-proper requirements."),
    ("Exclusivity", "For 60 days from signature, Simon will not negotiate a similar arrangement "
                   "with any other fintech, crypto, or investment-management company."),
    ("Governing Law", "Republic of South Africa; disputes to arbitration under AFSA rules."),
    ("Confidentiality", "Both parties keep the contents of this term sheet strictly confidential."),
    ("Non-Binding", "Non-binding and subject to contract; no obligations arise until definitive "
                   "agreements are signed. Lapses if not countersigned within 30 days."),
])

hr(doc)

# ── Signature ──
add_heading(doc, "15.  Signature")
add_paragraph(doc,
    "By signing below the parties confirm they have read and understood these heads of terms "
    "and agree to proceed to definitive agreements on this basis. This signature does not "
    "create a binding legal obligation.")

sig = doc.add_table(rows=6, cols=2)
sig.style = "Table Grid"
for row in sig.rows:
    row.cells[0].width = Cm(8.2)
    row.cells[1].width = Cm(8.2)
sig_rows = [
    ("FOR AND ON BEHALF OF BITWEALTH (PTY) LTD\nand in his personal capacity as Founder",
     "AGREED TO BY " + PARTNER.upper() + "\nin his personal capacity"),
    ("Signature: _____________________________", "Signature: _____________________________"),
    ("Full Name:  " + FOUNDER, "Full Name:  " + PARTNER),
    ("Capacity:   Director / Founder", "Capacity:   Proposed Business Development Partner"),
    ("Date:       ___________________________", "Date:       ___________________________"),
    ("Witness:    ___________________________", "Witness:    ___________________________"),
]
for r_idx, (left, right) in enumerate(sig_rows):
    for c_idx, val in enumerate([left, right]):
        cell_text(sig.cell(r_idx, c_idx), val, bold=(r_idx == 0), size=10,
                  align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(sig.cell(r_idx, c_idx), LGREY_HEX if r_idx == 0 else "FFFFFF")

out_path = OUT_DIR / "BitWealth_Simon_Equity_Proposal_v3.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
