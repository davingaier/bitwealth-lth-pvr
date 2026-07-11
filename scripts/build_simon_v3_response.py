"""
BitWealth (Pty) Ltd — Simon Hobday Equity Proposal v3 (Response to Comments)
Output: docs/Shareholding/BitWealth_Simon_Equity_Proposal_v3.docx

Single combined document with four parts:
  Part A — Attorney's Response to Reviewer Comments (all 17 comments addressed)
  Part B — Revised Term Sheet (clauses updated per agreed positions)
  Part C — New Ventures & Innovation Revenue framework
  Part D — Revised Cap Table Model & Earn-In Schedule (with near-miss vesting)

NOTE: not legal advice — for negotiation support only; final agreements to be
drafted/reviewed by a qualified SA attorney (FAIS, Companies Act, tax).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

OUT_DIR = Path("docs") / "Shareholding"
OUT_DIR.mkdir(parents=True, exist_ok=True)

try:
    TODAY = date.today().strftime("%d %B %Y").lstrip("0")
except Exception:
    TODAY = "11 July 2026"

# Colours
GOLD  = RGBColor(0xC9, 0xA2, 0x27)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x1B, 0x5E, 0x20)
RED   = RGBColor(0xB7, 0x1C, 0x1C)
GOLDD = RGBColor(0x8A, 0x6D, 0x0F)

L_GOLD_HEX = "F5EBC8"
GOLD_HEX   = "C9A227"
DARK_HEX   = "1A1A1A"
TEAL_HEX   = "006064"
RED_HEX    = "B71C1C"
GREEN_HEX  = "1B5E20"
LGREY_HEX  = "F5F5F5"
LGREEN_HEX = "E8F5E9"
LGOLD2_HEX = "FFF8E1"
LRED_HEX   = "FDECEA"


# ═══════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, size=10, color=None, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else DARK
    return run


def add_heading(doc, text, level=1, gold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper() if level == 1 else text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = GOLD if gold else DARK
    run.font.name = "Calibri"
    p.paragraph_format.keep_with_next = True
    return p


def add_paragraph(doc, text, size=10, color=None, italic=False, sb=2, sa=4, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    run = p.add_run(text)
    run.font.size = Pt(size); run.bold = bold
    run.font.color.rgb = color if color else DARK
    run.italic = italic
    return p


def add_bullet(doc, text, size=10, indent=0.6):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.color.rgb = DARK
    return p


def hr(doc, color=GOLD_HEX):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), color)
    pBdr.append(bottom); pPr.append(pBdr)
    return p


def doc_header(doc, title, subtitle, date_str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = GOLD; run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(11); run2.font.color.rgb = DARK
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)

    p3 = doc.add_paragraph()
    run3 = p3.add_run(f"Date: {date_str}  ·  PRIVATE & CONFIDENTIAL  ·  NON-BINDING  ·  SUBJECT TO CONTRACT")
    run3.font.size = Pt(9); run3.font.color.rgb = GREY; run3.italic = True
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(10)
    hr(doc)


def ts_section(doc, title, rows):
    """Term-sheet style 2-column table. rows = [(label, detail_with_optional_newlines)]"""
    add_heading(doc, title, level=1)
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        row.cells[0].width = Cm(5.0)
        row.cells[1].width = Cm(11.5)
    for i, (label, detail) in enumerate(rows):
        bg = LGREY_HEX if i % 2 == 0 else "FFFFFF"
        cell_text(table.cell(i, 0), label, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_bg(table.cell(i, 0), bg)
        cell = table.cell(i, 1); set_cell_bg(cell, bg); cell.text = ""
        for p_idx, para in enumerate(detail.split("\n")):
            p = cell.paragraphs[0] if p_idx == 0 else cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(para.strip())
            run.font.size = Pt(10); run.font.color.rgb = DARK
            if para.strip().startswith(("(a)", "(b)", "(c)", "(d)", "(e)", "(f)",
                                        "Option", "Default rule", "Exception",
                                        "Recommended", "Position:")):
                run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = Cm(2.0)
    section.left_margin = section.right_margin = Cm(2.3)

doc_header(
    doc,
    "BITWEALTH (PTY) LTD",
    "Revised Equity Proposal v3 — Response to Simon Hobday's Comments",
    TODAY,
)

# ── Preamble / disclaimer ──────────────────────────────────────────────────
add_paragraph(doc,
    "This document responds to the seventeen (17) comments raised by Simon Hobday on "
    "the v2 Equity Proposal dated June 2026, and introduces a framework for new products "
    "and non-AUM revenue. It is structured in four parts:", sa=4)
add_bullet(doc, "Part A — Point-by-point response to each of Simon's comments, with the "
                "Company's negotiating position and rationale.")
add_bullet(doc, "Part B — Revised Term Sheet clauses giving effect to the agreed positions.")
add_bullet(doc, "Part C — New Ventures & Innovation Revenue framework (Simon's new-revenue request).")
add_bullet(doc, "Part D — Revised Cap Table Model, including near-miss / pro-rata vesting.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run("IMPORTANT — NOT LEGAL ADVICE.  ")
run.bold = True; run.font.size = Pt(9); run.font.color.rgb = RED
run2 = p.add_run(
    "This document is prepared as negotiation support only and does not constitute legal, "
    "regulatory, or tax advice. It must be reviewed and the definitive agreements drafted by "
    "a qualified South African attorney, with FAIS input from Finova Capital and independent "
    "tax advice for both parties, before signature.")
run2.font.size = Pt(9); run2.font.color.rgb = GREY; run2.italic = True

hr(doc)

# ══════════════════════════════════════════════════════════════════════════
# PART A — RESPONSE TO COMMENTS
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART A — Response to Simon's Comments")
add_paragraph(doc,
    "Position key:  ACCEPT = agreed as requested.  ACCEPT (MODIFIED) = agreed in principle "
    "with adjustments to keep the deal balanced.  COUNTER = alternative proposed. "
    "Of the 17 comments, the Company accepts 9 outright, accepts 6 with modification, and "
    "counters 2.", italic=True, color=GREY, size=9, sa=6)

# Comments data: (num, topic, simon_point, position, position_color_hex, attorney_note)
COMMENTS = [
    (1, "IP valuation evidence",
     "Wants independent evidence of the R2.5m IP value, proof of ownership, developer "
     "assignments, dependency list, and no competing claims.",
     "ACCEPT (MODIFIED)", LGOLD2_HEX,
     "Fair diligence. The Company will provide a disclosure pack: signed IP transfer "
     "agreement, developer/contractor assignment confirmations, an open-source/third-party "
     "dependency list, and a warranty of title with no known competing claims. However, the "
     "R2.5m is an agreed commercial figure, not an independent valuation — Davin warrants "
     "ownership and title, not the market value. An independent valuation can be commissioned "
     "at Simon's cost if he requires one."),
    (2, "Introduced Client Register",
     "Wants a formal register; any party introduced by Simon recorded within 5 business days "
     "and credited to him for milestone and revenue-share purposes.",
     "ACCEPT (MODIFIED)", LGOLD2_HEX,
     "Agreed, with guardrails: (i) an 'Introduced Client' is one where Simon makes the first "
     "meaningful contact and the client is not already in the Company's pipeline; (ii) each "
     "client is credited once; (iii) entries are logged within 5 business days and are open to "
     "reasonable dispute/verification by the Board. This protects both sides from double-counting."),
    (3, "FAIS 3-month hard deadline",
     "Says the 3-month FAIS compliance breach clause is too harsh where delays are caused by "
     "Finova, FSCA, paperwork, or the Company.",
     "ACCEPT", LGREEN_HEX,
     "Agreed. Changed to a 'reasonable endeavours' obligation, expressly excluding delays "
     "outside Simon's control (regulator, FSP, or Company-caused). Only a wilful or negligent "
     "failure by Simon to pursue accreditation will count as a breach."),
    (4, "Define role scope",
     "Wants defined time commitment, role scope, sales/marketing/compliance support, budget, "
     "and whether this is employment, consulting, or non-exec.",
     "ACCEPT", LGREEN_HEX,
     "Agreed and sensible for both parties. Proposed: a non-executive business-development "
     "consulting role with defined quarterly deliverables, an agreed minimum time commitment, "
     "and a marketing/sales budget set annually by the Board. Avoids open-ended obligations "
     "on Simon and open-ended cost on the Company."),
    (5, "Objective AUM verification",
     "Wants milestones deemed achieved automatically once month-end AUM meets the threshold, "
     "confirmed by platform/custodian records; Board confirmation administrative only.",
     "ACCEPT (MODIFIED)", LGOLD2_HEX,
     "Agreed, with one safeguard: AUM must meet the threshold at a month-end and be sustained "
     "to the following month-end (two consecutive month-ends) to avoid a one-day spike "
     "triggering a permanent equity transfer. Verified objectively from platform, custodian, "
     "exchange, and accounting records. Board role is administrative confirmation only."),
    (6, "Board cannot withhold confirmation",
     "Wants a clause that Board confirmation may not be unreasonably withheld; deemed achieved "
     "if not confirmed within 10 business days despite evidence.",
     "ACCEPT", LGREEN_HEX,
     "Agreed — consistent with making verification objective. If the evidence shows the "
     "threshold met and sustained, and the Board does not confirm within 10 business days, "
     "the milestone is deemed achieved."),
    (7, "Near-miss protection",
     "Says pure lapse is too founder-friendly; wants pro-rata vesting, an extension, or "
     "revenue-share if 80%+ of a milestone is reached by deadline.",
     "ACCEPT (MODIFIED)", LGOLD2_HEX,
     "Agreed in principle via a new shortfall mechanism (Part B, Clause 7): if AUM reaches "
     "≥ 80% of the target by the deadline, Simon receives an automatic single 6-month "
     "extension; if still short after the extension but ≥ 80%, he vests a pro-rata portion of "
     "that tranche equal to (AUM achieved ÷ target) × tranche %. Below 80%, the tranche lapses. "
     "This rewards genuine near-misses without giving away equity for material underperformance."),
    (8, "Extend Option C to all tranches",
     "Counter-proposal that the one-time extension apply to all tranches, plus pro-rata "
     "recognition for narrowly-missed milestones.",
     "ACCEPT", LGREEN_HEX,
     "Agreed and now superseded by the single unified shortfall mechanism in Clause 7, which "
     "applies equally to all three tranches (extension + pro-rata vesting)."),
    (9, "Fee model + reserved matters",
     "Wants a full post-Finova profitability model, and consent rights (reserved matters) over "
     "FSP/JR changes, platform fees, performance fees, and custody/execution once he holds equity.",
     "COUNTER", LRED_HEX,
     "Fee transparency: ACCEPTED — the Company will share the post-Finova fee/profitability "
     "model. Reserved-matters veto: COUNTERED. A 10% holder should not have a veto over "
     "operational and strategic decisions (choice of FSP, exchange, custody). Instead the "
     "Company offers: (i) consultation rights on material changes; (ii) an anti-prejudice "
     "covenant — no change may be made for the purpose of, or with the effect of, reducing "
     "Simon's economic entitlement below its then-current basis without his consent; and "
     "(iii) genuine reserved-matter consent rights that switch on only once Simon holds ≥ 25%. "
     "This protects Simon's economics without handing a minority a veto over the business."),
    (10, "Expand Good Leaver",
     "Wants Good Leaver expanded to include constructive dismissal, material role reduction, "
     "Company failure to support, loss of FSP/JR not caused by Simon, or Company/Davin breach.",
     "ACCEPT", LGREEN_HEX,
     "Agreed — all reasonable and now included. If Simon leaves for any of these reasons he is "
     "a Good Leaver and retains earned shares."),
    (11, "Bad Leaver too severe",
     "Says R1 buyback of ALL shares is too severe; wants Bad Leaver limited to proven "
     "fraud/dishonesty/wilful misconduct/gross negligence/final regulatory breach, earned "
     "shares at fair value except in fraud, and independent determination first.",
     "ACCEPT (MODIFIED)", LGOLD2_HEX,
     "Largely agreed. Bad Leaver is now limited to proven fraud, theft, dishonesty, wilful "
     "misconduct, gross negligence, or a final (non-appealable) regulatory finding against "
     "Simon. Consequence: unearned tranches lapse; EARNED shares are bought at fair market "
     "value EXCEPT where the trigger is fraud/theft/dishonesty, in which case earned shares are "
     "bought at par. Any forced transfer requires prior independent determination "
     "(auditor/expert or arbitrator)."),
    (12, "Non-compete too broad",
     "Wants the broad non-compete replaced with confidentiality + non-solicitation, and any "
     "restraint limited to clients he dealt with, 12 months, South Africa.",
     "ACCEPT", LGREEN_HEX,
     "Agreed — and legally prudent (SA restraints must be reasonable to be enforceable). The "
     "broad non-compete is replaced with (i) confidentiality, (ii) non-solicitation of clients "
     "Simon actually dealt with and of staff, for 12 months post-exit, within South Africa."),
    (13, "Drag-along minority protections",
     "Wants same price/terms, cash-equivalent consideration, warranties limited to "
     "title/capacity, release from guarantees, and no related-party undervalue sale.",
     "ACCEPT", LGREEN_HEX,
     "Agreed — standard and fair minority protections on a drag. All included in Clause 10."),
    (14, "Limit clawback + reciprocity",
     "Wants clawback limited to Simon's wilful misconduct/fraud/gross negligence/proven "
     "regulatory breach, plus reciprocal protection if Davin/Company compromises the JR/FSP.",
     "ACCEPT", LGREEN_HEX,
     "Agreed. Clawback is narrowed to those causes, and a reciprocal Good-Leaver trigger is "
     "added: if Davin's or the Company's acts/omissions cause loss of JR/FSP status or damage "
     "client relationships, Simon is treated as a Good Leaver and his earned equity is protected."),
    (15, "Anti-dilution clarification",
     "Notes the 'anti-dilution' clause is really a pre-emptive right; wants clarity and a ban "
     "on issuing cheap shares/options to Davin/associates/staff to dilute him.",
     "ACCEPT (MODIFIED)", LGOLD2_HEX,
     "Clarified: Simon has PRO-RATA pre-emptive rights (a right of first offer on new issues at "
     "the same price), not a full ratchet or weighted-average anti-dilution. Added an "
     "anti-circumvention protection: no shares or options may be issued to Davin, associates, "
     "or staff below independently-determined fair value for the purpose of diluting Simon, "
     "without his consent."),
    (16, "No-circumvention / tail",
     "Wants introductions that convert within 24-36 months after introduction or termination "
     "to count toward milestones and/or revenue share.",
     "ACCEPT (MODIFIED)", LGOLD2_HEX,
     "Agreed with a shorter tail: genuine Introduced Clients (per Clause 4) who convert within "
     "18 months of introduction, or within 12 months of Simon's exit, count toward milestones "
     "and any applicable revenue share. A 24-36 month tail is too long and would burden the "
     "Company well after Simon's involvement ends."),
    (17, "Tax advice on sweat equity",
     "Wants independent tax advice before any nil-consideration share transfer, as the 10% "
     "sweat equity may have tax consequences.",
     "ACCEPT", LGREEN_HEX,
     "Agreed — and important for both parties. Independent tax advice is now a condition "
     "precedent. Nil-consideration and below-value share transfers can trigger income-tax, "
     "CGT, and donations-tax questions that must be cleared first."),
]

# Build the comment response table
tbl = doc.add_table(rows=len(COMMENTS) + 1, cols=4)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ["#", "Simon's Comment", "Position", "Company's Response & Rationale"]
widths = [Cm(0.8), Cm(4.5), Cm(2.6), Cm(9.0)]
for i, col in enumerate(tbl.columns):
    for cell in col.cells:
        cell.width = widths[i]
for i, h in enumerate(hdrs):
    cell_text(tbl.cell(0, i), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl.cell(0, i), TEAL_HEX)

for r, (num, topic, point, pos, pos_bg, note) in enumerate(COMMENTS, 1):
    row_bg = LGREY_HEX if r % 2 == 0 else "FFFFFF"
    cell_text(tbl.cell(r, 0), str(num), bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl.cell(r, 0), row_bg)

    # Simon's comment: topic (bold) + point
    c1 = tbl.cell(r, 1); set_cell_bg(c1, row_bg); c1.text = ""
    p = c1.paragraphs[0]
    run = p.add_run(topic + " — "); run.bold = True; run.font.size = Pt(9)
    run.font.color.rgb = DARK
    run2 = p.add_run(point); run2.font.size = Pt(9); run2.font.color.rgb = GREY

    # Position badge cell
    pos_color = GREEN if pos == "ACCEPT" else (RED if pos == "COUNTER" else GOLDD)
    cell_text(tbl.cell(r, 2), pos, bold=True, size=9, color=pos_color, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(tbl.cell(r, 2), pos_bg)

    cell_text(tbl.cell(r, 3), note, size=9, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_bg(tbl.cell(r, 3), row_bg)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# PART B — REVISED TERM SHEET
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART B — Revised Term Sheet Clauses")
add_paragraph(doc,
    "The clauses below supersede the corresponding clauses in the v2 proposal and give effect "
    "to the positions in Part A.", italic=True, color=GREY, size=9, sa=6)

ts_section(doc, "1.  Parties & Company", [
    ("Founder",  "Davin Cloete, in his personal capacity and as Director of the Company."),
    ("Partner",  "Simon Hobday, in his personal capacity."),
    ("Company",  "BitWealth (Pty) Ltd, operating a Bitcoin DCA investment management platform."),
    ("FSP",      "Finova Capital (Pty) Ltd — BitWealth to be appointed a Juristic "
                 "Representative (JR) under Finova's FSP licence (20% performance-fee share + "
                 "50% platform-fee share; no fixed monthly hosting fee)."),
])

ts_section(doc, "2.  IP Foundation & Evidence  (addresses Comment 1)", [
    ("IP Asset", "The BitWealth LTH PVR platform, transferred to the Company under a signed IP "
                 "Transfer Agreement; recorded on the balance sheet at an agreed value of R2,500,000."),
    ("Evidence Pack",
     "Before signing, the Company will make available to Simon:\n"
     "(a) the executed IP Transfer Agreement;\n"
     "(b) developer/contractor IP-assignment confirmations;\n"
     "(c) a list of third-party and open-source dependencies and their licences;\n"
     "(d) a warranty of title and no known competing claims."),
    ("Valuation Basis",
     "R2.5m is the agreed commercial consideration, not an independent valuation. Davin "
     "warrants ownership and title, not market value. Simon may commission an independent "
     "valuation at his own cost."),
])

ts_section(doc, "3.  Simon's Role & FAIS  (addresses Comments 3 & 4)", [
    ("Role", "Non-executive Head of Business Development (consulting engagement), with defined "
             "quarterly client-acquisition deliverables agreed by the Board."),
    ("Support & Budget",
     "The Company will provide reasonable sales, marketing, and compliance/admin support, and "
     "an annual marketing/sales budget approved by the Board. Obligations on both sides are "
     "defined, not open-ended."),
    ("Time Commitment", "An agreed minimum time commitment per month, recorded in the engagement letter."),
    ("FAIS Accreditation",
     "Simon will use reasonable endeavours to obtain and maintain the FAIS accreditation "
     "required for his representative role. Delays caused by the FSCA, Finova, paperwork, or "
     "the Company are expressly excluded. Only a wilful or negligent failure by Simon "
     "constitutes a breach."),
])

ts_section(doc, "4.  Introduced Client Register  (addresses Comments 2 & 16)", [
    ("Register", "The Company maintains an 'Introduced Client Register'. Each client Simon "
                 "introduces is logged within 5 business days of first meaningful contact."),
    ("Definition",
     "An 'Introduced Client' is one where Simon makes the first meaningful contact and who is "
     "not already in the Company's active pipeline. Each client is credited once. Disputed "
     "entries are resolved by the Board acting reasonably, failing which by independent expert."),
    ("Conversion Tail",
     "An Introduced Client who converts within 18 months of introduction, or within 12 months "
     "of Simon's exit, counts toward milestones and any applicable revenue share (Part C)."),
])

ts_section(doc, "5.  Founding Equity Grant (10%)  (addresses Comment 17)", [
    ("Grant", "Davin transfers 100 ordinary shares (10%) to Simon for nil consideration on "
              "signing. No performance condition attaches to the founding 10%."),
    ("Tax Advice",
     "Both parties will obtain independent tax advice before any nil- or below-value transfer. "
     "The parties acknowledge potential income-tax, CGT, and donations-tax consequences of "
     "sweat equity, to be cleared before implementation."),
])

ts_section(doc, "6.  Earn-In Milestones (additional 30%)  (addresses Comments 5 & 6)", [
    ("Tranches", "Three equal 10% tranches (100 shares each), each a transfer of existing shares "
                 "from Davin. Targets: R50m / R100m / R200m AUM. Deadlines: 24 / 48 / 72 months "
                 "from the Commencement Date."),
    ("Objective Verification",
     "A milestone is achieved when total AUM meets the target at a month-end AND remains at or "
     "above it to the next month-end (two consecutive month-ends), verified from platform, "
     "custodian, exchange, and accounting records. This prevents a one-day spike from "
     "triggering a permanent equity transfer."),
    ("Administrative Confirmation",
     "Board confirmation is administrative only and may not be unreasonably withheld or delayed. "
     "If the evidence shows the threshold met and sustained and the Board has not confirmed "
     "within 10 business days, the milestone is deemed achieved."),
    ("Sequence", "Milestones are achieved in sequence; a later tranche cannot vest before the "
                 "earlier tranche has been earned (or its shortfall resolved under Clause 7)."),
])

ts_section(doc, "7.  Milestone Shortfall — Near-Miss Vesting  (addresses Comments 7 & 8)", [
    ("Replaces v2 options",
     "This unified mechanism replaces the four options in the v2 proposal and applies equally "
     "to all three tranches."),
    ("Step 1 — Extension",
     "If AUM reaches at least 80% of the target by the deadline, Simon automatically receives a "
     "single 6-month extension for that tranche."),
    ("Step 2 — Pro-rata vesting",
     "If, at the end of the extension, AUM is still below target but at least 80% of it, Simon "
     "vests a pro-rata portion of that tranche = (AUM achieved ÷ target) × 10%. Example: "
     "R45m of R50m (90%) → Simon vests 9% instead of 10%; the residual 1% remains with Davin."),
    ("Step 3 — Lapse",
     "If AUM is below 80% of target at the end of the extension, the tranche lapses in full and "
     "the shares remain with Davin. Simon retains all previously-earned tranches and his "
     "founding 10%."),
])

# Part C sits as its own major heading (New Ventures)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# PART C — NEW VENTURES & INNOVATION REVENUE
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART C — New Ventures & Innovation Revenue")
add_paragraph(doc,
    "This Part addresses Simon's request for a revenue share on new products or ideas that "
    "generate new revenue. It balances three things: (i) your strong preference that genuinely "
    "new products are built in a separate company on freshly-negotiated terms; (ii) fairness to "
    "Simon where he personally originates a new, non-AUM revenue stream; and (iii) protection "
    "against Simon receiving extra revenue share on top of his equity for revenue generated by "
    "clients that you or others introduce.", sa=6)

add_heading(doc, "The core problem, stated plainly", level=2, gold=False)
add_paragraph(doc,
    "Simon will already be a shareholder. As a shareholder he automatically shares — through "
    "dividends and equity value — in ALL of the Company's revenue, including AUM revenue from "
    "clients introduced by you or anyone else. If he ALSO received a revenue share on top of "
    "that, he would be paid twice for the same revenue, and would benefit from others' efforts. "
    "The solution is to separate 'ownership reward' (his equity) from 'origination reward' (a "
    "ring-fenced share of only the specific new stream he personally creates and runs).", sa=6)

ts_section(doc, "C1.  Default Rule — New Ventures go into a NewCo", [
    ("Default rule",
     "Any materially new product, business line, or revenue stream that is separable from the "
     "core BitWealth AUM business is developed in a NEW company ('NewCo'), with a fresh "
     "shareholding structure negotiated specifically for that venture."),
    ("No automatic entitlement",
     "Simon has NO automatic right to equity or revenue in any NewCo by virtue of his BitWealth "
     "shareholding. Each NewCo is negotiated on its own merits, based on who originates, funds, "
     "builds, and runs it."),
    ("Right to negotiate",
     "Where Simon materially originates a NewCo concept, he is given a good-faith right of first "
     "opportunity to negotiate participation in that NewCo — but on fresh terms, not on his "
     "BitWealth terms."),
    ("Who decides",
     "The Board (with Davin's control vote) decides whether a new idea is pursued inside "
     "BitWealth or via a NewCo. This keeps strategic control with you."),
])

ts_section(doc, "C2.  Exception — Simon-originated non-AUM revenue kept inside BitWealth", [
    ("When this applies",
     "Only where ALL of the following are true:\n"
     "(a) Simon personally originates a genuinely new, NON-AUM revenue stream (i.e., not "
     "performance fees, platform fees, or exchange-fee share on managed assets);\n"
     "(b) the Board decides to keep it inside BitWealth rather than a NewCo; and\n"
     "(c) Simon actively leads the delivery of that stream."),
    ("Innovation Revenue Share",
     "Simon receives a ring-fenced share of the NET revenue of THAT specific stream only:\n"
     "(a) suggested 15% of net revenue of the stream (negotiable per stream);\n"
     "(b) for a fixed period of 48 months from first revenue;\n"
     "(c) payable only while Simon actively leads the stream;\n"
     "(d) the stream is separately accounted so the ring-fence is auditable."),
    ("Ring-fence — the key protection",
     "The Innovation Revenue Share attaches ONLY to the new non-AUM stream Simon creates. It "
     "NEVER touches AUM-based revenue (performance, platform, or exchange fees) from any "
     "client — including clients you or others introduce. This directly answers your concern."),
    ("No double-count",
     "The Innovation Revenue Share is a defined cost of the stream, taken before profit. Simon "
     "then also shares, via his equity %, in whatever residual profit that stream contributes "
     "to the Company — which is normal for any shareholder-employee and is acceptable because "
     "the stream only exists due to his origination and active leadership."),
    ("IP ownership",
     "All intellectual property in the new stream vests in the Company (or NewCo), not in Simon "
     "personally, regardless of the revenue share."),
    ("Exclusions",
     "This does NOT apply to enhancements or extensions of the core AUM business (e.g., new "
     "client tiers, pricing changes, or new asset pairs) — those are core Company revenue in "
     "which Simon already shares as a shareholder."),
])

ts_section(doc, "C3.  Reciprocity & Anti-Circumvention", [
    ("Reciprocal",
     "The same framework applies to Davin- or third-party-originated ventures: Simon has no "
     "revenue-share claim over a non-AUM stream he did not originate and does not lead."),
    ("Anti-circumvention",
     "Neither party may route what is really core BitWealth AUM revenue through a 'new stream' "
     "label to avoid the other's legitimate share. Characterisation disputes go to independent "
     "expert determination."),
])

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
run = p.add_run("Recommended stance:  ")
run.bold = True; run.font.size = Pt(10); run.font.color.rgb = GREEN
run2 = p.add_run(
    "Offer Simon the NewCo default (C1) as the primary answer, and the ring-fenced Innovation "
    "Revenue Share (C2) as the narrow, fair exception for genuine non-AUM ideas he both "
    "originates and runs. Resist any general revenue share across the whole business — his "
    "equity already gives him that.")
run2.font.size = Pt(10); run2.font.color.rgb = DARK

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# PART B continued — Protective provisions (Clauses 8-12) & Cap table (Part D)
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART B (continued) — Protective Provisions")

ts_section(doc, "8.  FSP, Fee Transparency & Reserved Matters  (addresses Comment 9)", [
    ("Fee transparency",
     "ACCEPTED. The Company will share a full fee/profitability model showing net economics "
     "after the Finova revenue share (20% of performance fees + 50% of platform fees)."),
    ("Reserved matters — countered",
     "A 10% holder will NOT have a veto over operational/strategic decisions (FSP, exchange, "
     "custody, execution). Instead:\n"
     "(a) consultation rights on material changes;\n"
     "(b) an anti-prejudice covenant — no change may be made for the purpose of, or with the "
     "effect of, reducing Simon's economic entitlement below its then-current basis without "
     "his consent;\n"
     "(c) genuine reserved-matter consent rights that apply only once Simon holds 25% or more."),
])

ts_section(doc, "9.  Good Leaver & Bad Leaver  (addresses Comments 10, 11 & 14)", [
    ("Good Leaver — expanded",
     "Simon is a Good Leaver on: death/permanent disability; retirement; termination without "
     "cause; constructive dismissal; material reduction of his role; the Company/Davin failing "
     "to support onboarding; loss of FSP/JR status not caused by Simon; or material breach by "
     "Davin/the Company. Consequence: retains all earned shares; unearned tranches follow "
     "Clause 7."),
    ("Bad Leaver — narrowed",
     "Limited to proven fraud, theft, dishonesty, wilful misconduct, gross negligence, or a "
     "final (non-appealable) regulatory finding against Simon."),
    ("Bad Leaver consequence",
     "Unearned tranches lapse. EARNED shares are bought at fair market value, EXCEPT where the "
     "trigger is fraud/theft/dishonesty (then at par, R1/share). Any forced transfer requires "
     "prior independent determination by an auditor/expert or arbitrator."),
    ("Clawback — narrowed + reciprocal",
     "Clawback applies only to losses caused by Simon's wilful misconduct, fraud, gross "
     "negligence, or proven regulatory breach. Reciprocally, if Davin's/the Company's acts or "
     "omissions cause loss of JR/FSP status or damage client relationships, Simon is treated as "
     "a Good Leaver."),
])

ts_section(doc, "10.  Confidentiality, Non-Solicitation, Drag & Tag  (addresses Comments 12 & 13)", [
    ("Non-compete removed",
     "The broad non-compete is removed. Replaced with (i) confidentiality and (ii) "
     "non-solicitation of clients Simon actually dealt with and of staff, for 12 months "
     "post-exit, within South Africa."),
    ("Drag-along — with protections",
     "On a bona-fide sale of ≥ 75%, Simon must sell pro-rata on the SAME price and terms, with: "
     "cash-equivalent consideration; warranties limited to title and capacity; release from "
     "guarantees/liabilities; and no sale to a related party at an undervalue."),
    ("Tag-along",
     "If Davin sells more than 50% of his shares, Simon may include his shares on equivalent "
     "terms and the same price per share."),
])

ts_section(doc, "11.  Pre-emptive Rights & Anti-Circumvention  (addresses Comment 15)", [
    ("Nature clarified",
     "Simon has PRO-RATA pre-emptive rights (a right of first offer on new share issues at the "
     "same price as the incoming subscriber) — not a full ratchet or weighted-average "
     "anti-dilution."),
    ("Anti-circumvention",
     "No shares or options may be issued to Davin, associates, or staff below "
     "independently-determined fair value for the purpose of diluting Simon, without his "
     "consent."),
])

ts_section(doc, "12.  Conditions Precedent & General", [
    ("Conditions precedent",
     "(a) Shareholders' Agreement and Equity Participation Agreement in agreed form;\n"
     "(b) independent legal advice obtained by Simon;\n"
     "(c) independent tax advice obtained by both parties (Comment 17);\n"
     "(d) Finova Capital JR appointment or letter of intent;\n"
     "(e) confirmation Simon meets minimum FAIS fit-and-proper requirements."),
    ("Governing law", "Republic of South Africa; disputes to arbitration under AFSA rules."),
    ("Non-binding", "This document is non-binding and subject to contract; no obligations arise "
                    "until definitive agreements are signed."),
])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# PART D — REVISED CAP TABLE
# ══════════════════════════════════════════════════════════════════════════
add_heading(doc, "PART D — Revised Cap Table Model & Earn-In")

def cap_table(title, rows, note=None, header_bg=TEAL_HEX):
    add_heading(doc, title, level=2, gold=False)
    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [Cm(5.0), Cm(2.6), Cm(2.4), Cm(2.6), Cm(4.4)]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = widths[i]
    hdrs = ["Shareholder", "Shares", "% Holding", "Change", "Notes"]
    for i, hdr in enumerate(hdrs):
        cell_text(table.cell(0, i), hdr, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(table.cell(0, i), header_bg)
    for r_idx, row in enumerate(rows, 1):
        bg = L_GOLD_HEX if row[0].startswith("TOTAL") else (LGREY_HEX if r_idx % 2 == 0 else "FFFFFF")
        bold = row[0].startswith("TOTAL")
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT]
        for c_idx, (val, al) in enumerate(zip(row, aligns)):
            cell_text(table.cell(r_idx, c_idx), str(val), bold=bold, size=10, align=al)
            set_cell_bg(table.cell(r_idx, c_idx), bg)
    if note:
        p = doc.add_paragraph()
        run = p.add_run(f"★  {note}")
        run.italic = True; run.font.size = Pt(9); run.font.color.rgb = GREY
        p.paragraph_format.space_after = Pt(8)

cap_table("On Signing (Founding 10%)", [
    ["Davin Cloete", "900", "90%", "−100", "Technology, strategy, IP"],
    ["Simon Hobday", "100", "10%", "+100", "Founding grant (nil consideration)"],
    ["TOTAL", "1,000", "100%", "—", ""],
], note="Founding 10% carries no performance condition.", header_bg=TEAL_HEX)

cap_table("If All Milestones Achieved (Full Earn-In)", [
    ["Davin Cloete", "600", "60%", "−300", "After three 10% tranches"],
    ["Simon Hobday", "400", "40%", "+300", "R50m + R100m + R200m AUM all achieved"],
    ["TOTAL", "1,000", "100%", "—", ""],
], note="Maximum earn-in reached: Simon 40%, Davin 60%.", header_bg=GOLD_HEX)

add_heading(doc, "Near-Miss / Pro-Rata Vesting (new — Clause 7)", level=2, gold=False)
add_paragraph(doc,
    "Worked example showing how a narrowly-missed tranche now vests pro-rata instead of "
    "lapsing in full:", sa=4)

pr = doc.add_table(rows=5, cols=4)
pr.style = "Table Grid"; pr.alignment = WD_TABLE_ALIGNMENT.CENTER
pr_h = ["AUM achieved vs R50m target", "% of target", "Outcome", "Tranche vested"]
pr_w = [Cm(5.0), Cm(2.8), Cm(5.5), Cm(3.2)]
for i, col in enumerate(pr.columns):
    for cell in col.cells:
        cell.width = pr_w[i]
for i, h in enumerate(pr_h):
    cell_text(pr.cell(0, i), h, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(pr.cell(0, i), TEAL_HEX)
pr_rows = [
    ["R50m or more", "≥100%", "Full tranche vests", "10%"],
    ["R45m", "90%", "6-mo extension, then pro-rata", "9%"],
    ["R40m", "80%", "6-mo extension, then pro-rata", "8%"],
    ["R35m", "70%", "Below 80% — tranche lapses", "0%"],
]
for r, row in enumerate(pr_rows, 1):
    bg = LGREY_HEX if r % 2 == 0 else "FFFFFF"
    for c, val in enumerate(row):
        color = GREEN if ("vests" in val or "10%" == val) else (RED if "lapses" in val or val == "0%" else DARK)
        cell_text(pr.cell(r, c), val, size=9, color=color if c in (2, 3) else DARK,
                  bold=(c == 3), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(pr.cell(r, c), bg)

add_paragraph(doc,
    "Note: pro-rata vesting only applies at or above 80% of target, and only after the "
    "automatic 6-month extension has run. Below 80%, the tranche lapses and the shares remain "
    "with Davin.", italic=True, color=GREY, size=9, sb=4)

out_path = OUT_DIR / "BitWealth_Simon_Equity_Proposal_v3.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
