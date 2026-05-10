"""Populate the Response column of the CAEP Preliminary Licensing Questionnaire."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = Path(r"docs\FSCA Compliance\CAEP Licensing - Preliminary Licensing Questionnaire v4.0 202605.docx")
OUT = Path(r"docs\FSCA Compliance\CAEP Licensing - Preliminary Licensing Questionnaire v4.0 202605 - BitWealth Responses.docx")

NEED = "[NEED FROM YOU]"

# Map row index -> response text
# Row indices match the document table:
# R0=header, R1=Q1, R2=Q2, R3=Q3, R4=Q4, R5=Q5,
# R6=Q6 main, R7=Q6a entity, R8=Q6b bank, R9=Q6c mgmt accounts,
# R10=Q7 (count of Reps),
# R11=Q8 main (identified?), R12=Q8a names+prior, R13=Q8b qualifications, R14=Q8c employment elsewhere,
# R15=Q9 target date,
# R16=Q10 (advice/intermediary, Cat II),
# R17=Q11 product categories,
# R18=Q12 high-level strategy,
# R19=Q13 main (manage investments?), R20-R29=Q13a-j sub-questions,
# R30=Q14 target audience, R31=Q15 client numbers, R32=Q16 marketing,
# R33=Q17 placement agents, R34=Q18 COI, R35=Q19 regulatory,
# R36=Q23 header, R37=Q23 retail/prof/inst, R38=Q23 outside SA,
# R39=Q24 distribution, R40=Q25 product suppliers, R41=Q26 remuneration,
# R42=Q27 client funds, R43=Q28 outsourcing, R44=Q29 policies,
# R45=Q30 systems, R46=Q31 algo/digital,
# R47=Q35 support, R48=Q36 go-live
RESPONSES = {
    1: (
        "BitWealth (Pty) Ltd\n"
        "Registration number: 2026/090346/07"
    ),
    2: (
        "www.bitwealth.co.za — customer portal and marketing site."
    ),
    3: (
        "Davin Harald Gaier — Founder & CEO\n"
        "Email: davin.gaier@bitwealth.co.za\n"
        "Mobile: +27 82 407 9038"
    ),
    4: (
        "1. Davin Harald Gaier — Founder, sole shareholder and CEO. Responsible for overall "
        "strategy, technology architecture, on-chain research, ongoing software development, "
        "and product direction. 23 years of experience as a project manager in the financial "
        "services industry; PMP-certified through PMI.\n"
        "2. Simon Hobday (intended appointment as Distribution Director, subject to formal "
        f"agreement currently in progress) — responsible for client distribution, lead "
        f"generation and onboarding. Qualifications and prior FSP history: {NEED} — to be "
        "confirmed with Simon.\n\n"
        "There are no other shareholders, employees or contractors in BitWealth at present. "
        "The business is in pre-launch state. Under the CAEP appointment, BitWealth will "
        "operate as a Juristic Representative with two appointed Representatives (Davin and "
        "Simon), supported by CAEP's Key Individual and compliance infrastructure."
    ),
    5: (
        "BitWealth (Pty) Ltd is a South African fintech that has built a fully automated "
        "Bitcoin investment platform driven by a proprietary on-chain signal called the "
        "Long-Term Holder Price Variance Ratio (LTH PVR). The strategy applies dynamic "
        "Dollar-Cost Averaging — buying, holding or selling Bitcoin daily based on whether "
        "the price sits within statistically derived bands of long-term-holder cost basis. "
        "The platform is live on the VALR exchange via segregated client sub-accounts, "
        "fully back-tested against 10+ years of on-chain history, and benchmarked daily "
        "against a Standard DCA reference strategy.\n\n"
        "What it does, end-to-end:\n"
        "• Ingests on-chain LTH cost-basis data daily from CryptoQuant\n"
        "• Generates buy/sell/hold decisions per client, sized to available capital\n"
        "• Executes orders on VALR (LIMIT preferred, MARKET fallback) inside each "
        "client's dedicated VALR sub-account\n"
        "• Reconciles fills, posts a full ledger, and computes daily NAV per client\n"
        "• Sends client reporting and operational alerts; performance versus Standard DCA "
        "is published continuously\n\n"
        "Three-year objectives:\n"
        "• Year 1: Launch under CAEP as Juristic Representative; onboard initial 30–50 "
        "retail and HNW clients; establish operating track record and audited performance\n"
        "• Year 2: Grow AUM through partner-distribution channel and digital marketing; "
        "introduce additional digital-asset strategies (e.g. ETH variant, multi-asset)\n"
        "• Year 3: Either continue under CAEP or apply for own Cat II + CASP licence; "
        "expand to professional/institutional mandates and explore CIS structuring."
    ),
    6: (
        "Yes — BitWealth (Pty) Ltd will be appointed as a Juristic Representative under "
        "CAEP's FSP licence (Category II, and the pending CASP authorisation for the "
        "digital-asset component)."
    ),
    7: (
        "Already established. BitWealth (Pty) Ltd is a registered South African private "
        "company. No assistance required to establish a new entity."
    ),
    8: (
        "Yes — active business bank account held with First National Bank (FNB)."
    ),
    9: (
        "Not at present — the business is pre-operational and has not yet generated revenue. "
        "An external accountant has been engaged and the Founder will handle bookkeeping "
        "in-house until trading commences. Monthly management accounts will be prepared from "
        "go-live (target: 01 July 2026) onward."
    ),
    10: (
        "Two (2) authorised Representatives:\n"
        "1. Davin Harald Gaier — Founder & CEO\n"
        "2. Simon Hobday — intended Distribution Director (subject to confirmation of the "
        "partnership agreement, anticipated within the next month)"
    ),
    11: (
        "Yes — both candidates identified:\n\n"
        "1. Davin Harald Gaier — Founder. No prior FSP appointments.\n"
        f"2. Simon Hobday — intended Distribution Director. Prior FSP appointment history: "
        f"{NEED} — to be confirmed with Simon.\n"
    ),
    12: (
        "Davin Harald Gaier (Founder) — qualifications:\n"
        "• Highest formal qualification: Project Management Professional (PMP) certification "
        "through the Project Management Institute (PMI)\n"
        "• FAIS-recognised qualifications: none confirmed at this stage — would value CAEP's "
        "input on which existing qualifications/experience may be recognised under the FAIS "
        "fit-and-proper framework\n"
        "• RE5: not yet completed — willing to complete on CAEP's recommended timeline before "
        "or shortly after appointment\n"
        "• Industry experience: 23 years in the financial services industry (project management)\n"
        "• Professional designations: none\n\n"
        "Simon Hobday — qualifications:\n"
        f"• {NEED} — to be confirmed with Simon (highest qualification, RE5 status, FAIS "
        "qualifications, years of relevant experience, designations)\n\n"
        "Both appointees are willing to complete any RE5 / Class of Business / CPD "
        "requirements that CAEP identifies as part of the appointment process."
    ),
    13: (
        "Davin Harald Gaier:\n"
        "• Director and employee of J And D Equities (Pty) Ltd\n"
        "• Fixed-term contract with Absa Insurance Company\n"
        "• Not appointed as a Representative of any other FSP\n\n"
        f"Simon Hobday: {NEED} — to be confirmed with Simon."
    ),
    14: (
        "Target appointment date: 01 July 2026."
    ),
    15: (
        "Both — intermediary services and discretionary investment management.\n\n"
        "Category II (Discretionary FSP): YES — BitWealth will exercise discretion on a "
        "mandate basis. Each client signs a written discretionary investment-management "
        "mandate that authorises BitWealth to execute trades within agreed parameters in "
        "the client's segregated exchange sub-account.\n\n"
        "Mandate type: a single rules-based discretionary mandate ('LTH PVR Bitcoin DCA "
        "Mandate') with the following parameters disclosed to the client:\n"
        "• Universe: Bitcoin (BTC) only, traded against ZAR or USDT\n"
        "• Strategy: Daily LTH PVR signal-driven dollar-cost averaging (buy/hold/sell)\n"
        "• Discretion limits: trade sizing and timing within published rules; no leverage; "
        "no derivatives; no custody — funds remain in client's own VALR sub-account\n"
        "• Risk parameters: daily exposure caps, minimum cash buffer, kill-switch triggers\n"
        "• Reporting: daily NAV, monthly statements, real-time client portal access\n\n"
        "Category IIA (Hedge Fund) is not currently required — no pooled fund structure, "
        "no leverage, no shorting."
    ),
    16: (
        "Primary categories required:\n"
        "• Long-Term Deposits (where applicable for ZAR float — minor)\n"
        "• Securities and Instruments — Shares (potentially, for future strategies)\n"
        "• Crypto Assets — under the FSCA's Crypto Asset Service Provider (CASP) "
        "framework. This is the principal category required.\n\n"
        "We understand CAEP's CASP authorisation is pending application — confirmation of "
        "CASP timeline is critical to our go-live date and we would like to discuss this."
    ),
    17: (
        "High-level strategy:\n"
        "Provide South African retail and HNW investors with a regulated, transparent and "
        "fully automated way to gain Bitcoin exposure using a research-backed on-chain "
        "strategy that materially outperforms Standard DCA over 3+ year horizons (validated "
        "by 10+ years of back-testing).\n\n"
        "Target sector: Retail and HNW South African residents seeking long-term Bitcoin "
        "exposure with disciplined entry/exit signals rather than discretionary timing or "
        "buy-and-hold-only strategies.\n\n"
        "What we intend to do once appointed under CAEP's licence:\n"
        "1. Migrate the existing live operating platform to operate under CAEP's FSP/CASP "
        "umbrella (client mandates, disclosures, complaints process, COI register all "
        "aligned to CAEP standards under KI supervision).\n"
        "2. Activate distribution via the Founder's network and (subject to confirmation) "
        "a partner with a financial-services distribution book.\n"
        "3. Onboard initial 30–50 clients in the first 6–9 months, with full FICA/AML and "
        "FAIS suitability assessments per client.\n"
        "4. Operate, report, and build an audited performance track record under CAEP "
        "supervision."
    ),
    18: (
        "Yes — discretionary investment management, with a single rules-based mandate. "
        "Please see (a)–(j) below."
    ),
    19: (
        "Investments are NOT pooled. Each client's funds are held in their own dedicated "
        "VALR sub-account in the client's name (or in a structure beneficially owned by "
        "the client). BitWealth has trade authority over the sub-account by API key under "
        "the discretionary mandate but does not take custody of client assets.\n\n"
        "Structural summary:\n"
        "• Client opens an account with VALR (FSCA-registered exchange, separate FSP/CASP)\n"
        "• Client creates a dedicated sub-account for the BitWealth strategy\n"
        "• Client funds the sub-account in ZAR; BitWealth purchases/sells BTC on the "
        "client's behalf within that sub-account in line with the daily strategy signal\n"
        "• All cash and BTC remain in the client's name throughout — BitWealth never "
        "holds client funds or assets\n"
        "• Withdrawal authority remains exclusively with the client\n\n"
        "Future option: a CIS (unit trust / hedge fund) structure or single-strategy "
        "fund-of-one is on the roadmap for Year 2–3, but is not part of the initial "
        "appointment scope."
    ),
    20: (
        "Yes — introductions would be appreciated to the following:\n"
        "• Audit firm with crypto-asset experience for annual financial statements\n"
        "• Independent compliance officer / firm (if not provided by CAEP)\n"
        "• PI insurance broker familiar with discretionary FSPs and crypto-asset managers\n"
        "• (Future) CIS Manco that supports digital-asset strategies, when we move to a "
        "fund structure in Year 2/3\n\n"
        "Existing relationships already in place:\n"
        "• Exchange / execution: VALR (corporate account established; sub-account model "
        "for clients in use)\n"
        "• On-chain data provider: CryptoQuant (commercial subscription)\n"
        "• Cloud / database: Supabase\n"
        "• Email / notifications: Resend"
    ),
    21: (
        "Single instrument: Bitcoin (BTC) only.\n"
        "• Spot BTC, traded on VALR against ZAR (primary) and USDT (secondary)\n"
        "• Order types: LIMIT (post-only preferred) with a 5-minute MARKET fallback if "
        "unfilled\n"
        "• No leverage, no margin, no derivatives, no shorting, no staking, no DeFi "
        "exposure\n"
        "• Cash held in ZAR within the client's VALR sub-account between trades\n\n"
        "Future strategies on the roadmap (out of scope for initial appointment): "
        "ETH variant of LTH PVR, and a multi-asset crypto allocation."
    ),
    22: (
        "Long-term: 3 to 10+ years.\n"
        "The strategy is designed to outperform Standard DCA over multi-year holding "
        "periods (back-test horizon: 2014–present). Clients are screened on onboarding to "
        "confirm a minimum 3-year horizon and capacity to absorb significant interim "
        "drawdowns consistent with Bitcoin's historical volatility."
    ),
    23: (
        "Target return: outperform a Standard DCA Bitcoin benchmark over rolling 3-year "
        "and 5-year periods, after fees.\n\n"
        "Measurement framework (computed daily and reported to clients):\n"
        "• NAV per client (in ZAR)\n"
        "• Time-Weighted Return (TWR) and Money-Weighted Return (MWR)\n"
        "• Standard DCA benchmark (same cash-flow schedule, no signal) tracked alongside\n"
        "• Outperformance vs benchmark (TWR delta and ZAR delta)\n"
        "• Maximum drawdown\n\n"
        "Indicative back-tested outperformance (10-year, 2014–2024): meaningful TWR "
        "outperformance versus Standard DCA, with materially lower maximum drawdown. "
        "Detailed back-test report available on request — including 1, 3, 5, 7 and 10-year "
        "performance tables versus the Standard DCA benchmark.\n\n"
        "All marketing and client communication will use FAIS-compliant performance "
        "disclosure (past performance is not indicative; back-test caveats; net-of-fees; "
        "same-period benchmark)."
    ),
    24: (
        "Targeted total capital to be raised in Year 1: R 5m – R 15m AUM, raised across "
        "approximately 30–50 clients at an average ticket size of R 100k – R 300k."
    ),
    25: (
        "Launch AUM: R 1m – R 3m within the first 30 days of go-live (Founder pilots and "
        "early-adopter clients)."
    ),
    26: (
        "• 6 months:  R 3m – R 8m AUM\n"
        "• 12 months: R 8m – R 20m AUM\n"
        "• 24 months: R 25m – R 60m AUM"
    ),
    27: (
        "Target launch date: 01 July 2026, subject to CASP authorisation timing on CAEP's "
        "side, JR appointment processing by the FSCA, and completion of any outstanding "
        "fit-and-proper requirements (e.g. RE5)."
    ),
    28: (
        "Flow of funds (no client money handled by BitWealth at any point):\n\n"
        "1. Client signs the discretionary mandate, FAIS disclosure, and FICA pack.\n"
        "2. Client opens own VALR account, completes VALR's KYC, and creates a dedicated "
        "BitWealth sub-account.\n"
        "3. Client deposits ZAR directly into their own VALR sub-account from their own SA "
        "bank account (FICA-verified at VALR's level).\n"
        "4. BitWealth's daily automated pipeline issues BUY/SELL/HOLD signals; trades are "
        "placed on VALR via API in the client's sub-account, settling internally on VALR.\n"
        "5. ZAR and BTC remain in the client's sub-account at all times. BitWealth has "
        "trade-only API permissions — no withdrawal authority.\n"
        "6. Management and performance fees: BitWealth invoices the client (or, by client "
        "authority, deducts via VALR's sub-account fee mechanism) — to be confirmed against "
        "CAEP's preferred fee-collection model.\n"
        "7. Client withdrawals are initiated by the client directly from their own VALR "
        "account back to their own SA bank account.\n\n"
        "BitWealth at no point receives, holds, pools or has withdrawal authority over "
        "client cash or crypto."
    ),
    29: (
        "Primary target audience:\n"
        "• South African resident retail clients with investible surplus of R 100k+, a "
        "long-term horizon (3+ years), and risk capacity for material Bitcoin volatility\n"
        "• High-net-worth individuals (R 1m+ allocation) seeking a regulated, rules-based "
        "Bitcoin allocation as part of a diversified portfolio\n"
        "• Family offices and sophisticated investors comfortable with digital assets and "
        "looking for a non-discretionary, signal-driven exposure\n\n"
        "Excluded: clients with horizons under 3 years; clients unable to absorb 50%+ "
        "drawdowns; non-resident retail clients (initially)."
    ),
    30: (
        "• 12 months: 30 – 60 clients\n"
        "• 24 months: 80 – 150 clients\n"
        "• 36 months: 150 – 300 clients"
    ),
    31: (
        "Yes — assistance would be appreciated with:\n"
        "• Client-facing fact sheet / one-pager (we have a draft; would value CAEP review "
        "and FAIS-compliance sign-off)\n"
        "• FAIS-compliant marketing materials and disclosure templates (e.g. risk "
        "disclosures specific to crypto assets)\n"
        "• Mandate template review against CAEP's standard\n"
        "• Pitch deck for HNW / family-office introductions\n\n"
        "Existing collateral available for review:\n"
        "• Strategy one-pager (ABC One Pager)\n"
        "• 10-year back-test report with Standard DCA benchmark and performance tables\n"
        "• Customer portal walkthrough\n"
        "• Solution Design Document (technical architecture)"
    ),
    32: (
        "A distribution arrangement with Simon Hobday is in progress and is anticipated to "
        "be formalised within the next month. Once formalised, Simon will operate as an "
        "appointed Representative under CAEP/BitWealth (preferred structure) rather than as "
        "a separate introducer, so that all client-facing activity sits within the regulated "
        "perimeter.\n\n"
        "No other placement-agent, introducer or capital-raising arrangements are in place. "
        "Prior discussions with another FSP (RocketX) are paused indefinitely and there is "
        "no live arrangement.\n\n"
        "The Simon arrangement will be structured to comply with FAIS conflict-of-interest "
        "and remuneration rules, and disclosed to clients in line with CAEP's standard "
        "framework."
    ),
    33: (
        "None.\n\n"
        "For completeness:\n"
        "• No equity, fee-share or commercial arrangement is contemplated with any party "
        "that could constitute a material conflict of interest\n"
        "• Prior exploratory discussions with another FSP (RocketX) are paused; no agreement "
        "was concluded\n"
        "• BitWealth's relationship with VALR is a standard customer relationship\n"
        "• The Founder will adopt a personal-account-dealing policy under CAEP's framework "
        "to govern personal BTC holdings/trading and ensure separation from client activity\n"
        "• The single-product nature of the offering (Bitcoin only) and the AUM-based fee "
        "model will be disclosed transparently to clients from first contact"
    ),
    34: (
        "None known, for both Davin Harald Gaier and Simon Hobday:\n"
        "• No prior FSP appointment, debarment, or regulatory action\n"
        "• No insolvency, judgment, or sequestration history\n"
        "• No litigation pending against BitWealth (Pty) Ltd or either appointee\n"
        "• No criminal record related to dishonesty or financial-services offences\n\n"
        "Both appointees will complete CAEP's full fit-and-proper declaration when requested."
    ),
    36: (
        "Initially: retail clients (with FAIS suitability filtering for long horizons and "
        "risk capacity) and high-net-worth / sophisticated retail.\n\n"
        "In time: professional clients, family offices and sophisticated institutional "
        "investors as the AUM track record matures (Year 2–3).\n\n"
        "Not currently in scope: pension funds, retail collective-investment investors "
        "via CIS structures (CIS structuring is on the Year 2–3 roadmap if demand warrants)."
    ),
    37: (
        "Initial scope: South African residents only.\n\n"
        "Future (out of scope for initial appointment, subject to separate authorisations):\n"
        "• SADC residents, on a reverse-solicitation basis where permissible\n"
        "• Non-resident professional/HNW investors via appropriate cross-border structures\n\n"
        "All non-SA marketing or active solicitation would be raised with CAEP for "
        "regulatory review before any activity commences."
    ),
    38: (
        "Distribution model:\n\n"
        "1. Founder's direct network — initial pilot users and HNW introductions (Months 1–6)\n"
        "2. Distribution partner (subject to confirmation) — a financial-services "
        "professional with an existing book of suitable clients, operating as either an "
        "appointed Representative under CAEP/BitWealth or as a disclosed introducer with "
        "appropriate FAIS structuring\n"
        "3. Digital marketing — SEO, content marketing, performance reporting transparency "
        "(daily live performance vs Standard DCA published on the customer portal)\n"
        "4. Educational content — back-test data, on-chain research, and strategy "
        "explainers building organic credibility\n\n"
        "Sales journey (lead → onboarding):\n"
        "1. Lead enters via website, referral or partner introduction\n"
        "2. Discovery call with Founder/Representative — strategy walkthrough, suitability "
        "discussion, risk profiling\n"
        "3. Suitability assessment against documented FAIS criteria; unsuitable leads "
        "declined and recorded\n"
        "4. Issue of mandate, disclosure, fee schedule, and FICA pack\n"
        "5. Client signs digitally; FICA documents collected and verified\n"
        "6. Client opens/links VALR sub-account; BitWealth API permissions configured "
        "(trade-only, no withdrawal)\n"
        "7. Client funds the sub-account; BitWealth includes the client in the next daily "
        "execution cycle\n"
        "8. Welcome pack, portal credentials, and first-week monitoring touchpoint"
    ),
    39: (
        "• Execution venue: VALR (Pty) Ltd — FSCA-registered exchange and CASP applicant. "
        "Standard customer relationship; corporate account in place\n"
        "• On-chain data: CryptoQuant — paid commercial subscription for LTH cost-basis "
        "data\n"
        "• Cloud platform & database: Supabase\n"
        "• Edge runtime: Deno Deploy / Supabase Edge Functions\n"
        "• Email / notifications: Resend\n"
        "• Web hosting: Netlify (customer portal and marketing site)\n"
        "• Monitoring: in-house (Supabase logs + alert pipeline)\n\n"
        "Already in place: VALR (corporate + sub-account model), CryptoQuant, Supabase, "
        "Resend, Netlify. Introductions welcomed for (a) audit firm with crypto-asset "
        "experience, (b) PI insurance broker familiar with discretionary FSPs and crypto "
        "managers, (c) FAIS / RE5 training provider, and (d) any other specialist providers "
        "CAEP recommends for a regulated digital-asset manager."
    ),
    40: (
        "Remuneration model — two layers:\n\n"
        "1. Annual management fee: 1.50% – 2.00% per annum (charged monthly in arrears on "
        "average daily AUM) — to be finalised with CAEP based on CAEP fee-share and market "
        "comparables\n"
        "2. Performance fee: 15% – 20% of outperformance versus the Standard DCA benchmark, "
        "measured on a per-client high-water-mark basis and crystallised quarterly or "
        "annually (final structure to be agreed with CAEP)\n\n"
        "Fee split with CAEP: as per CAEP's standard JR fee-share — to be agreed.\n\n"
        "Indicative monthly/annual revenue (illustrative, at midpoint fee structure):\n"
        "• At R 10m AUM: ~R 17.5k/month management fee = R 210k/year, plus performance "
        "fees in positive years\n"
        "• At R 30m AUM: ~R 52.5k/month management fee = R 630k/year, plus performance fees\n"
        "• At R 60m AUM: ~R 105k/month management fee = R 1.26m/year, plus performance fees\n\n"
        "Main revenue drivers: (a) AUM growth via distribution partner, (b) client "
        "retention through transparent reporting and audited outperformance, (c) "
        "performance-fee crystallisation in BTC bull cycles."
    ),
    41: (
        "No — BitWealth does not collect, receive or hold client funds, premiums or "
        "investment monies at any point.\n\n"
        "All client cash (ZAR) is deposited by the client directly into the client's own "
        "VALR sub-account from the client's own FICA-verified SA bank account. All BTC "
        "purchased remains in that same client sub-account. Withdrawals are initiated by "
        "the client back to the client's own SA bank account.\n\n"
        "BitWealth's API access to each sub-account is restricted to TRADE permissions "
        "only — there is no WITHDRAWAL permission, technically or contractually. Each "
        "client's API key is created within their own VALR sub-account with explicit "
        "permission scopes, and is documented in the client mandate.\n\n"
        "Fee collection: management and performance fees will be collected via a method "
        "agreed with CAEP — either (a) direct invoicing of the client with payment to a "
        "CAEP-supervised business account, or (b) authorised in-sub-account fee debit "
        "subject to client mandate and clear disclosure."
    ),
    42: (
        "Functions intended to be outsourced under the CAEP appointment:\n"
        "• FAIS regulatory compliance administration — to CAEP\n"
        "• Key Individual oversight and supervision — to CAEP\n"
        "• Annual financial statements / accounting — external accountant (already engaged)\n"
        "• Annual statutory audit — external audit firm (introduction welcomed)\n"
        "• PI insurance — external broker (introduction welcomed)\n"
        "• Trade execution and exchange — VALR\n"
        "• On-chain data — CryptoQuant\n"
        "• Email/SMS comms — Resend\n\n"
        "Retained in-house: software development and operation, signal generation, daily "
        "trading-pipeline operation, client portal, client communication, monthly "
        "bookkeeping (handled by the Founder until volumes warrant outsourcing), and "
        "complaints intake (with escalation to CAEP per CAEP's complaints policy)."
    ),
    43: (
        "Yes — extensive operational documentation already exists (developed during the "
        "build phase). Available for CAEP review:\n\n"
        "• Solution Design Document (SDD) — comprehensive technical architecture, data "
        "flows, daily pipeline operation, schema design, deployment procedures (~1,700+ "
        "lines, current version 0.6)\n"
        "• Daily operating pipeline — 6-step automated pipeline with documented steps, "
        "timing, guards, and recovery procedures\n"
        "• Pipeline Resume runbook — procedures for failure recovery and pipeline restart\n"
        "• Alert system specification — alert categories, severity levels, escalation paths\n"
        "• Test-case libraries — Pipeline_Resume_Test_Cases, Alert_System_Test_Cases, "
        "WebSocket_Order_Monitoring_Test_Cases, Production Test Plan\n"
        "• Back-test methodology and results documentation (10+ years)\n"
        "• Standard DCA benchmark methodology\n"
        "• ABC One-Pager (client-facing strategy summary)\n\n"
        f"FAIS-specific policies still to be developed/aligned with CAEP standards "
        "(welcome CAEP guidance on templates):\n"
        "• Client onboarding & FICA procedures (under CAEP framework)\n"
        "• Suitability assessment templates\n"
        "• Conflicts-of-interest register and policy\n"
        "• Personal account dealing policy\n"
        "• Complaints policy\n"
        "• Record-keeping policy\n"
        "• Risk management framework\n"
        "• Business continuity / disaster recovery plan\n"
        "• Marketing and disclosure approval workflow"
    ),
    44: (
        "Production stack (already operational):\n\n"
        "Trading & strategy:\n"
        "• Daily automated pipeline on Supabase Edge Functions (Deno runtime)\n"
        "• PostgreSQL on Supabase (multi-schema: public / lth_pvr / lth_pvr_bt for "
        "back-testing isolation)\n"
        "• Pipeline orchestration via pg_cron (scheduled jobs at 03:00 UTC and 30-min "
        "guards)\n"
        "• VALR REST + WebSocket integration (HMAC-SHA512 signed requests; sub-account "
        "routing per client)\n"
        "• CryptoQuant API for on-chain LTH data ingestion\n\n"
        "Client onboarding & CRM:\n"
        "• Customer portal (web app) — KYC document upload, mandate signing, account "
        "linking, daily NAV reporting\n"
        "• In-house CRM via Supabase tables (customer_details, customer_portfolios, "
        "exchange_accounts) — would migrate to or integrate with CAEP's preferred CRM if "
        "required\n"
        "• Document storage: Supabase Storage (KYC documents, mandates, statements)\n\n"
        "Communications:\n"
        "• Resend for transactional and reporting email\n"
        "• Email logo and template suite professionally designed and standardised\n\n"
        "Back-office & reporting:\n"
        "• Daily ledger, fills reconciliation and NAV calculation per client\n"
        "• Standard DCA benchmark computed daily for transparent client comparison\n"
        "• Audit trail of every signal, decision, intent, order, fill and ledger line "
        "(immutable, time-stamped)\n\n"
        "Security:\n"
        "• Row-Level Security on all multi-tenant tables\n"
        "• Service-role keys held in environment variables; client API keys encrypted at "
        "rest\n"
        "• Trade-only API permissions on every client VALR sub-account (no withdrawal)\n"
        "• Centralised alert system with daily digest email to operations\n\n"
        "All of the above is in production today. We can demo the platform end-to-end at "
        "the next meeting."
    ),
    45: (
        "Yes — the entire BitWealth offering is an automated discretionary investment-"
        "management service. Material features:\n\n"
        "• Online onboarding: KYC document upload, FICA verification handoff, digital "
        "mandate signing\n"
        "• Algorithmic trading: 100% rules-based daily pipeline. No human discretion is "
        "applied to individual trade decisions; all decisions follow the documented LTH "
        "PVR strategy logic. Every signal, decision and trade is auditable end-to-end\n"
        "• No 'robo-advice' in the FAIS sense: BitWealth does not provide personalised "
        "financial advice. It provides a single discretionary mandate that the client "
        "elects to enter into after a Representative-led suitability discussion\n"
        "• Daily digital reporting: real-time portal access to NAV, holdings, performance "
        "vs Standard DCA benchmark, transaction history\n\n"
        "Algorithm governance:\n"
        "• Strategy logic is version-controlled and change-managed; any material change "
        "would be documented, back-tested, and (under the CAEP framework) submitted for "
        "KI/compliance review before deployment\n"
        "• Kill-switch and manual override available to suspend automated trading at any "
        "time\n"
        "• Daily reconciliation between expected (model) and actual (executed) trades\n\n"
        "We would welcome CAEP's input on the additional FAIS/COFI controls expected "
        "around algorithmic discretionary trading, and would look to CAEP's KI to have "
        "documented oversight over strategy changes."
    ),
    46: (
        "Primary support required from CAEP:\n\n"
        "1. Regulatory hosting — appointment of BitWealth (Pty) Ltd as Juristic "
        "Representative under CAEP's Cat II FSP licence and (when authorised) the CASP "
        "framework, including FSCA lodgement and ongoing supervision\n"
        "2. Key Individual oversight — assigned KI with documented capacity to supervise "
        "an algorithmic discretionary Bitcoin strategy\n"
        "3. Compliance infrastructure — FAIS-compliant policy templates, mandate template, "
        "FICA workflow, COI register, complaints and record-keeping framework, marketing-"
        "approval workflow\n"
        "4. Training & qualifications — guidance on RE5 / Class of Business / CPD "
        "requirements for the Founder and any second appointee, with referrals where "
        "needed\n"
        "5. PI insurance — guidance on minimum requirements and ideally inclusion under or "
        "alongside CAEP's existing arrangements\n"
        "6. Ongoing regulatory reporting — FSCA returns, financial-soundness monitoring, "
        "fit-and-proper recertification\n"
        "7. Strategic and commercial mentorship — particularly around CIS structuring and "
        "fund formation in Year 2/3\n"
        "8. Introductions — audit firm, accounting firm (if needed), and any specialist "
        "service providers relevant to a regulated digital-asset manager"
    ),
    47: (
        "Target go-live: 01 July 2026.\n\n"
        "Key milestones / dependencies:\n"
        "• Sign hosting / engagement agreement with CAEP\n"
        "• Confirm CAEP CASP authorisation status and indicative timeline (critical path)\n"
        "• Fit-and-proper documentation for both appointees and any qualification gap-fill "
        "(e.g. RE5)\n"
        "• Formalise the Simon Hobday appointment as Distribution Director (anticipated "
        "within the next month)\n"
        "• FSCA lodgement of JR appointment for BitWealth (Pty) Ltd\n"
        "• Operational alignment: client mandate template, FICA pack, disclosure pack, "
        "COI register, complaints policy — all mapped to CAEP standards under KI sign-off\n"
        "• PI insurance in place\n"
        "• Marketing material reviewed and approved by CAEP compliance\n"
        "• First-client onboarding dry run with KI present\n"
        "• Go-live: open to first cohort of clients\n\n"
        "Critical-path dependency: CASP authorisation timing on CAEP's side. We would "
        "appreciate confirmation of indicative timing at the next meeting so we can plan "
        "around it."
    ),
}


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def write_response(cell, text):
    # Clear existing content
    cell.text = ""
    p = cell.paragraphs[0]
    # Render with [NEED FROM YOU] in red bold
    parts = text.split(NEED)
    for i, part in enumerate(parts):
        if part:
            r = p.add_run(part)
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
        if i < len(parts) - 1:
            r = p.add_run("[NEED FROM YOU]")
            r.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
    # Light yellow background for cells that need user input
    if NEED in text:
        shade(cell, "FFF8E1")


def main():
    doc = Document(SRC)
    table = doc.tables[0]
    # Remap keys 11..47 to 12..48 to correct off-by-one (rows 11-47 in original
    # mapping referred to question groups; the actual table has 1-row offset
    # because Q8 occupies rows 11-14 not 11-13).
    remapped = {}
    for k, v in RESPONSES.items():
        if k <= 10:
            remapped[k] = v
        else:
            remapped[k + 1] = v
    # Add a short response to R11 (Q8 lead-in: "Have you identified the individuals?")
    remapped[11] = (
        "Yes — both appointees identified: Davin Harald Gaier (Founder) and Simon Hobday "
        "(intended Distribution Director, subject to formalising the partnership within the "
        "next month). Details for each provided in 8(a)–(c) below."
    )
    # Add a short response to R36 (Q23 header line "Client type and geography:")
    remapped[36] = "Detailed responses provided in the two sub-questions below."
    for row_idx, response in remapped.items():
        cell = table.rows[row_idx].cells[2]
        write_response(cell, response)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
