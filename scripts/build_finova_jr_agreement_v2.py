"""
BitWealth Asset Managers (Pty) Ltd  ×  Finova (Pty) Ltd
Juristic Representative (Mandate) Agreement — s13(1)(b) FAIS Act
Output: docs/FSCA Compliance/BitWealth_Finova_JR_Agreement_v2.docx

Built from Finova's template, updated and aligned with:
  · FAIS Act 37 of 2002 and the General Code of Conduct (BN 80 of 2003)
  · Determination of Fit & Proper Requirements (BN 194 of 2017)
  · Crypto Assets declared a financial product (GN 1350, 19 Oct 2022)
  · FICA 38 of 2001 (crypto asset service providers = accountable institutions)
  · POPIA 4 of 2013 (client confidentiality / data protection)

Key facts captured:
  · JR  = BitWealth Asset Managers (Pty) Ltd (Reg 2026/090346/07)
  · FSP = Finova (Pty) Ltd (Reg 2004/018433/07, FSP No. 21095)
  · Product = Crypto Assets — CATEGORY II (discretionary) confirmed by Finova
  · Custody = client assets in segregated VALR sub-accounts; JR never holds client funds
  · Fee split (FSP collects, pays JR its share):
        Performance fee  — JR 80% / FSP 20%
        Platform fee     — JR 50% / FSP 50%
        Management fee   — JR 50% / FSP 50%
        Exchange-fee share (VALR) — JR 100% (not shared)
  · Reverse referral — FSP pays JR 25% of revenue FSP receives on other
        products/services referred to it by the JR
  · Placeholders (highlighted) for items to confirm with Finova / compliance officer
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path("docs") / "FSCA Compliance"
OUT_DIR.mkdir(parents=True, exist_ok=True)
FINOVA_NAME = "Finova (Pty) Ltd"
FINOVA_REG  = "2004/018433/07"

DARK  = RGBColor(0x1A, 0x1A, 0x1A)
NAVY  = RGBColor(0x0A, 0x2A, 0x43)
GREY  = RGBColor(0x55, 0x55, 0x55)
RED   = RGBColor(0xB7, 0x1C, 0x1C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX  = "0A2A43"
LGREY_HEX = "F2F2F2"

BODY_FONT = "Calibri"
BODY_SIZE = 10.5


# ═══════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_runs(p, text):
    """Add text to paragraph, rendering [[...]] placeholders as yellow-highlighted runs."""
    import re
    parts = re.split(r"(\[\[.*?\]\])", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("[[") and part.endswith("]]"):
            r = p.add_run(part[2:-2])
            r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            r.font.color.rgb = RED
            r.bold = True
        else:
            r = p.add_run(part)
            r.font.color.rgb = DARK
        r.font.name = BODY_FONT
        r.font.size = Pt(BODY_SIZE)


def clause(doc, number, text, level=0, bold_lead=None, space_after=4):
    """Numbered clause paragraph. level controls indentation.
    bold_lead: optional leading text rendered bold (e.g. a defined-term label)."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8 + level * 1.1)
    pf.first_line_indent = Cm(-0.8) if number else Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if number:
        rn = p.add_run(f"{number}\t")
        rn.bold = True; rn.font.name = BODY_FONT; rn.font.size = Pt(BODY_SIZE)
        rn.font.color.rgb = DARK
    if bold_lead:
        rb = p.add_run(bold_lead)
        rb.bold = True; rb.font.name = BODY_FONT; rb.font.size = Pt(BODY_SIZE)
        rb.font.color.rgb = DARK
    _add_runs(p, text)
    return p


def section(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}.\t{title}" if number else title)
    r.bold = True; r.font.name = BODY_FONT; r.font.size = Pt(12)
    r.font.color.rgb = NAVY
    return p


def plain(doc, text, italic=False, size=BODY_SIZE, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_after=4, color=DARK, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.alignment = align
    _add_runs(p, text) if "[[" in text else None
    if "[[" not in text:
        r = p.add_run(text)
        r.italic = italic; r.bold = bold
        r.font.name = BODY_FONT; r.font.size = Pt(size); r.font.color.rgb = color
    return p


def hrule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "1"); b.set(qn("w:color"), NAVY_HEX)
    pBdr.append(b); pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════
# DOCUMENT
# ═══════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles["Normal"]
style.font.name = BODY_FONT
style.font.size = Pt(BODY_SIZE)
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.0)
    s.left_margin = s.right_margin = Cm(2.4)

# ── Title block ──────────────────────────────────────────────────────────────
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("JURISTIC REPRESENTATIVE AGREEMENT")
r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY; r.font.name = BODY_FONT

st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = st.add_run("(Mandate Agreement in terms of section 13(1)(b) of the "
               "Financial Advisory and Intermediary Services Act, 37 of 2002)")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY; r.font.name = BODY_FONT
st.paragraph_format.space_after = Pt(10)
hrule(doc)

# ── Parties ──────────────────────────────────────────────────────────────────
plain(doc, "Made and entered into by and between:", space_after=6, bold=True)

pa = doc.add_paragraph(); pa.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pa.add_run("FINOVA (PTY) LTD")
r.bold = True; r.font.size = Pt(12); r.font.name = BODY_FONT; r.font.color.rgb = DARK
plain(doc, "Registration number: 2004/018433/07",
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
plain(doc, "FSP Licence number: 21095", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
plain(doc, "(an authorised Financial Services Provider, hereinafter referred to as "
           "\u201cthe Appointing FSP\u201d)", align=WD_ALIGN_PARAGRAPH.CENTER,
      italic=True, space_after=6)

andp = doc.add_paragraph(); andp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = andp.add_run("and"); r.italic = True; r.font.name = BODY_FONT; r.font.size = Pt(11)

pb = doc.add_paragraph(); pb.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pb.add_run("BITWEALTH ASSET MANAGERS (PTY) LTD")
r.bold = True; r.font.size = Pt(12); r.font.name = BODY_FONT; r.font.color.rgb = DARK
plain(doc, "Registration number: 2026 / 090346 / 07",
      align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
plain(doc, "(hereinafter referred to as \u201cthe Juristic Representative\u201d)",
      align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=8)

plain(doc, "The Appointing FSP and the Juristic Representative are hereinafter "
           "collectively referred to as \u201cthe Parties\u201d and individually as a \u201cParty\u201d.",
      space_after=6)

# ── Recitals ─────────────────────────────────────────────────────────────────
section(doc, None, "PREAMBLE")
clause(doc, "A.", "The Appointing FSP is an authorised Financial Services Provider "
        "(FSP No. 21095) licensed under the FAIS Act to render financial services in respect "
        "of, among others, Crypto Assets;", space_after=4)
clause(doc, "B.", "Crypto Assets were declared a \u201cfinancial product\u201d under the FAIS Act "
        "by General Notice 1350 published on 19 October 2022, and the rendering of financial "
        "services in respect of Crypto Assets is accordingly regulated under the FAIS Act;",
        space_after=4)
clause(doc, "C.", "The Juristic Representative is a private company that operates a "
        "Bitcoin dollar-cost-averaging (DCA) investment execution platform and wishes to render "
        "intermediary services in respect of Crypto Assets to clients under the licence of the "
        "Appointing FSP;", space_after=4)
clause(doc, "D.", "The Juristic Representative will build and service its own client base "
        "and no sale of business or transfer of an existing client book is contemplated by this "
        "Agreement;", space_after=4)
clause(doc, "E.", "The Parties wish to record the terms upon which the Appointing FSP "
        "appoints the Juristic Representative as its juristic representative in terms of section "
        "13(1)(b) of the FAIS Act.", space_after=6)

# ── 1. Introduction / appointment ────────────────────────────────────────────
section(doc, "1", "INTRODUCTION AND APPOINTMENT")
clause(doc, "1.1", "The Appointing FSP hereby mandates and appoints the Juristic "
        "Representative as its juristic representative and authorises it to render the financial "
        "services indicated in Schedule A to clients, in respect of the financial product "
        "\u201cCrypto Assets\u201d (collectively, \u201cthe financial services\u201d).")
clause(doc, "1.2", "The Juristic Representative is authorised to render Category II "
        "DISCRETIONARY FSP SERVICES in respect of Crypto Assets, as indicated in Schedule A. "
        "The Juristic Representative is not authorised to furnish financial advice unless and "
        "until this Agreement and Schedule A are amended in writing to include advice and the "
        "applicable fit-and-proper competency requirements are met.")
clause(doc, "1.3", "The Juristic Representative accepts the appointment, agrees to be "
        "bound by the terms of this Agreement, and undertakes to abide by the FAIS Act, the "
        "General Code of Conduct for Authorised Financial Services Providers and Representatives "
        "(BN 80 of 2003), the Determination of Fit and Proper Requirements (BN 194 of 2017), "
        "the Financial Intelligence Centre Act, 38 of 2001 (\u201cFICA\u201d), the Protection of "
        "Personal Information Act, 4 of 2013 (\u201cPOPIA\u201d), and all other applicable legislation.")
clause(doc, "1.4", "The Appointing FSP accepts responsibility for those activities of the "
        "Juristic Representative that are performed within the scope of, or in the course of "
        "implementing, this Mandate Agreement, as contemplated in section 13 of the FAIS Act.")
clause(doc, "1.5", "This Mandate Agreement is confined to the rendering of financial "
        "services in respect of the financial product aforesaid, and nothing herein shall be "
        "construed as creating an agency for any matter not directly related to the provision of "
        "those financial services.")
clause(doc, "1.6", "The Juristic Representative is not authorised to bind the Appointing "
        "FSP in any manner whatsoever, save with the prior written authority of the Appointing FSP.")
clause(doc, "1.7", "The Juristic Representative undertakes to comply with all lawful "
        "instructions of the Appointing FSP, and the Appointing FSP shall not be liable for any "
        "act of the Juristic Representative in excess of its authority.")

# ── 2. Obligations of the JR ─────────────────────────────────────────────────
section(doc, "2", "OBLIGATIONS OF THE JURISTIC REPRESENTATIVE")
jr = [
    ("2.1", "The Juristic Representative agrees to give full effect to all reasonable and "
            "lawful directives, instructions and guidelines issued by the Appointing FSP from "
            "time to time in regard to the mandate."),
    ("2.2", "The Juristic Representative shall at all times comply with the lawful "
            "instructions of the appointed compliance officer of the Appointing FSP: "
            "Talita Olivier (tel: +27 21 879 3064). Any change to the designated compliance "
            "officer will be notified to the Juristic Representative in writing."),
    ("2.3", "The Juristic Representative undertakes to abide by all processes and procedures "
            "the Appointing FSP reasonably deems necessary to ensure compliance with the FAIS "
            "Act and any applicable subordinate legislation."),
    ("2.4", "The Juristic Representative must at all times maintain financial soundness and "
            "credit standing to the reasonable satisfaction of the Appointing FSP, as "
            "contemplated in the fit-and-proper requirements."),
    ("2.5", "The Juristic Representative undertakes to bring any complaint it receives "
            "relating to a financial service rendered by it to the attention of the Appointing "
            "FSP as soon as possible, and to co-operate in the resolution thereof in accordance "
            "with the Appointing FSP\u2019s complaints management framework."),
    ("2.6", "The Juristic Representative undertakes to maintain and keep in safe custody, for "
            "a period of at least 5 (five) years, all appropriate records pertaining to the "
            "financial services rendered, as required by the General Code of Conduct."),
    ("2.7", "The Juristic Representative undertakes to abide by the terms of any insurance "
            "contract effected by the Appointing FSP, including any professional indemnity cover, "
            "insofar as it relates to the rendering of financial services by the Juristic "
            "Representative."),
    ("2.8", "The Juristic Representative shall appoint at least one Key Individual who will "
            "manage and oversee the Juristic Representative and be responsible for the due "
            "performance of all functions relating to the financial services rendered. The "
            "current appointed Key Individual is Guy Algeo (ID 7305205050087). "
            "Davin Harald Gaier intends to be appointed as Key Individual once he has "
            "successfully completed the required regulatory examinations (RE) and satisfied all "
            "applicable fit-and-proper requirements. Until such time, Guy Algeo shall act as "
            "the Key Individual of the Juristic Representative, whereafter the Parties shall "
            "record the change in writing and effect the requisite FSCA notification."),
    ("2.9", "The Key Individual shall meet the experience, qualification and regulatory "
            "examination (RE) requirements applicable under the Determination of Fit and Proper "
            "Requirements (BN 194 of 2017), maintain the characteristics of honesty and "
            "integrity, complete the required class-of-business and product-specific training and "
            "continuous professional development (CPD), and sign the periodic fit-and-proper "
            "declaration issued by the Appointing FSP, at the Key Individual\u2019s own expense."),
    ("2.10", "The Key Individual must maintain, at all times, a working knowledge and "
             "understanding of all industry regulation and legislation applicable to a Juristic "
             "Representative rendering Category II discretionary financial services, including "
             "in respect of Crypto Assets, and comply therewith."),
]
for n, txt in jr:
    clause(doc, n, txt)

# 2.11 operational ability with sub-items
clause(doc, "2.11", "The Juristic Representative shall maintain Operational Ability at all "
        "times, and the Appointing FSP shall be entitled to perform periodic due-diligence "
        "checks to confirm such Operational Ability. Should it be compromised, the Appointing "
        "FSP may require the defect to be remedied within a reasonable time, failing which it "
        "may terminate this Agreement. Operational Ability includes:")
for n, txt in [
    ("2.11.1", "a fixed business address;"),
    ("2.11.2", "adequate access to communication facilities, including at least a full-time "
               "telephone or cell-phone service and document-duplication facilities;"),
    ("2.11.3", "adequate storage and filing systems for the safe-keeping of records, business "
               "communications and correspondence, including secure electronic records;"),
    ("2.11.4", "an account with a registered bank; and"),
    ("2.11.5", "all necessary policies, procedures and systems to ensure full compliance with "
               "the FAIS Act, its subordinate legislation, FICA, POPIA, and any other applicable "
               "anti-money-laundering or counter-terrorist-financing legislation."),
]:
    clause(doc, n, txt, level=1)

clause(doc, "2.12", "Where the Juristic Representative employs the services of a third-party "
        "administrator or technology provider to render administrative or system functions on "
        "its behalf (including the [[VALR exchange / other providers — confirm]]), it shall "
        "remain responsible for such functions and ensure that appropriate outsourcing controls "
        "and agreements are in place.")

clause(doc, "2.13", "The Juristic Representative shall ensure that internal control "
        "structures, procedures and controls are in place, including at least:")
for n, txt in [
    ("2.13.1", "segregation of duties, roles and responsibilities where appropriate from an "
               "operational-risk perspective;"),
    ("2.13.2", "application of logical access security;"),
    ("2.13.3", "access rights and data security on electronic data;"),
    ("2.13.4", "physical security of assets and records;"),
    ("2.13.5", "documentation of business processes, policies, controls and technical "
               "requirements;"),
    ("2.13.6", "system and application testing; and"),
    ("2.13.7", "disaster-recovery and back-up procedures on electronic data."),
]:
    clause(doc, n, txt, level=1)

for n, txt in [
    ("2.14", "The Juristic Representative shall ensure that the necessary system controls and "
             "compliance measures are in place to manage and monitor the relevant systems in use."),
    ("2.15", "The Juristic Representative shall record all financial and system procedures to "
             "ensure it is able to report in terms of applicable accounting requirements."),
    ("2.16", "The Juristic Representative shall maintain general administration, accounting and "
             "risk-control measures to ensure accurate, complete and timeous processing of data "
             "and reporting, and the Appointing FSP shall be entitled to access all records of "
             "the Juristic Representative relating to any financial service rendered."),
    ("2.17", "The Juristic Representative shall report to the Appointing FSP in such format and "
             "at such intervals as the Appointing FSP may reasonably determine in respect of the "
             "services rendered under this mandate."),
]:
    clause(doc, n, txt)

clause(doc, "2.18", "The Appointing FSP shall be entitled to monitor and review the Juristic "
        "Representative\u2019s performance and compliance with this mandate, the FAIS Act and all "
        "subordinate legislation, at such intervals as it may determine, in order to:")
for n, txt in [
    ("2.18.1", "identify regulatory risks;"),
    ("2.18.2", "assess the impact of such risks;"),
    ("2.18.3", "instruct the Juristic Representative to implement required procedures to manage "
               "the risk; and"),
    ("2.18.4", "assign responsibility for the on-going management of these risks to specific "
               "individuals within the Juristic Representative."),
]:
    clause(doc, n, txt, level=1)

# ── Custody / PI — adapted to VALR sub-account model ──────────────────────────
clause(doc, "2.19", "PROFESSIONAL INDEMNITY AND FIDELITY COVER. The Appointing FSP shall "
        "increase its professional indemnity cover and/or fidelity insurance cover "
        "appropriately to cover the assets under management (AUM) and clients of the Juristic "
        "Representative. The cost of such increase shall be borne by the Juristic Representative "
        "and shall be recovered by the Appointing FSP by way of deduction from the monthly fee "
        "amounts invoiced by, and payable to, the Juristic Representative in respect of the "
        "financial services rendered under this Agreement (as set out in Schedule C). The "
        "Appointing FSP shall provide the Juristic Representative with satisfactory evidence of "
        "the increased cover (including the insurer’s policy schedule and the cover amount) upon "
        "each such increase and on reasonable request. Such cover shall include the matters "
        "referred to below.", bold_lead=None)

clause(doc, "2.20", "CLIENT ASSETS \u2014 CUSTODY MODEL. The Parties record that client crypto "
        "assets and client funds are held in segregated client sub-accounts at VALR (a licensed "
        "crypto asset exchange / service provider), and that the Juristic Representative does "
        "NOT receive, hold, or take control of client funds or client financial products in the "
        "course of rendering the financial services. Accordingly, the trust-account obligations "
        "in clause 2.21 do not ordinarily apply. The Juristic Representative shall nevertheless:")
for n, txt in [
    ("2.20.1", "ensure that client assets remain segregated and readily discernible from the "
               "assets of the Juristic Representative at all times;"),
    ("2.20.2", "maintain full and proper accounting records, brought up to date monthly, and "
               "prepare annual financial statements in accordance with applicable financial "
               "reporting standards; and"),
    ("2.20.3", "cause such financial statements to be audited, where required, by an external "
               "auditor and submit them to the Appointing FSP not later than four months after "
               "its financial year-end, or such longer period as the Appointing FSP may allow."),
]:
    clause(doc, n, txt, level=1)

clause(doc, "2.21", "CONDITIONAL TRUST-ACCOUNT OBLIGATIONS. Should the Juristic "
        "Representative at any time receive or hold client funds or financial products, it shall "
        "immediately notify the Appointing FSP and comply with section 10 of the General Code of "
        "Conduct, including to: account properly and promptly for such funds or products; issue "
        "written confirmation of receipt; open and maintain a separate designated client bank "
        "account; pay client funds into that account within one business day of receipt; ensure "
        "the account holds only client funds; ensure interest accrues to the client; and ensure "
        "clients have ready access to their funds, less authorised deductions and lawful charges.")

for n, txt in [
    ("2.22", "The Juristic Representative shall have and employ resources, procedures and "
             "technological systems that eliminate, as far as reasonably possible, the risk that "
             "clients, product suppliers, other providers and representatives will suffer "
             "financial loss through theft, fraud, dishonest acts, poor administration, "
             "negligence, professional misconduct or culpable omissions."),
    ("2.23", "Without limiting clause 2.22, the Juristic Representative shall structure its "
             "internal control procedures to provide reasonable assurance that: (a) the business "
             "is carried on in an orderly and efficient manner; (b) financial and other "
             "information used or provided is reliable; and (c) all applicable laws are complied "
             "with."),
    ("2.24", "The Juristic Representative indemnifies and agrees, on demand, to make good to "
             "the Appointing FSP the amount of any claim paid by the Appointing FSP to any client "
             "as compensation for loss or financial prejudice suffered due to the Juristic "
             "Representative having contravened or failed to comply with any provision of the "
             "FAIS Act, or having wilfully or negligently rendered a financial service that "
             "caused, or is likely to cause, prejudice or damage to the client."),
    ("2.25", "If a client submits a complaint to the Office of the FAIS Ombud alleging that "
             "the Juristic Representative contravened the FAIS Act, or wilfully or negligently "
             "rendered a financial service causing prejudice, or treated the client unfairly, "
             "and a determination is made against the Appointing FSP, the Juristic Representative "
             "agrees that the Appointing FSP shall be entitled to recover the costs and the "
             "amount of any award from the Juristic Representative."),
    ("2.26", "The provisions of clauses 2.24 and 2.25 shall apply mutatis mutandis to any "
             "similar situation involving an employee of the Juristic Representative who is also "
             "mandated as a representative of the Appointing FSP."),
    ("2.27", "The Juristic Representative shall inform the Appointing FSP immediately, in "
             "writing, of any person employed by it who is also mandated as a representative of "
             "the Appointing FSP, or who leaves its employ, or when notice of termination of "
             "service is received \u2014 whichever occurs first."),
    ("2.28", "The Juristic Representative shall adopt and implement procedures approved by the "
             "Appointing FSP to maintain client confidentiality at all times, in accordance with "
             "POPIA. It shall ensure that all client agreements contain an acknowledgement that "
             "client information may be disclosed to the Appointing FSP or the relevant product "
             "provider without such disclosure being a breach of confidentiality."),
    ("2.29", "FICA COMPLIANCE. The Juristic Representative acknowledges that crypto asset "
             "service providers are accountable institutions under Schedule 1 to FICA, and "
             "undertakes to comply with all applicable customer due-diligence, record-keeping, "
             "risk-management, and reporting obligations under FICA and the Appointing FSP\u2019s "
             "Risk Management and Compliance Programme."),
]:
    clause(doc, n, txt)

# ── 3. Obligations of FSP ────────────────────────────────────────────────────
section(doc, "3", "OBLIGATIONS OF THE APPOINTING FSP")
clause(doc, "3.1", "The Appointing FSP shall provide the Juristic Representative with an "
        "appropriate policy and procedure manual, disclosure documents, and any other "
        "documentation required under the FAIS Act and the General Code of Conduct. The "
        "Appointing FSP may amend such policies and procedures from time to time, acting "
        "reasonably, and shall notify the Juristic Representative of material changes.")
clause(doc, "3.2", "The Appointing FSP shall provide the Juristic Representative with a "
        "certificate in terms of section 13(1)(b) of the FAIS Act and a certified copy of its "
        "licence as proof of authorisation, which certified copy must be displayed prominently "
        "at the premises of the Juristic Representative.")
clause(doc, "3.3", "The Appointing FSP shall collect all client fees generated by the "
        "financial services and shall pay the Juristic Representative its share of such fees in "
        "accordance with Schedule C (Fee Schedule).")
clause(doc, "3.4", "The Appointing FSP shall render reasonable compliance oversight, "
        "supervision and support to enable the Juristic Representative to render the financial "
        "services in a compliant manner.")

# ── 4. Fees and commission ───────────────────────────────────────────────────
section(doc, "4", "FEES, REVENUE SHARE AND PAYMENT")
clause(doc, "4.1", "The fees, revenue share and payment terms applicable to this Agreement are "
        "set out in Schedule C (Fee Schedule), which forms part of this Agreement.")
clause(doc, "4.2", "The Parties agree that Schedule C may be amended, replaced or updated from "
        "time to time by agreement in writing signed and dated by the authorised signatories of "
        "both Parties, WITHOUT any need to amend, re-execute or re-sign the balance of this "
        "Agreement. A duly signed replacement Schedule C shall supersede the previous Schedule C "
        "with effect from the date stated therein.")
clause(doc, "4.3", "All fees charged to clients shall be disclosed to clients in accordance "
        "with the General Code of Conduct, and neither Party shall charge any fee not permitted "
        "by law or not disclosed to the client.")

# ── 5. Disputes ──────────────────────────────────────────────────────────────
section(doc, "5", "DISPUTES")
clause(doc, "5.1", "If any dispute arises as to the validity, interpretation, effect, or the "
        "rights and obligations of either Party under this Agreement, either Party may require "
        "that the dispute be referred to arbitration before a single arbitrator.")
for n, txt in [
    ("5.1.1", "The arbitration shall be held informally at Johannesburg, and the arbitrator "
              "shall be agreed between the Parties within 5 (five) business days, failing which "
              "the arbitrator shall be nominated by the Chairperson of the relevant Bar Council "
              "or the President of the relevant Law Society (or its successor). The arbitrator "
              "shall be an attorney or advocate of at least 10 (ten) years\u2019 standing with "
              "experience and knowledge of financial services law and the FAIS Act, and with no "
              "interest in the proceedings."),
    ("5.1.2", "The Parties shall keep the arbitration, its subject matter and evidence "
              "confidential."),
    ("5.1.3", "The decision of the arbitrator shall be final and binding and not subject to "
              "appeal."),
    ("5.1.4", "The arbitrator shall include in the award an order as to the costs of the "
              "arbitration and who shall bear them."),
    ("5.1.5", "The arbitrator shall be guided at all times by the requirements of the FAIS Act "
              "and all applicable ancillary legislation."),
    ("5.1.6", "This clause shall not prevent a Party from applying to court for urgent relief "
              "in appropriate circumstances."),
]:
    clause(doc, n, txt, level=1)
clause(doc, "5.2", "Notwithstanding clause 5.1, the Parties agree that all terms of this "
        "Agreement are material, and that a material breach by the Juristic Representative shall "
        "entitle the Appointing FSP to cancel this Agreement in accordance with clause 6.")
clause(doc, "5.3", "Should the Juristic Representative cease to operate as a representative of "
        "the Appointing FSP, it shall notify the Appointing FSP timeously and do all things "
        "necessary to enable the Appointing FSP to take reasonable steps, in consultation with "
        "affected clients, to ensure that outstanding business is completed or transferred to "
        "the Appointing FSP or another representative.")
clause(doc, "5.4", "The Juristic Representative acknowledges having received a copy of the "
        "Appointing FSP\u2019s Debarment Procedures and understands that, should the Appointing FSP "
        "be required to debar the Juristic Representative, this Agreement shall be cancelled, "
        "save that the Juristic Representative shall remain bound by obligations already incurred "
        "(including under clauses 2.20, 2.21, 2.22 and 2.24) and in respect of outstanding fees "
        "under Schedule C (Fee Schedule).")
clause(doc, "5.5", "The Juristic Representative indemnifies the Appointing FSP against any "
        "claims arising out of any breach of this Agreement by the Juristic Representative.")

# ── 6. Termination ───────────────────────────────────────────────────────────
section(doc, "6", "TERMINATION")
clause(doc, "6.1", "This Agreement, or the status of the Juristic Representative as a "
        "mandatory of the Appointing FSP, may be terminated by the Appointing FSP on reasonable "
        "grounds, or in the event of failure by the Juristic Representative to comply with any "
        "provision of the FAIS Act or any lawful directive or guideline of the Appointing FSP.")
clause(doc, "6.2", "Without detracting from clause 6.1, the Appointing FSP may terminate this "
        "Agreement summarily, without compensation or payment in lieu of notice, if:")
for n, txt in [
    ("6.2.1", "the Juristic Representative commits a material breach of its obligations;"),
    ("6.2.2", "the Juristic Representative, a Key Individual, or any employee fails to comply, "
              "or no longer complies, with any relevant fit-and-proper requirement, including "
              "honesty and integrity;"),
    ("6.2.3", "the Juristic Representative makes any false declaration concerning the "
              "fit-and-proper status of any employee or Key Individual;"),
    ("6.2.4", "circumstances arise justifying termination on grounds of breach of any Code of "
              "Conduct or the common law; or"),
    ("6.2.5", "a Compliance Officer submits to the Appointing FSP a written report of material "
              "non-compliance with the FAIS Act by the Juristic Representative."),
]:
    clause(doc, n, txt, level=1)
clause(doc, "6.3", "Either Party may terminate this Agreement on 3 (three) calendar months\u2019 "
        "prior written notice to the other, provided that such notice may not be given during "
        "the first 12 (twelve) months of this Agreement, and subject to clause 6.2.")
clause(doc, "6.4", "Upon termination for any reason, the Juristic Representative shall "
        "immediately take reasonable steps, in consultation with the Appointing FSP and affected "
        "clients, to notify all affected clients and to ensure that outstanding business is "
        "completed or transferred to the Appointing FSP or another representative.")
clause(doc, "6.5", "CLIENT OWNERSHIP ON TERMINATION (LTH PVR PRODUCT). Notwithstanding any "
        "other provision of this Agreement, upon termination of this Agreement for any reason, "
        "all clients introduced by either the Juristic Representative or the Appointing FSP in "
        "respect of the proprietary BitWealth LTH PVR product shall transfer to, and vest in, "
        "the ownership of the Juristic Representative. The Appointing FSP shall, however, "
        "continue to be compensated with its agreed fee share (as set out in Schedule C) "
        "indefinitely in respect of any clients that the Appointing FSP itself introduced to the "
        "LTH PVR product, for so long as those clients remain invested in the product.")

# ── 7. Sequestration ─────────────────────────────────────────────────────────
section(doc, "7", "VOLUNTARY SEQUESTRATION, WINDING-UP AND CLOSURE")
clause(doc, "7.1", "In the event of the voluntary sequestration, winding-up or closure of "
        "either Party, that Party acknowledges that no application for voluntary surrender, and "
        "no special or written resolution relating to winding-up or voluntary closure, shall "
        "have legal force unless a copy or notice thereof has been lodged with the Registrar / "
        "Authority and the Authority has, by notice, declared that satisfactory arrangements "
        "have been made to meet all liabilities under transactions entered into with clients "
        "prior to such event, or unless the Authority declares that the application, resolution "
        "or closure is not contrary to the FAIS Act.")

# ── 8. General ───────────────────────────────────────────────────────────────
section(doc, "8", "GENERAL")
clause(doc, "8.1", "The Parties shall observe the utmost good faith in their dealings with "
        "one another and with all clients, and undertake at all times to render financial "
        "services honestly, fairly, with due skill, care and diligence, and in the interests of "
        "clients and the integrity of the financial services industry.")
clause(doc, "8.2", "This Agreement shall remain in force until lawfully terminated by either "
        "Party in accordance with clause 6.")
clause(doc, "8.3", "This Agreement is governed by the law of the Republic of South Africa.")
clause(doc, "8.4", "No alteration or variation of this Agreement shall be of any force or "
        "effect unless recorded in writing and signed by both Parties.")
clause(doc, "8.5", "This Agreement constitutes the entire agreement between the Parties in "
        "respect of its subject matter, and no representation or warranty not recorded herein "
        "shall be binding.")
clause(doc, "8.6", "INTELLECTUAL PROPERTY RIGHTS. All intellectual property rights in and to "
        "the BitWealth LTH PVR strategy, algorithm, software, platform, models, methodologies, "
        "data, trade marks, branding and all related materials (the \u201cBitWealth IP\u201d) are and "
        "shall remain the sole and exclusive property of the Juristic Representative. Nothing in "
        "this Agreement transfers, assigns, licenses or grants to the Appointing FSP any right, "
        "title or interest in or to the BitWealth IP, save for a limited, non-exclusive, "
        "non-transferable right to use the BitWealth IP solely to the extent strictly necessary "
        "to perform its oversight and supervisory functions under this Agreement and the FAIS "
        "Act. Any improvements, modifications or derivative works of the BitWealth IP, whether "
        "made by either Party, shall vest in and remain the property of the Juristic "
        "Representative. Upon termination of this Agreement, the Appointing FSP shall cease all "
        "use of the BitWealth IP and, on request, return or destroy all materials containing the "
        "BitWealth IP in its possession or control.")
clause(doc, "8.7", "The Parties choose the following as their respective domicilia citandi et "
        "executandi and addresses for the service of notices:")
plain(doc, "The Appointing FSP: Finova (Pty) Ltd, Office 4, Needham House, Broadacres Shopping "
           "Centre, Cnr Valley and Cedar Road, Bryanston, Sandton. Email: info@finova.co.za",
      space_after=2)
plain(doc, "The Juristic Representative: BitWealth Asset Managers (Pty) Ltd, "
           "89 Bloekom Curve, Terenure ext. 16, Kempton Park, Gauteng, South Africa, 1619. "
           "Email: info@bitwealth.co.za", space_after=6)

clause(doc, "8.8", "COMMENCEMENT DATE. This Agreement comes into force and effect on the date "
        "of signature by the last Party to sign (“the Effective Date”), which date shall be "
        "inserted by each Party next to their respective signatures below. Each Party warrants "
        "authority to sign and is bound from the Effective Date.")

# ── Schedule A — product & services ──────────────────────────────────────────
doc.add_page_break()
section(doc, None, "SCHEDULE A — FINANCIAL PRODUCT AND SERVICES")
plain(doc, "The Juristic Representative is authorised to render the financial services marked "
           "below, in respect of the following financial product, under Category I of the FAIS "
           "Act licensing framework:", space_after=6)

sa = doc.add_table(rows=2, cols=4)
sa.style = "Table Grid"; sa.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in sa.rows:
    row.cells[0].width = Cm(1.6); row.cells[1].width = Cm(7.5)
    row.cells[2].width = Cm(3.3); row.cells[3].width = Cm(3.3)
sa_hdr = ["Nr", "Category I financial product", "Advice", "Intermediary service"]
for i, h in enumerate(sa_hdr):
    c = sa.cell(0, i); c.text = ""
    run = c.paragraphs[0].add_run(h)
    run.bold = True; run.font.size = Pt(9.5); run.font.name = BODY_FONT; run.font.color.rgb = WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, NAVY_HEX)
sa_row = ["1.1", "Crypto Assets", "\u2014 (not authorised)", "\u2713 Authorised"]
for i, val in enumerate(sa_row):
    c = sa.cell(1, i); c.text = ""
    run = c.paragraphs[0].add_run(val)
    run.font.size = Pt(10.5); run.font.name = BODY_FONT
    run.font.color.rgb = (GREEN := RGBColor(0x1B, 0x5E, 0x20)) if val.startswith("\u2713") else DARK
    run.bold = val.startswith("\u2713")
    c.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER)
plain(doc, "Note: Advice in respect of Crypto Assets is expressly excluded under this "
           "Agreement (clause 1.2). Should advice be required in future, this Schedule and the "
           "Agreement must be amended in writing and the applicable competency requirements met.",
      italic=True, size=9, color=GREY, space_after=6)
plain(doc, "Category note: This Schedule reflects Category I (intermediary services). If the "
           "BitWealth LTH PVR product exercises discretion on behalf of clients (i.e. automated "
           "buy/sell decisions taken without a prior client instruction for each transaction), "
           "the product may instead require Category II (discretionary FSP) licensing and a "
           "discretionary mandate compliant with the Code of Conduct for Administrative and "
           "Discretionary FSPs (BN 79 of 2003). [[Finova to confirm the correct FAIS category "
           "for the LTH PVR product — Category I (execution only) vs Category II (discretionary)]]",
      italic=True, size=9, color=GREY, space_after=6)

# ── Schedule B — items to confirm ────────────────────────────────────────────
section(doc, None, "SCHEDULE B — OUTSTANDING ITEMS TO CONFIRM")
plain(doc, "The following items (highlighted throughout the Agreement) must be confirmed with "
           "Finova and its compliance officer before signature:", space_after=6)
for item in [
    "Finova Capital (Pty) Ltd company registration number.",
    "Finova\u2019s appointed compliance officer \u2014 name and contact details (clause 2.2).",
    "Guy Algeo’s ID number (interim Key Individual), and confirmation of the transition to "
    "Davin Harald Gaier as Key Individual upon his completion of the RE examinations "
    "(clauses 2.8–2.9).",
    "Finova to confirm the increased professional indemnity / fidelity cover amount, provide "
    "evidence of cover, and confirm the monthly deduction mechanism (clause 2.19).",
    "Confirmation of third-party providers to be named in clause 2.12 (e.g. VALR).",
    "Arbitration city and the relevant Bar Council / Law Society region (clause 5.1.1).",
    "Finova registered address and email address for notices (clause 8.7).",
    "Effective / commencement date of the Agreement.",
    "Finova to confirm the correct FAIS category for the LTH PVR product — Category I "
    "(execution only) versus Category II (discretionary) — see Schedule A.",
    "Finova signatory name and capacity (Signatures and Schedule C).",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item); r.font.size = Pt(10); r.font.name = BODY_FONT; r.font.color.rgb = DARK

# ── Signatures ───────────────────────────────────────────────────────────────
doc.add_page_break()
section(doc, None, "SCHEDULE C \u2014 FEE SCHEDULE")
plain(doc, "This Schedule C forms part of the Juristic Representative Agreement between the "
           "Parties. In terms of clause 4.2, it may be amended, replaced or updated by agreement "
           "in writing signed and dated by both Parties, without amending or re-executing the "
           "balance of the Agreement.", italic=True, size=9, color=GREY, space_after=6)

clause(doc, "C.1", "NO MONTHLY FEE. The Juristic Representative shall not be required to pay "
        "the Appointing FSP any fixed monthly fee for the use of the Appointing FSP\u2019s "
        "operational oversight, licence and processes.")
clause(doc, "C.2", "REVENUE SHARE. The Juristic Representative renders intermediary services "
        "that generate client fees, which are collected by the Appointing FSP and shared with "
        "the Juristic Representative as set out below. The percentages reflect the share "
        "RETAINED by each Party of the relevant net fee actually received by the Appointing FSP:")

tblc = doc.add_table(rows=5, cols=3)
tblc.style = "Table Grid"; tblc.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in tblc.rows:
    row.cells[0].width = Cm(7.5); row.cells[1].width = Cm(4.2); row.cells[2].width = Cm(4.2)
hdrc = ["Fee type", "Juristic Representative (BitWealth)", "Appointing FSP (Finova)"]
for i, h in enumerate(hdrc):
    c = tblc.cell(0, i); c.text = ""
    run = c.paragraphs[0].add_run(h)
    run.bold = True; run.font.size = Pt(9.5); run.font.name = BODY_FONT; run.font.color.rgb = WHITE
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_bg(c, NAVY_HEX)
fee_rows_c = [
    ["Performance fee", "80%", "20%"],
    ["Platform fee", "50%", "50%"],
    ["Management fee", "50%", "50%"],
    ["Exchange-fee share (VALR)", "75%", "25%"],
]
for r, row in enumerate(fee_rows_c, 1):
    bg = LGREY_HEX if r % 2 == 0 else "FFFFFF"
    for cidx, val in enumerate(row):
        c = tblc.cell(r, cidx); c.text = ""
        run = c.paragraphs[0].add_run(val)
        run.font.size = Pt(10); run.font.name = BODY_FONT; run.font.color.rgb = DARK
        run.bold = (cidx > 0)
        c.paragraphs[0].alignment = (WD_ALIGN_PARAGRAPH.LEFT if cidx == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(c, bg)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

clause(doc, "C.3", "REVERSE REFERRAL. In addition, the Appointing FSP shall pay the Juristic "
        "Representative 25% (twenty-five percent) of the net revenue actually received by the "
        "Appointing FSP in respect of any other products or services referred by the Juristic "
        "Representative to the Appointing FSP.")
clause(doc, "C.4", "DEFINITIONS. \u201cPerformance fee\u201d, \u201cplatform fee\u201d and "
        "\u201cmanagement fee\u201d bear the meanings given to them in the client mandate/agreement "
        "concluded with each client. \u201cNet revenue\u201d and \u201cnet fee\u201d mean the amount actually "
        "received by the Appointing FSP after deduction of any product-provider or third-party "
        "charges and any statutory levies, but before income tax.")
clause(doc, "C.5", "PAYMENT. The Juristic Representative shall invoice the Appointing FSP "
        "monthly for its share of fees and revenue. The Appointing FSP shall pay such amount, "
        "less any professional indemnity / fidelity cover increase recoverable under clause 2.19, "
        "on or before the last business day of the month following the month in which the "
        "relevant fee or revenue is actually received by the Appointing FSP.")
clause(doc, "C.6", "PAYMENT DETAILS. Payments shall be made by electronic funds transfer to "
        "the Juristic Representative\u2019s nominated bank account:")
plain(doc, "Account holder: BitWealth Asset Managers (Pty) Ltd", space_after=1)
plain(doc, "Bank: First National Bank (FNB)", space_after=1)
plain(doc, "Account number: 63196928502", space_after=1)
plain(doc, "Account type: Cheque", space_after=1)
plain(doc, "Branch name: MY BRANCH", space_after=1)
plain(doc, "Branch / universal code: 255355", space_after=6)

plain(doc, "Agreed and accepted \u2014 Schedule C (Fee Schedule):", bold=True, space_after=10)
for line in ["For the Appointing FSP:  Signed __________________   Name: Guy Algeo   Capacity: [[Guy Algeo to complete]]   Date __________",
             "For the Juristic Representative:  Signed __________________   Name Davin Harald Gaier   Date __________"]:
    plain(doc, line, space_after=8)

doc.add_page_break()
section(doc, None, "SIGNATURES")
plain(doc, "Signed for and on behalf of the Appointing FSP, the signatory warranting authority:",
      space_after=10, bold=True)
for line in ["Signed: __________________________     Name: Guy Algeo",
             "Capacity: [[Guy Algeo to complete their own capacity/title]]",
             "At: Johannesburg            Date: (date of last signature = Effective Date)",
             "Capacity: ________________________     Date: __________________",
             "At: ______________________________",
             "Witness 1: _______________________     Witness 2: _______________________"]:
    plain(doc, line, space_after=6)

plain(doc, "Signed for and on behalf of the Juristic Representative, the signatory warranting "
           "authority:", space_after=10, bold=True)
for line in ["Signed: __________________________     Name: Davin Harald Gaier",
             "Capacity: Director                     Date: __________________",
             "At: ______________________________",
             "Witness 1: _______________________     Witness 2: _______________________"]:
    plain(doc, line, space_after=6)

# ── Footer disclaimer ────────────────────────────────────────────────────────
hrule(doc)
plain(doc, "PREPARED FOR REVIEW \u2014 This draft has been prepared to update Finova\u2019s template "
           "for BitWealth Asset Managers (Pty) Ltd and to align its clauses with the FAIS Act, "
           "the General Code of Conduct, the Fit & Proper Determination, FICA and POPIA. It is "
           "not legal advice and must be reviewed and approved by Finova\u2019s compliance officer "
           "and a qualified attorney before signature. Highlighted items require confirmation.",
      italic=True, size=8.5, color=GREY, space_after=2)

out = OUT_DIR / "BitWealth_Finova_JR_Agreement_v2.docx"
doc.save(out)
print(f"Saved: {out}")
