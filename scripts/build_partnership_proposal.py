"""Generate BitWealth_Partnership_Proposal.docx for RocketX/Simon discussion."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colours
GOLD = RGBColor(0xC9, 0xA2, 0x27)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
LIGHT_GOLD_HEX = "F5EBC8"

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_heading(text, level=1, color=GOLD):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h


def add_para(text, bold=False, italic=False, size=11, color=DARK, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_bullet(text, indent=0):
    p = doc.add_paragraph(text, style='List Bullet')
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table(headers, rows, header_fill=LIGHT_GOLD_HEX, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        shade_cell(hdr[i], header_fill)
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl


# ============ COVER ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("BitWealth (Pty) Ltd")
tr.bold = True
tr.font.size = Pt(28)
tr.font.color.rgb = GOLD

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Partnership & Equity Structure — Proposal Framework")
sr.font.size = Pt(16)
sr.font.color.rgb = DARK

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Prepared for discussion with the CEO of RocketX\nApril 2026  |  Discussion Document — Subject to Negotiation")
mr.font.size = Pt(10)
mr.italic = True
mr.font.color.rgb = GREY

doc.add_paragraph()

# ============ 1. PURPOSE ============
add_heading("1. Purpose of This Document", level=1)
add_para(
    "This document sets out a proposed two-phase structure for a partnership between BitWealth "
    "(Pty) Ltd, RocketX, and Simon [Surname]. Phase 1 is a formal Juristic Representative "
    "appointment under RocketX's FSP licence, running for an initial 12 months and renewable "
    "annually. Phase 2 — a permanent equity structure — is optional and only commences if and "
    "when all parties mutually agree to it."
)
add_para(
    "All numbers, percentages, and terms in this document are starting positions for discussion. "
    "Final terms will be negotiated and documented in formal agreements drafted by attorneys.",
    italic=True, color=GREY
)

# ============ 2. PARTIES & CONTRIBUTIONS ============
add_heading("2. Parties and Contributions", level=1)
add_para(
    "Each party brings a distinct and complementary contribution. Recognising the relative "
    "weight of each contribution is the foundation for the equity discussion."
)

add_table(
    headers=["Party", "Primary Contribution", "Nature"],
    rows=[
        ["Founder ([Your Name])", "IP, technology, product, founding effort, ongoing CEO role", "Sweat equity + existing IP"],
        ["RocketX (or CEO personally)", "FSP & CASP licence, regulatory & compliance infrastructure, industry credibility", "Regulatory access + governance"],
        ["Simon [Surname]", "Distribution network, client acquisition, capital injection (TBC), full-time operating role", "Capital + sweat equity + sales"],
    ],
    col_widths=[4.5, 7.5, 4.5]
)

# ============ 3. PHASE 1 ============
add_heading("3. Phase 1 — Juristic Representative (JR) Appointment", level=1)
add_para(
    "BitWealth (Pty) Ltd is appointed as a Juristic Representative of RocketX in terms of "
    "section 13 of the Financial Advisory and Intermediary Services Act (FAIS). This is a "
    "formal, FSCA-registered appointment — not an informal commercial arrangement — that "
    "allows BitWealth to render financial services to clients under RocketX's FSP licence "
    "from day one. The appointment runs for an initial 12-month term and renews automatically "
    "on each anniversary unless either party gives 90 days' written notice of non-renewal."
)
add_para(
    "Phase 1 stands on its own. It is not a trial period that must convert into equity — it "
    "is a complete, sustainable operating model in its own right. Phase 2 (equity) is a "
    "separate, optional decision that requires the mutual agreement of all parties.",
    italic=True, color=GREY
)

add_heading("3.1 Proposed JR Terms", level=2)
add_table(
    headers=["Term", "Proposed Position"],
    rows=[
        ["Legal basis", "Juristic Representative appointment under section 13 of FAIS, registered with the FSCA"],
        ["Effective date", "Upon (a) execution of the JR Agreement and (b) FSCA approval of BitWealth as a JR under RocketX's FSP licence"],
        ["Initial term", "12 months from the Effective Date"],
        ["Renewal", "Automatic renewal for further 12-month periods unless either party gives 90 days' written notice of non-renewal before the renewal date"],
        ["Fee-sharing (FAIS s.13)", "40% BitWealth / 30% RocketX / 30% Simon of net management and performance fees, after platform/exchange costs"],
        ["Compliance services included", "Full FAIS compliance administration, KI supervision, regulatory reporting, FICA/AML programme, complaints handling, record-keeping under the FSP"],
        ["Key Individual", "Assigned KI nominated by RocketX (or mutually agreed) with documented capacity and supervision plan for the LTH PVR strategy"],
        ["PI insurance", "Either extended under RocketX's existing professional indemnity cover, or BitWealth carries its own policy meeting the FSP's minimum requirements — to be agreed"],
        ["Client relationship", "Client mandates are with BitWealth (as JR), supervised by RocketX (as FSP). Client relationship and data ownership remain with BitWealth."],
        ["Client portability", "On termination of the JR appointment, BitWealth retains the right to transfer client mandates to an alternative FSP (or its own FSP) subject to client consent and FAIS notification rules"],
        ["Right to seek alternative FSP", "BitWealth may, at any time, apply for its own FSP licence or negotiate appointment under another FSP. RocketX shall not unreasonably impede such efforts."],
        ["Conflicts of interest", "Documented conflicts-of-interest management plan, including disclosure to clients of the JR/FSP fee-share and any RocketX equity interests"],
        ["Early termination triggers", "Material breach (with cure period), regulatory action against either party, insolvency, loss of FSP licence, mutual written consent"],
        ["Phase 2 trigger", "Phase 2 (equity) commences only on the mutual written agreement of all parties. Discussions to be initiated 90 days before any renewal date."],
        ["Effect of Phase 2", "On the SHA Effective Date, the JR Agreement is superseded by intra-group operating arrangements. The JR appointment with the FSCA continues uninterrupted."],
    ],
    col_widths=[5, 11.5]
)

add_heading("3.2 FSCA Appointment Process", level=2)
add_para(
    "Appointment of BitWealth (Pty) Ltd as a JR requires RocketX to lodge the appointment "
    "with the FSCA. BitWealth must satisfy the fit-and-proper requirements applicable to a "
    "juristic representative:"
)
add_bullet("Honesty, integrity and good standing declarations from all directors")
add_bullet("Operational ability — documented systems, processes and controls (largely already in place)")
add_bullet("Financial soundness — solvency and liquidity demonstrated")
add_bullet("Confirmation that RocketX's FSP categories (and, if applicable, CASP authorisation) cover the LTH PVR strategy")
add_bullet("Indicative timeline: 30–60 days from lodgement to FSCA approval")

add_heading("3.3 Why 30% / 30% / 40%?", level=2)
add_para(
    "Industry-typical JR/FSP fee-sharing arrangements in South Africa range from 20% to 40% "
    "to the FSP. 30% sits in the middle and reflects the value of a fully outsourced "
    "compliance function plus the regulatory licence itself. The matching 30% to Simon "
    "compensates the distribution role on equal footing during Phase 1, when no party yet "
    "holds equity. BitWealth retains 40% to fund ongoing technology development, infrastructure "
    "and operations."
)

# ============ 4. PHASE 2 ============
add_heading("4. Phase 2 — Permanent Equity Structure (Optional)", level=1)
add_para(
    "Phase 2 commences only if and when all parties mutually agree to convert the JR "
    "relationship into a permanent shareholding. It is not automatic and is not a precondition "
    "of Phase 1. If Phase 2 is agreed, the JR fee-share falls away and is replaced by "
    "intra-group operating arrangements; BitWealth (Pty) Ltd becomes a jointly-owned operating "
    "company. Three illustrative options are set out below."
)

add_heading("4.1 Founder Non-Negotiables", level=2)
add_bullet("Founder retains a controlling stake (>50%) for the foreseeable future")
add_bullet("Right of First Refusal (ROFR) on any share sale or transfer by any shareholder")
add_bullet("All IP currently held personally by the Founder to be formally assigned to BitWealth (Pty) Ltd before any equity is issued — this protects all parties")
add_bullet("Vesting on all new shareholder equity (minimum 3-year vesting with a 1-year cliff)")
add_bullet("Reverse vesting: if Simon or RocketX exit before Year 3, unvested shares are forfeited")

add_heading("4.2 Three Equity Options for Discussion", level=2)

# Option A
opt_a = doc.add_paragraph()
oar = opt_a.add_run("Option A — Founder-Led (Conservative)")
oar.bold = True; oar.font.color.rgb = GOLD; oar.font.size = Pt(13)
add_table(
    headers=["Shareholder", "% Equity", "Rationale"],
    rows=[
        ["Founder", "65%", "Reflects existing IP, 1,000+ hours of build, ongoing CEO role"],
        ["Simon", "20%", "Distribution + capital + full-time operations"],
        ["RocketX", "15%", "Licence umbrella; lower % because Founder retains operational control and IP"],
    ],
    col_widths=[4.5, 3, 9]
)
add_para(
    "Best if: RocketX's role is limited to licence + compliance, with no material capital injection. "
    "JR-style economics persist (revenue share) but at a reduced rate of ~10% on top of the equity.",
    italic=True, color=GREY
)

# Option B
doc.add_paragraph()
opt_b = doc.add_paragraph()
obr = opt_b.add_run("Option B — Balanced Three-Way Partnership (Recommended Starting Point)")
obr.bold = True; obr.font.color.rgb = GOLD; obr.font.size = Pt(13)
add_table(
    headers=["Shareholder", "% Equity", "Rationale"],
    rows=[
        ["Founder", "55%", "Retains control; existing IP and ongoing CEO role"],
        ["Simon", "25%", "Distribution + capital + full-time operations"],
        ["RocketX", "20%", "Licence umbrella + active regulatory/governance involvement"],
    ],
    col_widths=[4.5, 3, 9]
)
add_para(
    "Best if: All three parties are actively engaged. RocketX takes a board seat and provides "
    "strategic regulatory guidance beyond pure licence cover. Simon commits both capital and "
    "full-time effort. JR revenue share falls away entirely.",
    italic=True, color=GREY
)

# Option C
doc.add_paragraph()
opt_c = doc.add_paragraph()
ocr = opt_c.add_run("Option C — Capital-Weighted (If RocketX Invests)")
ocr.bold = True; ocr.font.color.rgb = GOLD; ocr.font.size = Pt(13)
add_table(
    headers=["Shareholder", "% Equity", "Rationale"],
    rows=[
        ["Founder", "51%", "Retains bare control"],
        ["Simon", "24%", "Distribution + capital + full-time operations"],
        ["RocketX", "25%", "Licence + regulatory + material capital injection (e.g. R[X]m at the agreed pre-money valuation)"],
    ],
    col_widths=[4.5, 3, 9]
)
add_para(
    "Best if: RocketX commits real capital alongside the licence — for example to fund marketing, "
    "engineering, or operating runway. Equity tracks total economic contribution.",
    italic=True, color=GREY
)

# ============ 5. VALUATION ============
add_heading("5. Valuation Framework", level=1)
add_para(
    "Three independent methods are presented below. The defensible position is the average (or "
    "highest) of the three, depending on negotiating context. As a pre-revenue start-up, "
    "valuation is inherently subjective — the goal is to anchor the discussion with credible "
    "benchmarks rather than arbitrary numbers."
)

add_heading("5.1 Method 1 — Cost-to-Replicate (Sweat Equity)", level=2)
add_para(
    "What would it cost a third party to build what you have built?"
)
add_table(
    headers=["Input", "Value"],
    rows=[
        ["Founder development hours", "1,000+ hours over 7 months"],
        ["Reasonable senior-developer rate (SA, fintech)", "R 700/hour"],
        ["Sweat equity value", "R 700,000"],
        ["Cash invested in infrastructure", "R 15,000"],
        ["Strategy IP premium (proprietary signal logic, backtested)", "R 200,000 – R 500,000"],
        ["Total (range)", "R 915,000 – R 1,215,000"],
    ],
    col_widths=[9, 7.5]
)

add_heading("5.2 Method 2 — Discounted Future Revenue (Bottom-Up)", level=2)
add_para(
    "What is the business worth based on near-term realistic revenue?"
)
add_para("Year 1 base case (Simon's network only):", bold=True)
add_bullet("20 clients × R 100,000 average AUM = R 2,000,000 AUM")
add_bullet("Assumed strategy net return: 25% pa (conservative vs backtest)")
add_bullet("Client profit pool: R 500,000")
add_bullet("Performance fee (10% of profit): R 50,000")
add_bullet("Platform fee (0.75% on AUM): R 15,000")
add_bullet("Total Year 1 revenue: ~R 65,000")

add_para("Year 3 projection (organic growth + Simon's network):", bold=True)
add_bullet("80–120 clients × R 150,000 average AUM = R 12m – R 18m AUM")
add_bullet("Annual revenue at same fee structure: R 400,000 – R 600,000")

add_para(
    "Applying a fintech revenue multiple of 4–6× Year 3 revenue (typical for early-stage SA "
    "fintech with regulatory cover):"
)
add_table(
    headers=["Scenario", "Year 3 Revenue", "Multiple", "Implied Valuation"],
    rows=[
        ["Conservative", "R 400,000", "4×", "R 1,600,000"],
        ["Base", "R 500,000", "5×", "R 2,500,000"],
        ["Optimistic", "R 600,000", "6×", "R 3,600,000"],
    ],
    col_widths=[4, 4, 3, 5.5]
)

add_heading("5.3 Method 3 — Berkus Method (Pre-Revenue Milestone)", level=2)
add_para(
    "A standard pre-revenue valuation method that assigns up to R 500,000 per de-risking milestone:"
)
add_table(
    headers=["Milestone", "Status", "Value"],
    rows=[
        ["Sound idea (basic value, product risk)", "Achieved — validated by backtest", "R 500,000"],
        ["Prototype (technology risk reduced)", "Achieved — full system live, 2 pilot users", "R 500,000"],
        ["Quality management team", "Partial — Founder + KI; pending Simon", "R 250,000"],
        ["Strategic relationships (market risk)", "Achieved on partnership signing — RocketX + Simon", "R 500,000"],
        ["Product rollout / sales (production risk)", "Pending — pilot only", "R 100,000"],
        ["Total Berkus valuation", "", "R 1,850,000"],
    ],
    col_widths=[7, 6, 3.5]
)

add_heading("5.4 Recommended Valuation Range", level=2)
add_table(
    headers=["Method", "Valuation"],
    rows=[
        ["Cost-to-Replicate", "R 915k – R 1.2m"],
        ["Discounted Future Revenue", "R 1.6m – R 3.6m"],
        ["Berkus (Pre-Revenue)", "R 1.85m"],
        ["Recommended pre-money valuation for discussion", "R 2.0m – R 2.5m"],
    ],
    col_widths=[9, 7.5]
)
add_para(
    "Anchor the negotiation at R 2.5m and accept anywhere down to R 1.8m. Anything below that "
    "undervalues the IP, the regulatory groundwork already done, and the operational platform. "
    "A successful 12-month JR period strengthens this valuation materially at Phase 2 — proven "
    "revenue, operational track record and a clean regulatory record all support the upper end "
    "of the range.",
    bold=True
)

# ============ 6. QUESTIONS FOR ROCKETX ============
add_heading("6. Questions to Ask the RocketX CEO", level=1)
add_para(
    "Before the equity discussion can be finalised, the following questions should be answered."
)

add_heading("6.1 Capital and Investment", level=2)
add_bullet("Will RocketX inject capital into BitWealth, or is the contribution purely the licence and compliance services?")
add_bullet("If capital is on the table, what amount and on what terms (equity vs convertible loan vs revenue advance)?")
add_bullet("Would RocketX be willing to fund the first 12 months of marketing and operating costs in exchange for additional equity?")

add_heading("6.2 Entity Structure", level=2)
add_bullet("Will RocketX (Pty) Ltd be the shareholder, or will the CEO take shares personally?")
add_bullet("If corporate, are there existing shareholders/board members in RocketX who must approve this transaction?")
add_bullet("Are there any restrictions in RocketX's existing FSP licence that would limit it from holding equity in another FSP applicant?")

add_heading("6.3 Operational Role", level=2)
add_bullet("Will RocketX want a board seat? Voting or observer?")
add_bullet("Beyond the licence, what active role does the CEO see RocketX playing — strategic guidance, introductions, infrastructure sharing?")
add_bullet("Will RocketX provide the Key Individual on a permanent basis, or only during the JR phase?")

add_heading("6.4 JR Operations and Renewal", level=2)
add_bullet("Does RocketX's current FSP licence have the right categories (and CASP authorisation, if applicable) to host BitWealth as JR for the LTH PVR strategy?")
add_bullet("Who would you propose as the assigned KI, and what is their documented capacity for ongoing supervision of this strategy?")
add_bullet("What are RocketX's JR onboarding fees and ongoing supervision fees, if any?")
add_bullet("What is your preferred PI insurance arrangement — extension of RocketX's policy, or BitWealth carries its own?")
add_bullet("What is the formal renewal/non-renewal process and notice mechanism?")

add_heading("6.5 JR-to-Equity Transition (Phase 2)", level=2)
add_bullet("What is the trigger to open Phase 2 discussions — we propose 90 days before each annual renewal date?")
add_bullet("On Phase 2 effective date, the JR fee-share falls away and is replaced by intra-group operating arrangements — agreed in principle?")
add_bullet("Are there any RocketX board or shareholder constraints around taking equity in a JR's underlying IP company?")

add_heading("6.6 Exit and Protections", level=2)
add_bullet("Will RocketX accept reverse vesting (i.e. equity vests over 3 years; unvested shares forfeited if RocketX exits early)?")
add_bullet("Tag-along, drag-along, and ROFR clauses — acceptable in principle?")
add_bullet("Founder veto rights on changes to strategy, capital structure, or sale of business — acceptable?")

# ============ 7. RECOMMENDED PATH ============
add_heading("7. Recommended Path Forward", level=1)
add_table(
    headers=["#", "Action", "Owner", "Timing"],
    rows=[
        ["1", "Sign IP Assignment Agreement transferring all LTH PVR IP from Founder to BitWealth (Pty) Ltd", "Founder + Attorney", "Before any partnership discussion"],
        ["2", "Execute JR Agreement; RocketX lodges JR appointment with FSCA", "All parties + Attorney", "On agreement of terms"],
        ["3", "KI onboarding, PI insurance arrangement, client mandate templates approved by KI", "RocketX KI + BitWealth", "In parallel with FSCA approval (30–60 days)"],
        ["4", "Phase 1 go-live: begin onboarding clients under JR appointment", "BitWealth + Simon", "Immediately on FSCA approval"],
        ["5", "Operate Year 1 as JR; review performance, AUM, regulatory record", "All parties", "Months 1–9 of JR term"],
        ["6", "Phase 2 window opens: mutually decide to (a) renew JR for Year 2, (b) sign SHA and supersede JR, or (c) part ways at renewal", "All parties + Attorney", "90 days before JR renewal date"],
        ["7", "Submit own FSP licence application (running in parallel as continuity option)", "Founder + Compliance Officer", "Within 12 months"],
    ],
    col_widths=[1, 8, 4, 4]
)

# ============ 8. RISKS ============
add_heading("8. Key Risks to Manage", level=1)
add_table(
    headers=["Risk", "Mitigant"],
    rows=[
        ["Founder loses control through dilution at Phase 2", "Cap RocketX + Simon combined at <50% in initial round; supermajority required for further dilution"],
        ["RocketX uses JR-licence dependency as leverage at renewal", "Explicit right in JR Agreement for BitWealth to apply for its own FSP or move to another FSP; pursue own FSP licence in parallel from Year 1"],
        ["FSCA delays in approving the JR appointment", "Lodge appointment immediately after signing; agree interim arrangements; allow 30–60 day window in go-live planning"],
        ["Assigned KI lacks capacity to supervise the strategy properly", "Document KI capacity and supervision plan in the JR Agreement; named KI with stated hours allocated"],
        ["Non-renewal of JR by RocketX (single-point-of-failure)", "90-day notice period; client-portability clause; BitWealth's own FSP application running in parallel"],
        ["Conflicts of interest between FSP fee-share and client interests", "Documented COI management plan; full client disclosure of fee-share and any equity interests"],
        ["Simon underdelivers on distribution", "At Phase 2: vesting schedule with performance milestones; reverse vesting if AUM targets missed"],
        ["IP disputes if Founder exits", "All IP assigned to BitWealth (Pty) Ltd before any equity is issued"],
        ["Client portability if JR ends", "Client mandates explicitly with BitWealth as JR; written portability clause; clients informed of the JR/FSP structure on onboarding"],
    ],
    col_widths=[6, 10.5]
)

# ============ FOOTER ============
doc.add_paragraph()
foot = doc.add_paragraph()
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = foot.add_run(
    "This document is a discussion framework prepared by the Founder of BitWealth (Pty) Ltd. "
    "All terms are non-binding and subject to negotiation and formal legal documentation. "
    "Final agreements should be drafted and reviewed by attorneys experienced in financial "
    "services partnerships and shareholder agreements in South Africa."
)
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = GREY

# Save
out_path = r"docs\BitWealth_Partnership_Proposal.docx"
doc.save(out_path)
print(f"Document saved: {out_path}")
