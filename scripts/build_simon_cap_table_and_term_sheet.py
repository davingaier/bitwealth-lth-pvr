"""Generate two documents for the Simon shareholding negotiation:
1. BitWealth_Simon_Cap_Table_Model.docx - cap table at each AUM tranche
2. BitWealth_Simon_Term_Sheet.docx - 2-page term sheet for Simon
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Brand palette
GOLD = RGBColor(0xC9, 0xA2, 0x27)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xC0, 0x00, 0x00)
LIGHT_GOLD_HEX = "F5EBC8"
LIGHT_GREY_HEX = "F2F2F2"
LIGHT_GREEN_HEX = "E8F5E9"

OUT_DIR = Path(r"docs\Shareholding")
OUT_DIR.mkdir(parents=True, exist_ok=True)


# ---------- Shared helpers ----------
def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def setup_doc():
    doc = Document()
    for s in doc.sections:
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)
        s.top_margin = Cm(1.8)
        s.bottom_margin = Cm(1.8)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    return doc


def add_heading(doc, text, level=1, color=GOLD):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
        run.font.name = 'Calibri'
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=DARK, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, header_fill=LIGHT_GOLD_HEX, col_widths=None,
              highlight_rows=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        shade(hdr[i], header_fill)
    highlight_rows = highlight_rows or {}
    for ri, row_data in enumerate(rows):
        cells = tbl.add_row().cells
        fill = highlight_rows.get(ri)
        for ci, val in enumerate(row_data):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if fill:
                shade(cells[ci], fill)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl


# =============================================================
# DOC 1 — CAP TABLE MODEL
# =============================================================
def build_cap_table():
    doc = setup_doc()

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BitWealth (Pty) Ltd")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = GOLD
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("Simon Hobday — Equity Earn-In Model")
    r.font.size = Pt(16); r.font.color.rgb = DARK
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = m.add_run("Cap-Table Walk-Through at Each AUM Milestone\n"
                  "Discussion document — May 2026")
    r.font.size = Pt(10); r.italic = True; r.font.color.rgb = GREY
    doc.add_paragraph()

    # Disclaimer
    add_para(doc,
        "This document models the cap-table position of BitWealth (Pty) Ltd at each AUM "
        "milestone under the proposed equity earn-in arrangement for Simon Hobday. All "
        "numbers are illustrative; final terms to be documented in a Shareholders' "
        "Agreement reviewed by attorneys. Based on the structural advice received from "
        "Tremyne (CA(SA)) on 17 May 2026: single ordinary share class, no holding "
        "company, IP transferred into the company as a shareholder loan from the Founder.",
        italic=True, color=GREY)

    # ============ 1. STARTING POSITION ============
    add_heading(doc, "1. Starting Position (Pre-Simon Subscription)", level=1)
    add_para(doc,
        "BitWealth (Pty) Ltd currently has 100 ordinary shares in issue, all held by the "
        "Founder. The Founder's IP is transferred into the company at an agreed value, "
        "creating a shareholder loan repayable to the Founder.")
    add_table(doc,
        headers=["Item", "Value (ZAR)", "Comment"],
        rows=[
            ["Issued shares", "100", "Sole shareholder: Davin Gaier"],
            ["IP agreed value (recorded on balance sheet)", "R 2,500,000", "Working assumption — adjust by agreement with Simon"],
            ["Corresponding shareholder loan to Founder", "R 2,500,000", "Repayable per IP Transfer Agreement; amortised over 15 years for tax"],
            ["Implied pre-money equity value (going-concern)", "R 4,500,000", "Distinct from IP value; basis for Simon's subscription price"],
            ["Implied post-money equity value (after Simon's R 500k)", "R 5,000,000", "R 500k / 10% = R 5m post-money"],
        ],
        col_widths=[7.0, 3.5, 7.0])

    # ============ 2. AT INITIAL SUBSCRIPTION ============
    add_heading(doc, "2. At Initial Subscription (Day 1)", level=1)
    add_para(doc,
        "Simon subscribes for R 500,000 in equity and advances R 500,000 as a shareholder "
        "loan. New shares are issued to Simon, diluting the Founder to 90%. Both "
        "shareholders now hold loan accounts against the company.")
    add_table(doc,
        headers=["Shareholder", "Shares", "Equity %", "Capital contributed", "Loan account (ZAR)"],
        rows=[
            ["Founder (Davin Gaier)", "100", "90.00%", "IP @ R 2.5m + sweat", "R 2,500,000 (IP)"],
            ["Simon Hobday", "11.11 (rounded — see note)", "10.00%", "R 500,000 cash", "R 500,000 (cash loan)"],
            ["Total", "111.11", "100.00%", "—", "R 3,000,000"],
        ],
        col_widths=[5.5, 3.0, 2.5, 4.0, 4.0],
        highlight_rows={2: LIGHT_GOLD_HEX})
    add_para(doc,
        "Note on share-count mechanics: in practice we would re-denominate the share register "
        "(e.g. issue 1,000 shares to the Founder pre-subscription, then issue 111 to Simon to "
        "give 10%) to avoid fractional shares. The percentages are what matter; share counts "
        "are mechanical and can be set by the company secretary at incorporation of the SHA.",
        italic=True, color=GREY)

    # ============ 3. CAP TABLE AT EACH TRANCHE ============
    add_heading(doc, "3. Cap Table at Each AUM Milestone (Option 2 — Recommended)", level=1)
    add_para(doc,
        "The 30% earn-in is split into 6 tranches of 5% each, triggered by escalating "
        "qualifying AUM thresholds. Each tranche is issued as a new ordinary share "
        "subscription at nominal value at the first quarter-end after the threshold is "
        "met. Shares once issued do not reverse if AUM later declines.")
    add_table(doc,
        headers=["Tranche", "Qualifying AUM (ZAR)", "Equity issued in tranche", "Simon cumulative", "Founder cumulative"],
        rows=[
            ["Start (Day 1)", "n/a (subscription)", "10.00%", "10.00%", "90.00%"],
            ["1", "R 10,000,000", "5.00%", "15.00%", "85.00%"],
            ["2", "R 20,000,000", "5.00%", "20.00%", "80.00%"],
            ["3", "R 35,000,000", "5.00%", "25.00%", "75.00%"],
            ["4", "R 55,000,000", "5.00%", "30.00%", "70.00%"],
            ["5", "R 80,000,000", "5.00%", "35.00%", "65.00%"],
            ["6 (cap)", "R 110,000,000", "5.00%", "40.00%", "60.00%"],
        ],
        col_widths=[2.5, 4.5, 4.0, 3.5, 3.5],
        highlight_rows={6: LIGHT_GREEN_HEX})
    add_para(doc,
        "At full earn-in, Simon holds 40% and the Founder retains 60% — exactly the "
        "control threshold advised by Tremyne. Earn-in window: 5 years from SHA effective "
        "date; any unearned tranches lapse on expiry.")

    # ============ 4. ECONOMIC POSITION AT EACH TRANCHE ============
    add_heading(doc, "4. Economic Position at Each Tranche", level=1)
    add_para(doc,
        "The table below illustrates how the IP shareholder loan to the Founder protects "
        "his economics. In any liquidity event (dividend, sale, wind-up), shareholder "
        "loans are repaid in full before equity holders share in remaining proceeds.")
    add_table(doc,
        headers=["Stage", "Founder equity %", "Simon equity %", "Founder loan o/s", "Simon loan o/s"],
        rows=[
            ["Day 1", "90.00%", "10.00%", "R 2,500,000", "R 500,000"],
            ["After Tranche 1", "85.00%", "15.00%", "(amortising)", "(repaying)"],
            ["After Tranche 3", "75.00%", "25.00%", "(amortising)", "(repaying)"],
            ["After Tranche 6 (cap)", "60.00%", "40.00%", "Per IP schedule", "Per loan schedule"],
        ],
        col_widths=[4.0, 3.5, 3.5, 3.5, 3.5])
    add_para(doc, "Loan repayment mechanics (proposed):", bold=True)
    add_bullet(doc, "Founder IP loan: amortised over 15 years per tax/accounting schedule; cash repayments only from declared surpluses, subordinated to operating-capital needs")
    add_bullet(doc, "Simon cash loan: repaid at SA prime −1% over 24 months from first cash surplus; suspended if cash flow insufficient")
    add_bullet(doc, "Both loans rank ahead of dividend declarations; neither shareholder may force repayment if it would render the company commercially insolvent")

    # ============ 5. WORKED EXIT EXAMPLE ============
    add_heading(doc, "5. Worked Example — Exit Sale at Year 5", level=1)
    add_para(doc,
        "Illustrative scenario: at Year 5, Simon has hit Tranches 1–4 (Simon 30%, "
        "Founder 70%). Company is sold for R 30,000,000 enterprise value. Loan accounts "
        "have been partially repaid (Founder R 1m remaining, Simon R 0 remaining).")
    add_table(doc,
        headers=["Step", "Calculation", "Amount (ZAR)"],
        rows=[
            ["1. Enterprise value", "Sale price", "R 30,000,000"],
            ["2. Repay outstanding loans first", "Founder R 1m, Simon R 0", "(R 1,000,000)"],
            ["3. Net proceeds to equity holders", "R 30m − R 1m", "R 29,000,000"],
            ["4. Founder share (70%)", "70% × R 29m", "R 20,300,000"],
            ["5. Simon share (30%)", "30% × R 29m", "R 8,700,000"],
            ["Founder total recovery", "Loan repayment + equity", "R 21,300,000"],
            ["Simon total recovery", "Loan repayment + equity", "R 8,700,000"],
        ],
        col_widths=[5.5, 6.0, 5.0],
        highlight_rows={5: LIGHT_GREEN_HEX, 6: LIGHT_GREEN_HEX})
    add_para(doc,
        "Note: this assumes a clean asset sale with no CGT, dividends tax or transaction "
        "costs. In practice the after-tax position will differ. The point of the worked "
        "example is to show how the IP loan protects Founder economics even after dilution.",
        italic=True, color=GREY)

    # ============ 6. SENSITIVITY ============
    add_heading(doc, "6. Sensitivity — What Simon Earns at Different AUM Outcomes", level=1)
    add_para(doc,
        "Simon's eventual equity stake depends entirely on AUM delivered. The table "
        "below shows his final cumulative stake under different AUM scenarios reached "
        "within the 5-year earn-in window.")
    add_table(doc,
        headers=["Qualifying AUM achieved", "Tranches earned", "Simon final equity %", "Founder final equity %"],
        rows=[
            ["< R 10m", "None (only initial subscription)", "10.00%", "90.00%"],
            ["R 10m – R 19m", "1", "15.00%", "85.00%"],
            ["R 20m – R 34m", "2", "20.00%", "80.00%"],
            ["R 35m – R 54m", "3", "25.00%", "75.00%"],
            ["R 55m – R 79m", "4", "30.00%", "70.00%"],
            ["R 80m – R 109m", "5", "35.00%", "65.00%"],
            ["R 110m or more", "6 (cap)", "40.00%", "60.00%"],
        ],
        col_widths=[5.0, 4.5, 3.5, 3.5])

    # ============ 7. KEY ASSUMPTIONS TO CONFIRM ============
    add_heading(doc, "7. Key Assumptions for Discussion with Simon", level=1)
    add_table(doc,
        headers=["Assumption", "Proposed value", "Open for discussion?"],
        rows=[
            ["IP agreed value", "R 2,500,000", "Yes — material to Founder loan size"],
            ["Going-concern post-money equity valuation", "R 5,000,000", "Yes — sets Simon's effective entry price"],
            ["Simon equity ceiling", "40% (Founder retains ≥60% control)", "Hard cap per CA advice"],
            ["Tranche size", "5% per tranche × 6 tranches", "Yes — could be 7.5% × 4, etc."],
            ["AUM curve shape", "Accelerating (R 10m, 20m, 35m, 55m, 80m, 110m)", "Yes — Simon may prefer flatter steps"],
            ["Earn-in window", "5 years from SHA effective date", "Yes — 3 / 4 / 5 / 7 years all possible"],
            ["6-month retention rule for qualifying AUM", "Required", "Recommended — protects against churn-gaming"],
            ["Loan repayment cadence", "Founder: 15 yrs amortised. Simon: 24 mths from surplus", "Yes"],
            ["Bad-leaver share buyback (Yr 1–4)", "Vested = subscription price; unvested = nominal", "Yes — standard"],
            ["Good-leaver share treatment", "Vested = market value; unvested = forfeit", "Yes — standard"],
            ["Founder reverse vesting", "None — IP loan provides commitment lock", "Yes — Simon may push back"],
            ["Future opcos (Bitcoin RA, leverage, etc.)", "Separate equity discussions per opco", "Critical — agree principle now"],
        ],
        col_widths=[6.0, 6.0, 4.5])

    out = OUT_DIR / "BitWealth_Simon_Cap_Table_Model.docx"
    doc.save(out)
    print(f"Saved: {out}")


# =============================================================
# DOC 2 — 2-PAGE TERM SHEET
# =============================================================
def build_term_sheet():
    doc = setup_doc()

    # Header
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("BitWealth (Pty) Ltd")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = GOLD
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("SHAREHOLDERS' AGREEMENT — INDICATIVE TERM SHEET")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = m.add_run("Non-binding discussion document  |  May 2026  |  Davin Gaier  ↔  Simon Hobday")
    r.font.size = Pt(9); r.italic = True; r.font.color.rgb = GREY

    # Status banner
    p = doc.add_paragraph()
    r = p.add_run(
        "This term sheet is non-binding and is provided as a basis for negotiation only. "
        "It does not constitute an offer or commitment. Final terms to be documented in "
        "a Shareholders' Agreement drafted and reviewed by attorneys. Conditional on "
        "execution of the IP Transfer Agreement and finalisation of the CAEP Juristic "
        "Representative appointment."
    )
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = RED

    add_table(doc,
        headers=["Term", "Proposed Position"],
        rows=[
            ["Company",
             "BitWealth (Pty) Ltd (Registration No. 2026/090346/07), a South African private company."],
            ["Parties",
             "(a) Davin Harald Gaier (Founder); (b) Simon Hobday (Distribution Director); (c) BitWealth (Pty) Ltd."],
            ["Conditions Precedent",
             "(i) IP Transfer Agreement signed and dated on or before 20 February for FY2026 inclusion; "
             "(ii) FSCA appointment of BitWealth (Pty) Ltd as Juristic Representative under CAEP confirmed; "
             "(iii) Both parties' fit-and-proper documentation completed."],
            ["Share Capital",
             "Single class of ordinary shares. Existing 100 (or re-denominated 1,000) ordinary shares held by the Founder pre-subscription."],
            ["Simon's Subscription (Day 1)",
             "Simon subscribes for ordinary shares representing 10.00% of the post-subscription issued capital, for R 500,000 in cash."],
            ["Simon's Shareholder Loan",
             "Simon advances an additional R 500,000 to the Company as an interest-bearing shareholder loan at SA prime −1%, repayable over 24 months from first cash surplus, subordinated to operating-capital needs."],
            ["Founder's IP Loan",
             "Founder transfers the LTH PVR intellectual property (per Annexure A of the IP Transfer Agreement) into the Company at an agreed value of R 2,500,000, recorded as a shareholder loan repayable to the Founder, amortised over 15 years for tax purposes."],
            ["Earn-In Mechanism",
             "Simon may earn up to a further 30.00% (cumulative cap: 40.00%) by introducing Qualifying AUM, in 6 tranches of 5.00% each, triggered at the following Qualifying AUM thresholds: R 10m / R 20m / R 35m / R 55m / R 80m / R 110m. Tranches issued at nominal value at the first quarter-end after the threshold is achieved."],
            ["Qualifying AUM",
             "Aggregate ZAR value of client assets (i) introduced by Simon (CRM-evidenced within 48 hours of first contact), (ii) retained for ≥ 6 months, (iii) measured as a rolling 30-day average, (iv) net of withdrawals, (v) excluding Simon's own and related-party capital."],
            ["Earn-In Window",
             "5 years from the SHA effective date. Any unearned tranches lapse on expiry."],
            ["Founder Equity Floor",
             "Founder shall at all times hold at least 60.00% of the issued ordinary shares (the maximum Simon can hold is 40.00%)."],
            ["Reverse Vesting (Simon — Initial 10%)",
             "Simon's initial 10% reverse-vests over 4 years with a 1-year cliff. Bad-leaver: unvested shares forfeit for R 100; vested shares may be repurchased at subscription price. Good-leaver: unvested forfeit; vested at fair market value."],
            ["Reverse Vesting (Founder)",
             "None. Founder commitment is provided by (i) the IP shareholder loan (subordinated on Founder bad-leaver event), (ii) Founder restraint of trade, (iii) buy-and-sell life policy."],
            ["Bad Leaver",
             "Voluntary resignation within 4 years, dismissal for cause (fraud, gross negligence, material SHA breach), breach of restraint of trade."],
            ["Good Leaver",
             "Death, permanent disability, mutual termination, removal without cause, sale of company."],
            ["Working Arrangement",
             "Both parties to work part-time on the Company until steady revenue is achieved. No salary payable to either party until the Company has sustained positive operating cash flow, at which point market-related remuneration will be agreed."],
            ["Restraint of Trade",
             "Each party shall be bound by a 12-month post-departure restraint covering competing crypto-asset management activities in South Africa. Reasonable in scope and duration; to be confirmed by attorneys."],
            ["Governance",
             "Board comprises Founder (Chair, casting vote) and Simon. Monthly management meetings; quarterly board meetings. Founder retains operational control while equity floor (60%) is maintained."],
            ["Reserved Matters (require Simon's consent)",
             "Issue of new shares (other than earn-in tranches); change to share class structure; sale of all or substantially all assets; voluntary winding-up; material change in business activity; transactions with related parties exceeding R 250,000."],
            ["Dividend Policy",
             "Dividends declared at Board's discretion only after (i) shareholder loans are current, (ii) regulatory capital requirements are met, (iii) 12-month operating reserve is maintained. Pro-rata to equity holdings."],
            ["Right of First Refusal (ROFR)",
             "Neither party may transfer shares to a third party without first offering them to the other on the same terms. Founder approval required for any transfer."],
            ["Tag-Along",
             "If the Founder sells more than 50% of his shares to a third party, Simon may tag along on the same terms pro-rata."],
            ["Drag-Along",
             "If shareholders holding ≥ 75% accept a bona fide third-party offer for 100% of the Company, remaining shareholder(s) may be dragged on the same terms."],
            ["Investment via Entity",
             "Simon may hold his shares via a wholly-owned company (including foreign), provided (i) transfer restrictions and ROFR apply on a look-through basis, (ii) Simon personally signs restraints and key-person commitments, (iii) tax-residency disclosures are provided."],
            ["Future Products / Opcos",
             "Future strategies (Bitcoin Retirement Annuity, leverage strategies, RV trades, international fund) shall be housed in separate operating companies, each with its own IP and shareholding structure. Simon's participation in such opcos to be negotiated separately on a per-opco basis; no automatic right to participate."],
            ["Confidentiality",
             "Both parties bound by confidentiality during and for 3 years after the SHA terminates."],
            ["Costs",
             "Each party bears its own legal costs. Stamp duty / CIPC fees borne by the Company."],
            ["Governing Law",
             "Laws of the Republic of South Africa. Disputes referred to arbitration under AFSA rules, seat Johannesburg."],
            ["Exclusivity / Timing",
             "Both parties agree to negotiate exclusively for 30 days from signature of this term sheet, targeting full SHA execution within 60 days."],
        ],
        col_widths=[4.5, 12.5])

    # Signature block
    doc.add_paragraph()
    add_para(doc, "Acknowledged and agreed (non-binding):", bold=True, size=10)
    doc.add_paragraph()
    sig = doc.add_table(rows=4, cols=2)
    sig.rows[0].cells[0].text = "Davin Harald Gaier"
    sig.rows[0].cells[1].text = "Simon Hobday"
    sig.rows[1].cells[0].text = "Founder, BitWealth (Pty) Ltd"
    sig.rows[1].cells[1].text = "Distribution Director (proposed)"
    sig.rows[2].cells[0].text = "Signature: _____________________________"
    sig.rows[2].cells[1].text = "Signature: _____________________________"
    sig.rows[3].cells[0].text = "Date: _____________"
    sig.rows[3].cells[1].text = "Date: _____________"
    for row in sig.rows:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    out = OUT_DIR / "BitWealth_Simon_Term_Sheet.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    build_cap_table()
    build_term_sheet()
