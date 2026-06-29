"""
BitWealth (Pty) Ltd — Simon Hobday Revised Equity Term Sheet v2
Generates two documents:
  1. docs/Shareholding/BitWealth_Simon_Cap_Table_v2.docx
  2. docs/Shareholding/BitWealth_Simon_Term_Sheet_v2.docx

Key terms:
  · Simon contributes NO capital
  · 10% founding equity (business development / client acquisition role)
  · 30% earn-in: 3 × 10% tranches at R50m / R100m / R200m AUM milestones
  · Each milestone has 24-month rolling deadline
  · FSP: Finova Capital (not CAEP)
  · IP value on balance sheet: R2,500,000 (contributed by Davin Cloete)
  · Protections: good/bad leaver, non-compete, drag-along/tag-along, regulatory clawback
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date

OUT_DIR = Path("docs") / "Shareholding"
OUT_DIR.mkdir(parents=True, exist_ok=True)

try:
    TODAY = date.today().strftime("%d %B %Y").lstrip("0")
except Exception:
    TODAY = "29 June 2026"

GOLD  = RGBColor(0xC9, 0xA2, 0x27)
DARK  = RGBColor(0x1A, 0x1A, 0x1A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY  = RGBColor(0x55, 0x55, 0x55)
L_GOLD_HEX = "F5EBC8"
GOLD_HEX   = "C9A227"
DARK_HEX   = "1A1A1A"
TEAL_HEX   = "006064"
RED_HEX    = "B71C1C"
GREEN_HEX  = "1B5E20"
LGREY_HEX  = "F5F5F5"
MIDGREY_HEX= "DDDDDD"


# ═══════════════════════════════════════════════════════════════
# DOCX HELPERS
# ═══════════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        val = kwargs.get(side, {"val": "single", "sz": "4", "color": "CCCCCC"})
        el = OxmlElement(f"w:{side}")
        for k, v in val.items():
            el.set(qn(f"w:{k}"), str(v))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, size=10, color=None, align=None, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else DARK
    return run


def add_heading(doc, text, level=1, gold=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper() if level == 1 else text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 11)
    run.font.color.rgb = GOLD if gold else DARK
    if level == 1:
        run.font.name = "Calibri"
    p.paragraph_format.keep_with_next = True
    return p


def add_paragraph(doc, text, size=10, color=None, italic=False, space_before=2, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color if color else DARK
    run.italic = italic
    return p


def add_bullet(doc, text, size=10, indent=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    return p


def hr(doc):
    """Thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), GOLD_HEX)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def doc_header(doc, title, subtitle, date_str):
    """Cover / header block with gold banner."""
    # Gold title bar
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(title)
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = GOLD; run.font.name = "Calibri"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p2 = doc.add_paragraph()
    run2 = p2.add_run(subtitle)
    run2.font.size = Pt(11); run2.font.color.rgb = DARK
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)

    p3 = doc.add_paragraph()
    run3 = p3.add_run(f"Date: {date_str}  ·  PRIVATE & CONFIDENTIAL  ·  NON-BINDING")
    run3.font.size = Pt(9); run3.font.color.rgb = GREY; run3.italic = True
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(10)

    hr(doc)


# ═══════════════════════════════════════════════════════════════════════════
# DOCUMENT 1 — CAP TABLE MODEL
# ═══════════════════════════════════════════════════════════════════════════
def build_cap_table(doc):
    doc_header(
        doc,
        "BITWEALTH (PTY) LTD",
        "Part A — Cap Table Model & Earn-In Schedule",
        TODAY,
    )

    # ── Background ──────────────────────────────────────────────────────
    add_heading(doc, "1.  Background & IP Foundation")
    add_paragraph(doc,
        "BitWealth (Pty) Ltd ('the Company') has been incorporated to operate a Bitcoin "
        "Dollar-Cost Averaging (DCA) investment management platform regulated under the "
        "Financial Advisory and Intermediary Services Act, 37 of 2002 (FAIS Act). "
        "The Company will be appointed as a Juristic Representative (JR) under the FSP "
        "licence of Finova Capital (Pty) Ltd.")
    add_paragraph(doc,
        "The proprietary technology platform (BitWealth LTH PVR System) has been "
        "transferred to the Company's balance sheet at an agreed fair value of "
        "R2,500,000 (Two Million Five Hundred Thousand Rand), contributed by "
        "Davin Cloete as intellectual property consideration. The IP transfer agreement "
        "has been executed and the asset is recorded in the Company's books.")

    # Share structure note
    add_paragraph(doc,
        "The Company has authorised and issued 1,000 ordinary shares of R1.00 par value each. "
        "All issued shares carry equal voting rights, dividend rights, and rights upon winding-up. "
        "Davin Cloete currently holds all 1,000 issued shares.",
        italic=True, color=GREY, size=9)

    hr(doc)

    # ── Helper to add a cap table stage ─────────────────────────────────
    def cap_table(title, rows, note=None, header_bg=TEAL_HEX):
        add_heading(doc, title, level=2, gold=False)
        table = doc.add_table(rows=len(rows)+1, cols=5)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Set column widths
        widths = [Cm(5.5), Cm(3.0), Cm(2.5), Cm(2.5), Cm(4.0)]
        for i, col in enumerate(table.columns):
            for cell in col.cells:
                cell.width = widths[i]

        # Header row
        hdrs = ["Shareholder", "Shares Held", "% Holding", "Change", "Notes"]
        for i, hdr in enumerate(hdrs):
            cell_text(table.cell(0, i), hdr, bold=True, size=9,
                      color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_bg(table.cell(0, i), header_bg)

        # Data rows
        for r_idx, row in enumerate(rows, 1):
            bg = L_GOLD_HEX if row[0].startswith("TOTAL") else (LGREY_HEX if r_idx % 2 == 0 else "FFFFFF")
            bold = row[0].startswith("TOTAL")
            aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                      WD_ALIGN_PARAGRAPH.LEFT]
            for c_idx, (val, al) in enumerate(zip(row, aligns)):
                cell_text(table.cell(r_idx, c_idx), str(val), bold=bold, size=10,
                          color=DARK, align=al)
                set_cell_bg(table.cell(r_idx, c_idx), bg)

        if note:
            p = doc.add_paragraph()
            run = p.add_run(f"★  {note}")
            run.italic = True; run.font.size = Pt(9); run.font.color.rgb = GREY
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(8)

    # ── Stage 0: Current ────────────────────────────────────────────────
    cap_table(
        "Stage 0 — Current Shareholding (Pre-Signing)",
        [
            ["Davin Cloete (Founder)",  "1,000", "100%", "—",  "IP contributor; technology platform"],
            ["Simon Hobday",            "Nil",   "—",    "—",  "Not yet a shareholder"],
            ["TOTAL",                   "1,000", "100%", "—",  ""],
        ],
        note="IP value R2,500,000 contributed by Davin Cloete and recorded on the balance sheet.",
        header_bg=DARK_HEX,
    )

    # ── Stage 1: On Signing ─────────────────────────────────────────────
    cap_table(
        "Stage 1 — On Signing (Initial Equity Grant)",
        [
            ["Davin Cloete (Founder)", "900", "90%",  "−100 shares", "Technology, strategy, capital"],
            ["Simon Hobday",           "100", "10%",  "+100 shares", "Business development / client acquisition"],
            ["TOTAL",                  "1,000","100%", "—",          ""],
        ],
        note="Simon's 10% is issued for no monetary consideration (sweat equity). "
             "Notional value at IP-implied share price: R250,000.",
        header_bg=TEAL_HEX,
    )

    # ── Stage 2: Post Milestone 1 ────────────────────────────────────────
    cap_table(
        "Stage 2 — Post Milestone 1  (R50m AUM achieved within 24 months of signing)",
        [
            ["Davin Cloete (Founder)", "800", "80%", "−100 shares", "Transfers from own holding"],
            ["Simon Hobday",           "200", "20%", "+100 shares", "Tranche 1 earned"],
            ["TOTAL",                  "1,000","100%","—",          ""],
        ],
        note="Tranche 1: 100 shares transfer from Davin to Simon on verified AUM ≥ R50m.",
        header_bg=GREEN_HEX,
    )

    # ── Stage 3: Post Milestone 2 ────────────────────────────────────────
    cap_table(
        "Stage 3 — Post Milestone 2  (R100m AUM achieved within 48 months of signing)",
        [
            ["Davin Cloete (Founder)", "700", "70%", "−100 shares", "Transfers from own holding"],
            ["Simon Hobday",           "300", "30%", "+100 shares", "Tranche 2 earned"],
            ["TOTAL",                  "1,000","100%","—",          ""],
        ],
        note="Tranche 2: 100 shares transfer on verified AUM ≥ R100m.",
        header_bg=GREEN_HEX,
    )

    # ── Stage 4: Post Milestone 3 (Full Earn-In) ─────────────────────────
    cap_table(
        "Stage 4 — Post Milestone 3  (R200m AUM achieved within 72 months of signing)",
        [
            ["Davin Cloete (Founder)", "600", "60%", "−100 shares", "Final position if all milestones met"],
            ["Simon Hobday",           "400", "40%", "+100 shares", "Tranche 3 earned — maximum earn-in reached"],
            ["TOTAL",                  "1,000","100%","—",          ""],
        ],
        note="Full earn-in: Simon holds 40% on achievement of all three milestones.",
        header_bg=GOLD_HEX,
    )

    hr(doc)

    # ── Lapse Scenarios ──────────────────────────────────────────────────
    add_heading(doc, "2.  Milestone Lapse Scenarios", level=2, gold=False)
    add_paragraph(doc,
        "The table below shows the resulting shareholding for each combination of milestone "
        "outcomes. Each tranche is independent — missing one milestone does not automatically "
        "prevent earning a later tranche, provided the later milestone is achieved within its "
        "own 24-month window.")

    lapse_table = doc.add_table(rows=6, cols=6)
    lapse_table.style = "Table Grid"
    lapse_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    l_hdrs = ["M1 (R50m)", "M2 (R100m)", "M3 (R200m)", "Davin", "Simon", "Outcome"]
    l_widths = [Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.0), Cm(2.0), Cm(4.5)]
    for i, col in enumerate(lapse_table.columns):
        for cell in col.cells:
            cell.width = l_widths[i]

    for i, h in enumerate(l_hdrs):
        cell_text(lapse_table.cell(0, i), h, bold=True, size=9,
                  color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(lapse_table.cell(0, i), TEAL_HEX)

    lapse_rows = [
        ["✓ Earned", "✓ Earned", "✓ Earned", "60%", "40%", "Best case — full earn-in"],
        ["✓ Earned", "✓ Earned", "✗ Lapsed", "70%", "30%", "M3 missed or lapsed"],
        ["✓ Earned", "✗ Lapsed", "✗ Lapsed", "80%", "20%", "Only M1 achieved"],
        ["✗ Lapsed", "✗ Lapsed", "✗ Lapsed", "90%", "10%", "No milestones achieved"],
        ["✗ Lapsed", "✓ Earned", "✓ Earned", "70%", "30%", "M1 missed but M2 & M3 achieved"],
    ]
    for r_idx, row in enumerate(lapse_rows, 1):
        bg = L_GOLD_HEX if r_idx == 1 else (LGREY_HEX if r_idx % 2 == 0 else "FFFFFF")
        for c_idx, val in enumerate(row):
            color = DARK
            if "✓" in val:   color = RGBColor(0x1B, 0x5E, 0x20)
            if "✗" in val:   color = RGBColor(0xB7, 0x1C, 0x1C)
            cell_text(lapse_table.cell(r_idx, c_idx), val,
                      size=9, color=color, bold=(c_idx in [3,4]),
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_bg(lapse_table.cell(r_idx, c_idx), bg)

    add_paragraph(doc,
        "Note: Minimum shareholding — Simon retains his initial 10% regardless of milestone outcomes, "
        "subject only to bad-leaver buyback provisions. Unearned tranches revert to Davin's free float.",
        italic=True, color=GREY, size=9)

    hr(doc)

    # ── Milestone Deadlines Summary ───────────────────────────────────────
    add_heading(doc, "3.  Milestone Deadlines", level=2, gold=False)

    ml_table = doc.add_table(rows=4, cols=5)
    ml_table.style = "Table Grid"
    ml_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    ml_widths = [Cm(2.5), Cm(3.5), Cm(3.0), Cm(3.5), Cm(3.5)]
    for i, col in enumerate(ml_table.columns):
        for cell in col.cells:
            cell.width = ml_widths[i]

    ml_hdrs = ["Tranche", "AUM Target", "Equity Earned", "Deadline", "Trigger Date"]
    for i, h in enumerate(ml_hdrs):
        cell_text(ml_table.cell(0, i), h, bold=True, size=9,
                  color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_bg(ml_table.cell(0, i), GOLD_HEX)

    ml_rows = [
        ["Tranche 1", "R50,000,000 AUM",  "+10% (total 20%)", "24 months from signing", "~June 2028"],
        ["Tranche 2", "R100,000,000 AUM", "+10% (total 30%)", "48 months from signing", "~June 2030"],
        ["Tranche 3", "R200,000,000 AUM", "+10% (total 40%)", "72 months from signing", "~June 2032"],
    ]
    for r_idx, row in enumerate(ml_rows, 1):
        bg = LGREY_HEX if r_idx % 2 == 0 else "FFFFFF"
        aligns = [WD_ALIGN_PARAGRAPH.CENTER]*2 + [WD_ALIGN_PARAGRAPH.CENTER]*3
        for c_idx, val in enumerate(row):
            cell_text(ml_table.cell(r_idx, c_idx), val, size=10,
                      bold=(c_idx == 2), color=DARK,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_bg(ml_table.cell(r_idx, c_idx), bg)

    add_paragraph(doc,
        "AUM is verified by reference to the Company's monthly NAV report and the "
        "aggregate AUM figure reported in the BitWealth platform. The Board has final "
        "authority to confirm milestone achievement.",
        italic=True, color=GREY, size=9)

    hr(doc)

    # ── Assumptions ──────────────────────────────────────────────────────
    add_heading(doc, "4.  Key Assumptions & Notes", level=2, gold=False)
    bullets = [
        "All 1,000 issued ordinary shares carry equal economic and voting rights.",
        "Simon's founding 10% is issued as sweat equity (nil monetary consideration); "
        "the notional value at the IP-implied price of R2,500 per share is R250,000.",
        "Earn-in tranches represent a transfer of existing shares from Davin Cloete — "
        "no new shares will be issued. Total share count remains 1,000.",
        "AUM includes assets under management across all client portfolios managed "
        "by the Company via the Finova Capital FSP sub-licence.",
        "The Company operates under Finova Capital's FSP licence as a Juristic "
        "Representative. Simon may be appointed as a Key Individual or Representative "
        "once FAIS compliance requirements are met.",
        "The earn-in schedule is subject to the full terms set out in the Shareholders' "
        "Agreement and the Equity Participation Agreement, which will take precedence.",
        "This model is for discussion purposes only and does not constitute a binding "
        "legal agreement.",
    ]
    for b in bullets:
        add_bullet(doc, b)

# ═══════════════════════════════════════════════════════════════════════════
# PART B — TERM SHEET
# ═══════════════════════════════════════════════════════════════════════════
def build_term_sheet(doc):
    doc_header(
        doc,
        "BITWEALTH (PTY) LTD",
        "Part B — Non-Binding Heads of Terms — Business Development Equity Participation",
        TODAY,
    )

    add_paragraph(doc,
        "This term sheet sets out the proposed principal terms of an equity participation "
        "arrangement between Davin Cloete and Simon Hobday in respect of BitWealth (Pty) Ltd "
        "('the Company'). This document is non-binding, subject to contract, and intended "
        "for discussion purposes only. Final terms will be recorded in a Shareholders' "
        "Agreement and Equity Participation Agreement drafted by the Company's attorneys.")

    hr(doc)

    # ── Term sheet table helper ──────────────────────────────────────────
    def ts_section(title, rows, header_bg=TEAL_HEX):
        """Add a term sheet section with header and rows.
        rows = list of (term_label, detail_text)
        """
        add_heading(doc, title, level=1)

        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Column widths: 1/3 | 2/3
        for row in table.rows:
            row.cells[0].width = Cm(5.5)
            row.cells[1].width = Cm(11.0)

        for i, (label, detail) in enumerate(rows):
            bg = LGREY_HEX if i % 2 == 0 else "FFFFFF"
            cell_text(table.cell(i, 0), label, bold=True, size=10,
                      color=DARK, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_bg(table.cell(i, 0), bg)

            # Detail may contain newlines — split and add paragraphs
            cell = table.cell(i, 1)
            set_cell_bg(cell, bg)
            cell.text = ""
            paras = detail.split("\n")
            for p_idx, para_text in enumerate(paras):
                if p_idx == 0:
                    p = cell.paragraphs[0]
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(para_text.strip())
                run.font.size = Pt(10)
                run.font.color.rgb = DARK
                # Bold sub-headings that start with a letter in parentheses
                if para_text.strip().startswith(("(a)", "(b)", "(c)", "(d)", "Option A", "Option B", "Option C", "Option D")):
                    run.bold = True

        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── 1. PARTIES ────────────────────────────────────────────────────────
    ts_section("1.  Parties", [
        ("Founder",  "Davin Cloete, ID [●], Republic of South Africa."),
        ("Partner",  "Simon Hobday, ID [●], Republic of South Africa ('Simon')."),
        ("Company",  "BitWealth (Pty) Ltd (Registration No. [●] / 2025), a private company "
                     "incorporated in the Republic of South Africa, operating a Bitcoin "
                     "Dollar-Cost Averaging investment management platform."),
        ("FSP",      "Finova Capital (Pty) Ltd, FSP No. [●] — BitWealth will be appointed "
                     "as a Juristic Representative (JR) under Finova Capital's FSP licence "
                     "in terms of the FAIS Act, 37 of 2002."),
    ])

    # ── 2. COMPANY BACKGROUND ──────────────────────────────────────────────
    ts_section("2.  Company & IP Foundation", [
        ("Business",     "Bitcoin DCA investment management for retail and HNWI clients. "
                         "The Company's proprietary LTH PVR algorithm drives automated "
                         "daily buy/sell decisions executed via the VALR exchange."),
        ("IP Asset",     "The BitWealth LTH PVR technology platform has been transferred "
                         "to the Company by Davin Cloete pursuant to a signed IP Transfer "
                         "Agreement. Agreed fair value: R2,500,000 (Two Million Five Hundred "
                         "Thousand Rand), recorded on the Company's balance sheet."),
        ("Share Capital","1,000 ordinary shares of R1.00 par value each are in issue. "
                         "All shares carry equal voting, dividend, and liquidation rights. "
                         "Davin Cloete currently holds all 1,000 shares."),
        ("Valuation Basis","No external valuation has been obtained. The IP value reflects "
                           "the agreed commercial consideration for the Founder's contribution. "
                           "No institutional capital has been raised at this stage."),
    ])

    # ── 3. SIMON'S ROLE ────────────────────────────────────────────────────
    ts_section("3.  Simon's Role & Contribution", [
        ("Role Title",    "Head of Business Development"),
        ("Capital",       "NIL — Simon will contribute no monetary capital to the Company. "
                          "His equity participation is earned exclusively through performance "
                          "against the AUM milestones set out below."),
        ("Responsibilities",
         "Simon will be primarily responsible for:\n"
         "(a) Origination and onboarding of new clients onto the BitWealth platform;\n"
         "(b) Management and growth of existing client relationships;\n"
         "(c) Achievement of the AUM milestones set out in Clause 5;\n"
         "(d) Such other business development activities as the Board may agree from time to time."),
        ("FAIS Compliance",
         "Simon will be required to satisfy all applicable FAIS fitness and propriety "
         "requirements to operate as a Representative under Finova Capital's FSP licence. "
         "Failure to obtain or maintain the required regulatory approvals within 3 months "
         "of the signing date will constitute a material breach of this arrangement."),
        ("Time Commitment", "Simon is expected to devote sufficient time to the business to "
                            "achieve the milestones. The Board may, by resolution, agree to "
                            "a formal employment or consulting agreement on terms to be negotiated."),
    ])

    # ── 4. FOUNDING EQUITY GRANT ──────────────────────────────────────────
    ts_section("4.  Founding Equity Grant (Initial 10%)", [
        ("Shares Granted", "On the signing date, Davin Cloete will transfer 100 ordinary shares "
                           "(10% of issued share capital) to Simon Hobday for nil monetary consideration."),
        ("Rationale",      "The 10% founding equity reflects Simon's commitment to the business "
                           "development role and his appointment as an executive partner. "
                           "No performance condition attaches to the initial 10%."),
        ("Notional Value", "At the IP-implied price of R2,500 per share, the notional value "
                           "of the 10% founding grant is R250,000 (Two Hundred and Fifty Thousand Rand). "
                           "This does not represent a guaranteed or market-derived valuation."),
        ("Effective Date", "Transfer of the founding shares will be effected within 5 business "
                           "days of signature of the Shareholders' Agreement."),
        ("Conditions",     "(a) Execution of the Shareholders' Agreement by both parties.\n"
                           "(b) Simon's written acceptance of this term sheet.\n"
                           "(c) Confirmation of Simon's FAIS eligibility with Finova Capital."),
    ])

    # ── 5. EARN-IN SCHEDULE ───────────────────────────────────────────────
    ts_section("5.  Earn-In Milestones (Additional 30%)", [
        ("Mechanism",
         "Simon may earn up to an additional 30% of the Company in three equal tranches "
         "of 10% each (100 shares per tranche), contingent on achieving the AUM milestones "
         "set out below. Each tranche is a transfer of existing shares from Davin Cloete."),
        ("Tranche 1",
         "10% (100 shares) transferred to Simon upon the Company achieving verified total "
         "AUM of R50,000,000 (Fifty Million Rand) or more.\n"
         "Deadline: within 24 months of the Commencement Date (~June 2028)."),
        ("Tranche 2",
         "10% (100 shares) transferred to Simon upon the Company achieving verified total "
         "AUM of R100,000,000 (One Hundred Million Rand) or more.\n"
         "Deadline: within 48 months of the Commencement Date (~June 2030)."),
        ("Tranche 3",
         "10% (100 shares) transferred to Simon upon the Company achieving verified total "
         "AUM of R200,000,000 (Two Hundred Million Rand) or more.\n"
         "Deadline: within 72 months of the Commencement Date (~June 2032)."),
        ("AUM Definition",
         "AUM means the aggregate Rand value of assets under management across all active "
         "client portfolios managed by the Company, as reported in the BitWealth platform's "
         "NAV report at month-end. The Board confirms AUM achievement by resolution within "
         "10 business days of the relevant month-end."),
        ("Independence",
         "Each tranche is independent. Missing one milestone does not disqualify Simon from "
         "earning subsequent tranches, provided the later milestone is achieved within its "
         "own prescribed period. However, AUM milestones must be achieved in sequence — "
         "Tranche 3 cannot be triggered without Tranche 2 having been earned."),
        ("Maximum Earn-In", "If all three milestones are achieved: Davin Cloete 60%, Simon Hobday 40%."),
        ("Minimum Position",
         "Regardless of milestone outcome, Simon retains his founding 10% (subject to "
         "good/bad leaver provisions in Clause 8)."),
    ])

    # ── 6. MILESTONE FAILURE OPTIONS ─────────────────────────────────────
    ts_section("6.  Milestone Failure — Options for Negotiation", [
        ("Overview",
         "The parties acknowledge that milestones may not always be achieved within the "
         "prescribed deadlines. Four options are set out below for negotiation. The parties "
         "will agree which option(s) apply in the final Shareholders' Agreement."),

        ("OPTION A — Pure Lapse (Simplest)",
         "If the relevant AUM milestone is not achieved by the Deadline:\n"
         "(a) The unearned tranche automatically lapses with no further action required.\n"
         "(b) Simon retains all previously earned tranches.\n"
         "(c) The lapsed shares remain with Davin Cloete.\n"
         "(d) Simon has no right to compensation or further claims in respect of the lapsed tranche.\n"
         "ADVANTAGE: Clean, simple, no disputes about valuation.\n"
         "DISADVANTAGE: No flexibility for near-miss situations."),

        ("OPTION B — Company Call Option at Par",
         "If the relevant AUM milestone is not achieved by the Deadline:\n"
         "(a) The unearned tranche lapses automatically.\n"
         "(b) BitWealth (or Davin Cloete) receives a 90-day call option to purchase the "
         "unearned shares from Simon at nominal/par value (R1.00 per share).\n"
         "(c) If the option is exercised, the shares transfer to Davin at R1 per share.\n"
         "(d) If the option lapses unexercised, the shares remain with Simon as unvested equity "
         "with no further earn-in rights attaching.\n"
         "ADVANTAGE: Gives the Company control over the share register.\n"
         "DISADVANTAGE: Creates administration and potential dispute."),

        ("OPTION C — One-Time Extension with Penalty Reduction",
         "In respect of any ONE tranche per lifetime of this agreement:\n"
         "(a) The Board may (by unanimous resolution) grant a single 12-month extension.\n"
         "(b) The earn-in percentage for the extended tranche reduces from 10% to 7.5% "
         "(i.e., a 25% penalty for late achievement; Simon earns 75 shares instead of 100).\n"
         "(c) The remaining 25 shares are retained by Davin.\n"
         "(d) This option is available once only and cannot be applied to more than one tranche.\n"
         "ADVANTAGE: Rewards near-miss performance; maintains Simon's motivation.\n"
         "DISADVANTAGE: Complexity; potential disagreement on whether to grant extension."),

        ("OPTION D — Formula Buyback on Lapse",
         "If the relevant milestone lapses and the parties wish to provide Simon some "
         "compensation for AUM growth (even if milestone missed):\n"
         "(a) The lapsed tranche is bought back by BitWealth at a formula price.\n"
         "(b) Formula: (Total AUM at Deadline × 0.5%) ÷ 100 shares = price per share.\n"
         "    Example: AUM = R40m → price = (R40m × 0.5%) ÷ 100 = R2,000 per share "
         "(R200,000 for the tranche).\n"
         "(c) Payment deferred for 12 months from buy-back date.\n"
         "ADVANTAGE: Fair acknowledgement of AUM contribution even without full milestone.\n"
         "DISADVANTAGE: Cash outflow from Company; creates valuation arguments."),

        ("Recommendation",
         "Davin's preference is OPTION A (Pure Lapse) for Tranches 2 and 3, with "
         "OPTION C (One-time extension) available for Tranche 1 only, given the early "
         "stage of the business. This is subject to negotiation."),
    ])

    # ── 7. FSP ARRANGEMENT ───────────────────────────────────────────────
    ts_section("7.  FSP Arrangement — Finova Capital", [
        ("FSP Holder",       "Finova Capital (Pty) Ltd (FSP No. [●])."),
        ("JR Appointment",   "BitWealth (Pty) Ltd will be appointed as a Juristic Representative "
                             "under Finova Capital's FSP licence.\n"
                             "Terms: 20% of all performance fees + 50% of all platform fees "
                             "payable to Finova Capital as revenue share. No fixed monthly hosting fee."),
        ("Simon's FAIS Role","Simon may be appointed as a Representative under Finova Capital's "
                             "licence, subject to completion of all required FAIS fit-and-proper "
                             "requirements (RE1, RE5, relevant qualifications as required by Finova)."),
        ("CAEP Transition",  "The Company has decided NOT to proceed with CAEP Compliance as FSP host. "
                             "The Finova Capital arrangement is the Company's chosen regulatory structure. "
                             "Simon's compensation does not change based on FSP selection."),
        ("Regulatory Risk",  "Simon acknowledges that the Finova Capital JR arrangement is subject "
                             "to regulatory approval and ongoing compliance. Loss of JR status would "
                             "materially impact the Company's ability to operate."),
    ])

    # ── 8. PROTECTIVE PROVISIONS ──────────────────────────────────────────
    ts_section("8.  Protective Provisions", [
        ("Good Leaver Definition",
         "Simon is a 'Good Leaver' if his association with the Company ends by reason of:\n"
         "(a) Death or permanent total disability;\n"
         "(b) Retirement at or after age 65;\n"
         "(c) Termination by the Company without cause or material breach by the Company.\n"
         "Good Leaver consequence: Simon retains all earned tranches; unearned tranches lapse per Clause 6."),

        ("Bad Leaver Definition",
         "Simon is a 'Bad Leaver' if his association with the Company ends by reason of:\n"
         "(a) Voluntary resignation with less than 3 months' written notice;\n"
         "(b) Dismissal for misconduct, fraud, dishonesty, or material breach of fiduciary duty;\n"
         "(c) Any act that causes the Company to lose its Finova Capital JR status;\n"
         "(d) Competitive activity in breach of the non-compete clause below.\n"
         "Bad Leaver consequence: All shares (including founding 10%) subject to compulsory "
         "transfer to Davin Cloete at par value (R1.00 per share) within 30 days of the "
         "leaver event."),

        ("Non-Compete",
         "During the period Simon holds equity AND for 24 months after any transfer or "
         "forfeiture of his equity, Simon shall not, directly or indirectly:\n"
         "(a) Be employed by, consult for, or hold any interest in a competing Bitcoin "
         "DCA or cryptocurrency investment management business in South Africa;\n"
         "(b) Solicit any client of the Company to transfer their assets to a competitor;\n"
         "(c) Solicit any employee or contractor of the Company to leave their engagement.\n"
         "The non-compete radius is the Republic of South Africa."),

        ("Drag-Along",
         "If Davin Cloete (and/or his associates) receives a bona fide offer to acquire "
         "≥ 75% of the Company's shares, Simon is obliged to sell his pro-rata shares on "
         "the same terms and conditions as the Founder. This prevents Simon from blocking "
         "a sale that is supported by the majority shareholder."),

        ("Tag-Along",
         "If Davin Cloete proposes to sell more than 50% of his shares to a third party, "
         "Simon has the right (but not the obligation) to include his shares in the sale "
         "on equivalent terms and at the same price per share. This protects Simon from "
         "being left as a minority in a company with a new controlling shareholder."),

        ("Regulatory Clawback",
         "In the event that Simon's actions or omissions (including any act of misconduct, "
         "negligence, or regulatory breach in his capacity as a FAIS Representative) "
         "directly cause the Company to lose, suspend, or materially compromise its "
         "Finova Capital JR status, the Company may, by Board resolution, initiate a "
         "clawback of any earn-in tranches awarded within the preceding 24 months at "
         "par value (R1.00 per share). Davin Cloete's founding shares are not affected "
         "by this provision."),

        ("Anti-Dilution",
         "Simon's founding 10% and any earned tranche shares are protected against dilution "
         "arising from any future issuance of shares unless Simon provides written consent. "
         "If new shares are issued (e.g., to raise capital), Simon's anti-dilution rights "
         "entitle him to a pro-rata right of first offer to subscribe for new shares at "
         "the same price as the new investor."),
    ])

    # ── 9. SHARE TRANSFER MECHANISM ───────────────────────────────────────
    ts_section("9.  Share Transfer Mechanism", [
        ("Transfer Method",
         "All share transfers under this agreement will be effected by execution of a "
         "share transfer form and updating the Company's share register, filed with CIPC "
         "if required under the Companies Act, 71 of 2008."),
        ("Trigger Events",
         "A transfer is triggered upon:\n"
         "(a) Signing date (founding 10%), or\n"
         "(b) Board resolution confirming milestone achievement (earn-in tranches), or\n"
         "(c) Leaver event (as defined in Clause 8)."),
        ("Transfer Costs",
         "All costs associated with share transfers (including any securities transfer tax "
         "under the Securities Transfer Tax Act, 25 of 2007) will be borne by the Company."),
        ("Pre-emptive Rights",
         "Before any proposed transfer of shares to a third party, the other shareholder "
         "has a right of first refusal at the same price and terms as the proposed transfer."),
    ])

    # ── 10. CONDITIONS PRECEDENT ──────────────────────────────────────────
    ts_section("10.  Conditions Precedent", [
        ("CPs to Signing",
         "The equity participation arrangement is conditional on:\n"
         "(a) Execution of a Shareholders' Agreement in agreed form;\n"
         "(b) Execution of an Equity Participation Agreement in agreed form;\n"
         "(c) Simon confirming in writing that he has obtained independent legal advice;\n"
         "(d) Finova Capital confirming BitWealth's JR appointment or providing a letter "
         "of intent to appoint;\n"
         "(e) Confirmation that Simon meets the minimum FAIS fit-and-proper requirements "
         "applicable to his intended representative role."),
        ("Exclusivity",
         "From the date both parties sign this term sheet, Simon will not enter into "
         "discussions with any other fintech, crypto, or investment management company "
         "regarding a similar equity or business development arrangement for 60 days."),
    ])

    # ── 11. GENERAL ───────────────────────────────────────────────────────
    ts_section("11.  General", [
        ("Governing Law",    "Republic of South Africa. Disputes to be referred to arbitration "
                             "under the AFSA rules before any court proceedings."),
        ("Confidentiality",  "Both parties undertake to keep the contents of this term sheet "
                             "and all related discussions strictly confidential."),
        ("Attorneys",        "The Company will appoint attorneys to draft the definitive agreements. "
                             "Simon is advised to obtain independent legal advice at his own cost."),
        ("Non-Binding",      "This document is non-binding in its entirety and constitutes "
                             "heads of terms for negotiation only. No binding obligations "
                             "arise until definitive agreements are executed by both parties."),
        ("Validity",         "This term sheet lapses if not countersigned by both parties "
                             "within 30 days of the date hereof."),
    ])

    hr(doc)

    # ── Signature block ───────────────────────────────────────────────────
    add_heading(doc, "12.  Signature", level=1)

    add_paragraph(doc,
        "By signing below, the parties confirm that they have read and understood the "
        "above heads of terms and agree to proceed with the drafting of definitive "
        "agreements on this basis.  This signature does not create a binding legal obligation.")

    sig_table = doc.add_table(rows=6, cols=2)
    sig_table.style = "Table Grid"
    for row in sig_table.rows:
        row.cells[0].width = Cm(8)
        row.cells[1].width = Cm(8)

    sig_rows = [
        ("FOR AND ON BEHALF OF BITWEALTH (PTY) LTD\nand in his personal capacity as Founder",
         "AGREED TO BY SIMON HOBDAY\nin his personal capacity"),
        ("Signature: _____________________________",  "Signature: _____________________________"),
        ("Full Name:  Davin Cloete",                  "Full Name:  Simon Hobday"),
        ("Capacity:   Director / Founder",            "Capacity:   Proposed Business Development Partner"),
        ("Date:       ___________________________",   "Date:       ___________________________"),
        ("Witness:    ___________________________",   "Witness:    ___________________________"),
    ]
    for r_idx, (left, right) in enumerate(sig_rows):
        bg = LGREY_HEX if r_idx == 0 else "FFFFFF"
        for c_idx, val in enumerate([left, right]):
            cell_text(sig_table.cell(r_idx, c_idx), val,
                      bold=(r_idx == 0), size=10, color=DARK,
                      align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_bg(sig_table.cell(r_idx, c_idx), bg if r_idx == 0 else "FFFFFF")

# ═══════════════════════════════════════════════════════════════════════════
# COMBINED DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

build_cap_table(doc)
doc.add_page_break()
build_term_sheet(doc)

out_path = OUT_DIR / "BitWealth_Simon_Equity_Proposal_v2.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
