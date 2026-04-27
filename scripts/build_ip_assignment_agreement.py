"""Generate BitWealth_IP_Assignment_Agreement.docx — draft for attorney review."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GOLD = RGBColor(0xC9, 0xA2, 0x27)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x00, 0x00)
LIGHT_GOLD_HEX = "F5EBC8"
LIGHT_RED_HEX = "FBE9E7"

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)


def add_para(text, bold=False, italic=False, size=11, color=DARK, align=None, space_after=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_clause(num, title, body_paragraphs):
    """Add a numbered clause with title and one or more body paragraphs."""
    h = doc.add_paragraph()
    hr = h.add_run(f"{num}.  {title.upper()}")
    hr.bold = True
    hr.font.size = Pt(12)
    hr.font.color.rgb = GOLD
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(4)
    for body in body_paragraphs:
        if isinstance(body, tuple):
            sub_num, sub_text = body
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"{sub_num}  ")
            r.bold = True
            r.font.size = Pt(11)
            r2 = p.add_run(sub_text)
            r2.font.size = Pt(11)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(body)
            r.font.size = Pt(11)


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_table(headers, rows, header_fill=LIGHT_GOLD_HEX, col_widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        shade_cell(hdr[i], header_fill)
    for row_data in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = ''
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return tbl


# ============ COVER ============
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("INTELLECTUAL PROPERTY ASSIGNMENT AGREEMENT")
tr.bold = True; tr.font.size = Pt(20); tr.font.color.rgb = GOLD

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("(Founder to Company)")
sr.font.size = Pt(13); sr.italic = True; sr.font.color.rgb = DARK

# Disclaimer banner
doc.add_paragraph()
disc_tbl = doc.add_table(rows=1, cols=1)
disc_cell = disc_tbl.rows[0].cells[0]
shade_cell(disc_cell, LIGHT_RED_HEX)
disc_cell.text = ''
dp = disc_cell.paragraphs[0]
dr1 = dp.add_run("DRAFT FOR ATTORNEY REVIEW — NOT LEGALLY BINDING\n")
dr1.bold = True; dr1.font.size = Pt(11); dr1.font.color.rgb = RED
dr2 = dp.add_run(
    "This document is a starting framework prepared for review and finalisation by a "
    "qualified attorney admitted to practise in the Republic of South Africa. It is not legal "
    "advice. Do not sign or rely on this document until it has been reviewed, amended where "
    "necessary, and approved by your attorney."
)
dr2.italic = True; dr2.font.size = Pt(10); dr2.font.color.rgb = RED

doc.add_paragraph()

# ============ PARTIES ============
add_para("ENTERED INTO BY AND BETWEEN:", bold=True, size=11)
add_para("[FULL NAME OF FOUNDER]", bold=True)
add_para('Identity Number: [ID NUMBER]')
add_para('Residential Address: [PHYSICAL ADDRESS]')
add_para('(hereinafter referred to as "the Assignor" or "the Founder")')

add_para("AND", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

add_para("BITWEALTH (PTY) LTD", bold=True)
add_para('Registration Number: [CIPC REGISTRATION NUMBER]')
add_para('Registered Address: [REGISTERED OFFICE ADDRESS]')
add_para('Herein duly represented by [NAME], in his capacity as Director, '
         'duly authorised hereto')
add_para('(hereinafter referred to as "the Assignee" or "the Company")')

add_para("(collectively referred to as \"the Parties\")", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)

# ============ RECITALS ============
add_para("")
hr_p = doc.add_paragraph()
hrr = hr_p.add_run("RECITALS")
hrr.bold = True; hrr.font.size = Pt(12); hrr.font.color.rgb = GOLD
hr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para(
    "WHEREAS the Founder, in his personal capacity and prior to the incorporation of the "
    "Company, conceived, developed, and continues to develop certain proprietary intellectual "
    "property relating to a Bitcoin Dollar-Cost Averaging investment strategy known as the "
    '"Long-Term Holder Price-to-Value Ratio" or "LTH PVR" strategy, including all related '
    "software, algorithms, data models, technical infrastructure, brand assets, and "
    "documentation;"
)
add_para(
    "AND WHEREAS the Company was incorporated on [DATE OF INCORPORATION] for the purpose of "
    "commercialising the LTH PVR strategy and rendering related financial services in the "
    "Republic of South Africa under an appropriate Financial Services Provider licence;"
)
add_para(
    "AND WHEREAS the Founder wishes to assign, transfer, and convey to the Company all of his "
    "right, title, and interest in and to the said intellectual property, in order that the "
    "Company shall hold such intellectual property as its own asset for all commercial, "
    "regulatory, and operational purposes;"
)
add_para(
    "AND WHEREAS the Company wishes to accept such assignment on the terms and conditions set "
    "out in this Agreement;"
)
add_para(
    "NOW THEREFORE the Parties agree as follows:", bold=True
)

# ============ CLAUSES ============
add_clause("1", "Definitions and Interpretation", [
    'In this Agreement, unless the context indicates otherwise, the following terms shall have the meanings assigned to them:',
    ('1.1', '"Agreement" means this Intellectual Property Assignment Agreement, including all annexures hereto.'),
    ('1.2', '"Assigned IP" means all of the Intellectual Property described in clause 2 and Annexure A, including all rights, title, and interest therein, whether registered or unregistered, existing or future, and wherever in the world such rights may subsist.'),
    ('1.3', '"Effective Date" means the date of last signature of this Agreement.'),
    ('1.4', '"Intellectual Property" or "IP" means all intellectual property of any nature whatsoever, including without limitation: copyright (including software source and object code, documentation, designs, and written works); patents and patentable inventions; trade marks (registered and unregistered), trade names, service marks, and logos; domain names; designs (registered and unregistered); database rights; trade secrets; know-how; confidential information; algorithms and methodologies; data models; technical specifications; and any goodwill associated with any of the foregoing.'),
    ('1.5', '"LTH PVR Strategy" means the Long-Term Holder Price-to-Value Ratio investment methodology developed by the Founder, comprising signal logic, sigma band calculations, sizing rules, sell tier logic, and bear-market pause logic, as more fully described in Annexure A.'),
    ('1.6', '"Software" means all source code, object code, scripts, configuration files, database schemas, edge functions, deployment scripts, and related technical materials forming part of the Assigned IP, as listed in Annexure B.'),
    ('1.7', '"Third-Party IP" means any intellectual property owned by a third party which forms part of, is incorporated into, or is used in connection with the Assigned IP, including open-source software components, third-party APIs, and licensed data feeds.'),
    ('1.8', 'Headings are for convenience only and shall not be used in the interpretation of this Agreement.'),
    ('1.9', 'Words importing the singular include the plural and vice versa, and words importing one gender include all genders.'),
])

add_clause("2", "Description of Assigned IP", [
    'The Assigned IP includes, without limitation, the following categories of intellectual property created, developed, or acquired by the Founder prior to the Effective Date and relating to the LTH PVR Strategy and the business of the Company:',
    ('2.1', 'The LTH PVR Strategy itself, including all signal logic, decision algorithms, sigma band methodology, position sizing rules (B1–B5 buy tiers and S6–S11 sell tiers), and bear-market pause logic.'),
    ('2.2', 'All Software comprising the LTH PVR execution platform, including but not limited to: the Supabase database schema and all migrations; all Deno edge functions (including but not limited to ef_fetch_ci_bands, ef_generate_decisions, ef_create_order_intents, ef_execute_orders, ef_poll_orders, ef_post_ledger_and_balances, ef_alert_digest, ef_valr_ws_monitor, and ef_resume_pipeline); the administrative and client-facing user interfaces; the back-testing engine; and all associated PowerShell deployment scripts.'),
    ('2.3', 'All technical and product documentation, including the Solution Design Document (SDD), test case documents, deployment guides, and architectural specifications.'),
    ('2.4', 'All brand assets, including the "BitWealth" name, logos, trade dress, marketing collateral, the ABC One-Pager, and any other client-facing materials.'),
    ('2.5', 'All domain names, social media handles, exchange-account integrations, and third-party platform accounts (including but not limited to GitHub, Supabase, CryptoQuant, VALR, Resend, and Netlify) used in connection with the LTH PVR Strategy and the business of the Company.'),
    ('2.6', 'All historical back-test data, model outputs, performance reports, and analytical work product produced in the course of developing the LTH PVR Strategy.'),
    ('2.7', 'A complete inventory of the Assigned IP is set out in Annexure A. To the extent that any item of intellectual property created by the Founder prior to the Effective Date and relating to the business of the Company is not specifically listed in Annexure A, such item shall nevertheless form part of the Assigned IP and shall be deemed assigned to the Company under this Agreement.'),
])

add_clause("3", "Assignment", [
    ('3.1', 'With effect from the Effective Date, the Founder hereby irrevocably and unconditionally assigns, transfers, and conveys to the Company, free of any encumbrance, all of his right, title, and interest in and to the Assigned IP, together with all goodwill associated therewith, for the full term of any rights subsisting therein and any extensions or renewals thereof, in all jurisdictions worldwide.'),
    ('3.2', 'The assignment in clause 3.1 includes the right to sue for and recover damages, profits, and any other remedies in respect of any infringement of the Assigned IP occurring before, on, or after the Effective Date.'),
    ('3.3', 'The Founder hereby waives, to the maximum extent permitted by law, all moral rights he may have in any copyright works forming part of the Assigned IP.'),
    ('3.4', 'The Company accepts the assignment on the terms set out in this Agreement.'),
])

add_clause("4", "Consideration", [
    ('4.1', 'In consideration for the assignment effected by this Agreement, the Company shall pay or provide to the Founder the consideration set out in this clause 4.'),
    ('4.2', '[OPTION A — NOMINAL CONSIDERATION] The Company shall pay to the Founder a nominal cash consideration of R 100.00 (One Hundred Rand), the receipt of which is hereby acknowledged.'),
    ('4.3', '[OPTION B — SHARE-BASED CONSIDERATION] The Company shall issue to the Founder [NUMBER] ordinary shares in the share capital of the Company, credited as fully paid up, on the Effective Date.'),
    ('4.4', '[OPTION C — DEFERRED / LOAN ACCOUNT] The Company shall credit the Founder\'s shareholder loan account in the amount of R [AMOUNT], being the agreed value of the Assigned IP as at the Effective Date, repayable on terms to be agreed between the Parties from time to time.'),
    ('4.5', 'The Parties acknowledge that the consideration set out above represents fair and adequate value for the assignment, having regard to the Founder\'s ongoing role as a shareholder, director, and beneficiary of the Company\'s commercialisation of the Assigned IP.'),
    '[NOTE TO ATTORNEY: Select ONE of Options A, B, or C above based on the Founder\'s tax position and shareholder structure. The choice has material capital gains tax (CGT), VAT, and securities transfer tax implications. Refer to a tax practitioner.]',
])

add_clause("5", "Founder's Warranties", [
    'The Founder warrants and represents to the Company that, as at the Effective Date:',
    ('5.1', 'He is the sole legal and beneficial owner of the Assigned IP, and has full power, authority, and right to assign the Assigned IP to the Company free of any encumbrance.'),
    ('5.2', 'The Assigned IP is original to the Founder and, to the best of his knowledge after reasonable enquiry, does not infringe the intellectual property rights of any third party, save for the Third-Party IP identified in Annexure C.'),
    ('5.3', 'No third party has any claim, right, licence, option, or interest in or to the Assigned IP other than as disclosed in Annexure C.'),
    ('5.4', 'No part of the Assigned IP was created during the course of any employment, consultancy, or contractor engagement with any third party that would entitle such third party to claim ownership of, or rights in, the Assigned IP.'),
    ('5.5', 'There are no pending or threatened legal proceedings, disputes, or claims relating to the Assigned IP.'),
    ('5.6', 'All open-source software components incorporated into the Software have been used in compliance with their respective licence terms, and a list of such components is set out in Annexure C.'),
])

add_clause("6", "Third-Party IP and Open-Source Components", [
    ('6.1', 'The Parties acknowledge that the Software incorporates certain Third-Party IP, including open-source libraries and third-party API integrations (including without limitation CryptoQuant, VALR, Supabase, and Resend), as listed in Annexure C.'),
    ('6.2', 'The Founder confirms that all Third-Party IP has been used in accordance with its applicable licence terms, and the Company shall continue to comply with such terms following the Effective Date.'),
    ('6.3', 'No Third-Party IP is assigned under this Agreement; the Company shall obtain its own licences to use such Third-Party IP from the relevant third parties on an ongoing basis.'),
])

add_clause("7", "Future IP and Further Assurance", [
    ('7.1', 'In respect of any Intellectual Property created, developed, or acquired by the Founder after the Effective Date in the course of his role as a director, employee, or shareholder of the Company, and which relates to the business of the Company, such Intellectual Property shall vest in and be the sole property of the Company on creation, without the need for any further act of assignment.'),
    ('7.2', 'The Founder undertakes, at the Company\'s reasonable request and at the Company\'s cost, to execute any further documents and to take any further actions reasonably necessary to give full effect to this Agreement and to perfect the Company\'s title to the Assigned IP, including without limitation: signing applications for the registration of any registrable intellectual property; transferring control of domain names, hosting accounts, and third-party platform accounts; and providing all source code, credentials, encryption keys, and technical documentation.'),
    ('7.3', 'The Founder hereby appoints the Company as his attorney and agent to execute any document and take any action contemplated by clause 7.2 in the event that the Founder fails or refuses to do so within a reasonable time after request.'),
])

add_clause("8", "Confidentiality", [
    ('8.1', 'The Founder shall keep all information relating to the Assigned IP, including without limitation the Software source code, the LTH PVR signal logic, client data, and the Company\'s business plans and strategy, strictly confidential, and shall not disclose such information to any third party without the prior written consent of the Company.'),
    ('8.2', 'This obligation of confidentiality shall survive the termination of any employment, directorship, or shareholder relationship between the Founder and the Company.'),
    ('8.3', 'The obligation of confidentiality shall not apply to information which is or becomes publicly available other than through a breach of this Agreement, or which the Founder is required to disclose by law or by order of a competent court.'),
])

add_clause("9", "Restraint of Trade", [
    ('9.1', 'The Founder shall not, for a period of [12 / 24] months after the date on which he ceases to be a director, employee, or shareholder of the Company, directly or indirectly, in the Republic of South Africa, carry on or be engaged or interested in any business which competes with the business of the Company in respect of automated, signal-driven Bitcoin or crypto asset accumulation strategies for retail or institutional clients.'),
    ('9.2', 'The Parties acknowledge that this restraint is reasonable in scope, duration, and geography, having regard to the value of the Assigned IP and the legitimate proprietary interests of the Company.'),
    '[NOTE TO ATTORNEY: Restraints of trade in South Africa are enforceable provided they are reasonable. The duration, geographic scope, and scope of activity should be tailored to be defensible. Consult a labour-law specialist if necessary.]',
])

add_clause("10", "Indemnity", [
    ('10.1', 'The Founder indemnifies and holds the Company harmless against any loss, damage, claim, or expense (including reasonable legal costs) suffered or incurred by the Company arising out of or in connection with any breach of the warranties given by the Founder in clause 5.'),
    ('10.2', 'The Founder\'s liability under this indemnity shall be capped at the value of the consideration received under clause 4, save in the case of fraud or wilful misconduct.'),
])

add_clause("11", "Tax", [
    ('11.1', 'Each Party shall be responsible for its own tax liabilities arising from this Agreement, including without limitation any capital gains tax, income tax, value-added tax, or securities transfer tax.'),
    ('11.2', 'The Parties shall co-operate in good faith in obtaining any tax clearances or making any tax filings required as a consequence of this Agreement.'),
    '[NOTE TO ATTORNEY/TAX PRACTITIONER: The assignment of IP from a natural person to a company is a CGT event. A tax opinion should be obtained, and consideration given to the use of section 42 of the Income Tax Act (asset-for-share transaction relief) where shares are issued as consideration.]',
])

add_clause("12", "General", [
    ('12.1', 'Entire Agreement. This Agreement constitutes the entire agreement between the Parties in respect of its subject matter and supersedes all prior negotiations, understandings, and agreements.'),
    ('12.2', 'Amendment. No amendment to this Agreement shall be of any force or effect unless reduced to writing and signed by both Parties.'),
    ('12.3', 'Severability. If any provision of this Agreement is held to be invalid or unenforceable, the remaining provisions shall continue in full force and effect.'),
    ('12.4', 'Governing Law. This Agreement shall be governed by and construed in accordance with the laws of the Republic of South Africa.'),
    ('12.5', 'Jurisdiction. The Parties submit to the non-exclusive jurisdiction of the High Court of South Africa, [Gauteng Division, Johannesburg].'),
    ('12.6', 'Dispute Resolution. Any dispute arising out of or in connection with this Agreement shall first be referred to the Parties\' respective representatives for good-faith negotiation, failing which the dispute shall be referred to mediation, and failing which to arbitration in accordance with the rules of the Arbitration Foundation of Southern Africa (AFSA).'),
    ('12.7', 'Notices. All notices under this Agreement shall be in writing and delivered to the addresses set out at the head of this Agreement, or such other address as a Party may from time to time notify the other Party in writing.'),
    ('12.8', 'Costs. Each Party shall bear its own legal costs in respect of the negotiation and conclusion of this Agreement, save that the Company shall bear the costs of any registrations or filings required to perfect title to the Assigned IP.'),
    ('12.9', 'Counterparts. This Agreement may be signed in counterparts, each of which shall constitute an original and which together shall constitute one and the same agreement. Electronic signatures shall be valid and binding.'),
])

# ============ SIGNATURES ============
doc.add_paragraph()
sig_h = doc.add_paragraph()
sigr = sig_h.add_run("SIGNED BY THE PARTIES:")
sigr.bold = True; sigr.font.size = Pt(12); sigr.font.color.rgb = GOLD
sig_h.paragraph_format.space_before = Pt(12)

# Founder signature block
add_para("")
add_para("FOR THE FOUNDER:", bold=True)
add_para("")
add_para("_________________________________________")
add_para("[FULL NAME OF FOUNDER]")
add_para("Date:  ____ / ____ / 2026")
add_para("Place: _________________________________")
add_para("Witness 1 Signature: ___________________________   Name: __________________________")
add_para("Witness 2 Signature: ___________________________   Name: __________________________")

# Company signature block
add_para("")
add_para("FOR AND ON BEHALF OF BITWEALTH (PTY) LTD:", bold=True)
add_para("")
add_para("_________________________________________")
add_para("Name: [DIRECTOR NAME]")
add_para("Designation: Director")
add_para("Duly authorised hereto")
add_para("Date:  ____ / ____ / 2026")
add_para("Place: _________________________________")
add_para("Witness 1 Signature: ___________________________   Name: __________________________")
add_para("Witness 2 Signature: ___________________________   Name: __________________________")

# ============ ANNEXURES ============
doc.add_page_break()
ann_h = doc.add_paragraph()
ar = ann_h.add_run("ANNEXURE A — INVENTORY OF ASSIGNED IP")
ar.bold = True; ar.font.size = Pt(14); ar.font.color.rgb = GOLD

add_para(
    "The following is a non-exhaustive inventory of the Assigned IP. To the extent that any "
    "item of intellectual property is omitted from this list but otherwise falls within the "
    "definition of Assigned IP, it shall nevertheless be assigned under this Agreement.",
    italic=True, color=GREY
)

add_table(
    headers=["Category", "Description", "Location / Reference"],
    rows=[
        ["Strategy IP", "LTH PVR signal logic, sigma band methodology, B1–B5 / S6–S11 sizing rules", "Documented in SDD_v0.6.md"],
        ["Source Code — Backend", "All Supabase edge functions (Deno/TypeScript)", "GitHub repo: [URL]"],
        ["Source Code — Database", "All PostgreSQL schemas, migrations, RPC functions", "supabase/migrations/"],
        ["Source Code — UI", "Admin and client-facing HTML/JS interfaces", "ui/ and docs/"],
        ["Source Code — Backtest Engine", "Back-testing simulator and reporting", "supabase/functions/_shared/"],
        ["Documentation", "Solution Design Document, test cases, deployment guides", "docs/"],
        ["Brand Assets", "BitWealth name, logos, marketing collateral, ABC One-Pager", "docs/ABC_One_Pager.html and brand files"],
        ["Domain Names", "[LIST ALL DOMAINS, e.g. bitwealth.co.za]", "Registrar: [REGISTRAR]"],
        ["Third-Party Accounts", "GitHub, Supabase, CryptoQuant, VALR, Resend, Netlify", "[ACCOUNT EMAIL/HANDLE]"],
        ["Historical Data", "Back-test outputs, performance reports, model outputs", "Stored in [LOCATION]"],
    ],
    col_widths=[3.5, 7.5, 5.5]
)

doc.add_page_break()
ann_b = doc.add_paragraph()
br = ann_b.add_run("ANNEXURE B — SOFTWARE INVENTORY")
br.bold = True; br.font.size = Pt(14); br.font.color.rgb = GOLD

add_para(
    "Detailed list of edge functions, scripts, and software components forming part of the Assigned IP.",
    italic=True, color=GREY
)

add_table(
    headers=["Component", "Type", "Path / Location"],
    rows=[
        ["ef_fetch_ci_bands", "Edge Function", "supabase/functions/ef_fetch_ci_bands/"],
        ["ef_generate_decisions", "Edge Function", "supabase/functions/ef_generate_decisions/"],
        ["ef_create_order_intents", "Edge Function", "supabase/functions/ef_create_order_intents/"],
        ["ef_execute_orders", "Edge Function", "supabase/functions/ef_execute_orders/"],
        ["ef_poll_orders", "Edge Function", "supabase/functions/ef_poll_orders/"],
        ["ef_post_ledger_and_balances", "Edge Function", "supabase/functions/ef_post_ledger_and_balances/"],
        ["ef_alert_digest", "Edge Function", "supabase/functions/ef_alert_digest/"],
        ["ef_valr_ws_monitor", "Edge Function", "supabase/functions/ef_valr_ws_monitor/"],
        ["ef_resume_pipeline", "Edge Function", "supabase/functions/ef_resume_pipeline/"],
        ["ef_run_lth_pvr_simulator", "Edge Function (Backtest)", "supabase/functions/ef_run_lth_pvr_simulator/"],
        ["lth_pvr_simulator.ts", "Shared Module", "supabase/functions/_shared/lth_pvr_simulator.ts"],
        ["Database schemas", "PostgreSQL", "public, lth_pvr, lth_pvr_bt"],
        ["Migrations", "SQL", "supabase/migrations/"],
        ["Admin UI", "HTML/JS", "ui/Advanced BTC DCA Strategy.html"],
        ["Deployment scripts", "PowerShell", "deploy-*.ps1, redeploy-all-functions.ps1"],
    ],
    col_widths=[6, 4.5, 6]
)

doc.add_page_break()
ann_c = doc.add_paragraph()
cr = ann_c.add_run("ANNEXURE C — THIRD-PARTY IP AND OPEN-SOURCE COMPONENTS")
cr.bold = True; cr.font.size = Pt(14); cr.font.color.rgb = GOLD

add_para(
    "Third-party intellectual property and open-source components incorporated into or used by the Assigned IP. "
    "Such Third-Party IP is NOT assigned under this Agreement; the Company shall hold its own licences.",
    italic=True, color=GREY
)

add_table(
    headers=["Component", "Type", "Licence", "Notes"],
    rows=[
        ["Supabase", "Backend platform / DB host", "Commercial subscription", "Hosted PostgreSQL + Edge Functions"],
        ["CryptoQuant API", "Third-party data feed", "Commercial subscription", "Source of LTH PVR signal data"],
        ["VALR API", "Exchange integration", "Exchange T&Cs", "Order execution and account data"],
        ["Resend", "Email API", "Commercial subscription", "Alert digest delivery"],
        ["Netlify", "Static site hosting", "Commercial subscription", "Marketing site hosting"],
        ["Deno standard library", "Runtime", "MIT", "Edge function runtime"],
        ["Supabase JS Client", "Library", "MIT", "Frontend integration"],
        ["Tailwind CSS", "Library", "MIT", "UI styling"],
        ["Chart.js", "Library", "MIT", "Charts in one-pager"],
        ["[ANY OTHER LIBRARIES]", "[TYPE]", "[LICENCE]", "[NOTES]"],
    ],
    col_widths=[4, 3.5, 3.5, 5.5]
)

# ============ NOTES TO ATTORNEY ============
doc.add_page_break()
notes_h = doc.add_paragraph()
nr = notes_h.add_run("NOTES TO THE REVIEWING ATTORNEY")
nr.bold = True; nr.font.size = Pt(14); nr.font.color.rgb = RED

note_items = [
    "This is a draft framework. It is intentionally comprehensive but must be tailored to the Founder's specific circumstances, tax position, and the company's intended capital structure.",
    "Clause 4 (Consideration) — three options are presented. The choice has material CGT, VAT, and STT implications. A tax opinion is recommended before execution. If an asset-for-share transaction under section 42 of the Income Tax Act is contemplated, additional documentation and SARS clearance may be required.",
    "Clause 9 (Restraint of Trade) — duration and scope must be reasonable to be enforceable in South African law. Consider the Founder's likely future activities and the legitimate interests of the Company. A separate restraint may also be appropriate in any future shareholders agreement.",
    "Clause 5 (Warranties) — these should be reviewed against the actual development history of the Software. If any third-party developers, contractors, or consultants contributed to any part of the codebase, additional assignment agreements from those parties may be required before this Agreement can be properly given.",
    "Annexures A, B, and C must be completed accurately. Any IP omitted from these annexures may give rise to ownership disputes later. Particular care should be taken with: (a) domain name registrations; (b) third-party platform account ownership (especially GitHub, Supabase); (c) any code that may have been written before BitWealth (Pty) Ltd was incorporated.",
    "If RocketX, Simon, or any other party will become shareholders of the Company, this Agreement should be signed BEFORE any such shareholding is granted, so that those parties subscribe for shares in a company that already owns the IP.",
    "Witnessing and place of signature should comply with the Electronic Communications and Transactions Act (ECTA) if signed electronically.",
    "Consider whether any moral rights or performer's rights subsist that require separate waiver under the Copyright Act, 1978 (as amended).",
    "If the Company intends to register any patent, trade mark, or design rights, separate registration applications will be required. This Agreement gives the Company the right to make those applications but does not effect the registrations themselves.",
]
for i, item in enumerate(note_items, 1):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(item)
    r.font.size = Pt(10)

# Save
out_path = r"docs\BitWealth_IP_Assignment_Agreement.docx"
doc.save(out_path)
print(f"Document saved: {out_path}")
