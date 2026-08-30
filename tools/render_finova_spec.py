"""Render docs/FINOVA_PARTNER_PORTAL_SPEC.md to a Word document for Finova.

The markdown file is the source of truth; re-run this after editing it:
    .venv\\Scripts\\python.exe tools/render_finova_spec.py

Supports the subset of markdown used by the spec: ATX headings, bullet lists,
ordered lists, pipe tables, blockquotes, `code`, **bold**, and *italics*.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parent.parent
SOURCE = REPO_ROOT / "docs" / "FINOVA_PARTNER_PORTAL_SPEC.md"
TARGET = REPO_ROOT / "docs" / "FINOVA_PARTNER_PORTAL_SPEC.docx"

BRAND_NAVY = RGBColor(0x0A, 0x2E, 0x4D)
CODE_GREY = RGBColor(0x44, 0x44, 0x44)

# **bold**, *italic*, `code` — kept in one pattern so we can walk matches in order.
INLINE_RE = re.compile(r"(\*\*.+?\*\*|(?<!\*)\*[^*]+?\*(?!\*)|`[^`]+?`)")


def add_inline(paragraph, text: str) -> None:
    """Append text to a paragraph, honouring bold / italic / code spans."""
    for token in INLINE_RE.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            paragraph.add_run(token[2:-2]).bold = True
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = CODE_GREY
        elif token.startswith("*") and token.endswith("*"):
            paragraph.add_run(token[1:-1]).italic = True
        else:
            paragraph.add_run(token)


def split_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_divider(line: str) -> bool:
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    header, *body = rows
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, header):
        cell.paragraphs[0].text = ""
        add_inline(cell.paragraphs[0], text)
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in body:
        cells = table.add_row().cells
        # Ragged rows would raise on zip-strict; pad instead so nothing is dropped.
        for cell, text in zip(cells, row + [""] * (len(header) - len(row))):
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], text)
    doc.add_paragraph()


def build(markdown: str) -> Document:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if not stripped:
            index += 1
            continue

        # Internal maintenance note — not for Finova.
        if stripped.startswith(">"):
            index += 1
            continue

        table_match = stripped.startswith("|") and index + 1 < len(lines) and is_divider(lines[index + 1])
        if table_match:
            rows = [split_row(stripped)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                rows.append(split_row(lines[index]))
                index += 1
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if heading:
            level = len(heading.group(1))
            para = doc.add_heading(level=min(level, 4))
            add_inline(para, heading.group(2))
            for run in para.runs:
                run.font.color.rgb = BRAND_NAVY
            index += 1
            continue

        bullet = re.match(r"^[-*]\s+(.*)$", stripped)
        if bullet:
            add_inline(doc.add_paragraph(style="List Bullet"), bullet.group(1))
            index += 1
            continue

        ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered:
            add_inline(doc.add_paragraph(style="List Number"), ordered.group(1))
            index += 1
            continue

        para = doc.add_paragraph()
        add_inline(para, stripped)
        if stripped.startswith("*") and stripped.endswith("*"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        index += 1

    return doc


def main() -> int:
    if not SOURCE.exists():
        print(f"Source not found: {SOURCE}", file=sys.stderr)
        return 1
    build(SOURCE.read_text(encoding="utf-8")).save(TARGET)
    print(f"Wrote {TARGET}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
